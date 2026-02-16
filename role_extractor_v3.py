"""
Role Extractor for Engineering Work Instructions
Version 3.5.0 - Production-Ready for Air-Gapped Networks

v3.5.0 Changes:
- Fixed single-word variant matching bug (e.g., "Validation" no longer verifies "Validation Engineer")
- Added ORGANIZATION_ENTITIES filter (NASA, Government, etc. no longer extracted as roles)
- Added 30+ missing aerospace/NASA roles to KNOWN_ROLES
- Expanded SINGLE_WORD_EXCLUSIONS with 50+ common English words
- Added post-verification filters: org entity removal, confidence threshold (0.4), stopword filter
- Removed 'nasa' and 'government' from KNOWN_ROLES, removed duplicate entries
- STRICT mode also skips organization entities

v3.4.0 Changes:
- Added ExtractionMode enum (STRICT vs DISCOVERY)
- STRICT mode: Whitelist-only extraction with 100% accuracy guarantee
- STRICT mode only extracts roles from KNOWN_ROLES that appear verbatim in text
- Each occurrence in STRICT mode is linked to exact character position
- Updated __init__ to accept extraction_mode parameter

v3.2.5 Changes:
- Added phone number and numeric pattern filtering in _clean_candidate()
- Added length validation (max 60 chars) in _is_valid_role()
- Added ZIP code pattern filtering
- Added FAA/Aviation-specific roles (accountable executive, certificate holder, etc.)
- Added OSHA/Safety-specific roles (process safety coordinator, plant manager, etc.)

Extracts organizational roles and their associated responsibilities from
engineering documents (Word docs, PDFs, text files) using pattern matching
and linguistic rules. No external AI/API dependencies.

Dependencies (commonly available on closed networks):
- re (standard library)
- collections (standard library)  
- dataclasses (standard library, Python 3.7+)
- python-docx (pip install python-docx) - for Word documents
- PyPDF2 or pdfplumber (pip install PyPDF2 pdfplumber) - for PDFs

Author: Nick / SAIC Systems Engineering
For use with: Technical Review Tool
"""

import re
from collections import defaultdict
from dataclasses import dataclass, field
from typing import List, Dict, Set, Tuple, Optional
import os
import csv
from enum import Enum

# Structured logging support
try:
    from config_logging import get_logger
    _logger = get_logger('role_extractor')
except ImportError:
    _logger = None

def _log(message: str, level: str = 'info', **kwargs):
    """Internal logging helper with fallback."""
    if _logger:
        getattr(_logger, level)(message, **kwargs)
    elif level in ('warning', 'error', 'critical'):
        print(f"[RoleExtractor] {level.upper()}: {message}")


# v3.0.100: ReDoS Protection (ISSUE-001)
# Maximum input length for regex operations to prevent CPU exhaustion
MAX_REGEX_INPUT_LENGTH = 10000
REGEX_CHUNK_SIZE = 5000  # Process in chunks if needed


def safe_regex_search(pattern, text: str, flags=0, max_length: int = MAX_REGEX_INPUT_LENGTH):
    """
    Safe wrapper for regex search with input length limiting.
    Prevents ReDoS attacks by truncating overly long inputs.
    
    Args:
        pattern: Compiled regex or pattern string
        text: Text to search
        flags: Regex flags (if pattern is string)
        max_length: Maximum input length to process
        
    Returns:
        Match object or None
    """
    if not text:
        return None
    
    # Truncate if too long
    if len(text) > max_length:
        _log(f"Input truncated from {len(text)} to {max_length} chars for regex safety", level='debug')
        text = text[:max_length]
    
    try:
        if isinstance(pattern, str):
            return re.search(pattern, text, flags)
        return pattern.search(text)
    except re.error as e:
        _log(f"Regex error: {e}", level='warning')
        return None


def safe_regex_findall(pattern, text: str, flags=0, max_length: int = MAX_REGEX_INPUT_LENGTH):
    """
    Safe wrapper for regex findall with input length limiting.
    
    Args:
        pattern: Compiled regex or pattern string
        text: Text to search
        flags: Regex flags (if pattern is string)
        max_length: Maximum input length to process
        
    Returns:
        List of matches (empty if error)
    """
    if not text:
        return []
    
    # Truncate if too long
    if len(text) > max_length:
        _log(f"Input truncated from {len(text)} to {max_length} chars for regex safety", level='debug')
        text = text[:max_length]
    
    try:
        if isinstance(pattern, str):
            return re.findall(pattern, text, flags)
        return pattern.findall(text)
    except re.error as e:
        _log(f"Regex error: {e}", level='warning')
        return []


def safe_regex_finditer(pattern, text: str, flags=0, max_length: int = MAX_REGEX_INPUT_LENGTH):
    """
    Safe wrapper for regex finditer with input length limiting.
    
    Args:
        pattern: Compiled regex or pattern string
        text: Text to search
        flags: Regex flags (if pattern is string)
        max_length: Maximum input length to process
        
    Yields:
        Match objects
    """
    if not text:
        return
    
    # Truncate if too long
    if len(text) > max_length:
        _log(f"Input truncated from {len(text)} to {max_length} chars for regex safety", level='debug')
        text = text[:max_length]
    
    try:
        if isinstance(pattern, str):
            yield from re.finditer(pattern, text, flags)
        else:
            yield from pattern.finditer(text)
    except re.error as e:
        _log(f"Regex error: {e}", level='warning')


class EntityKind(Enum):
    """Classification of extracted entity type."""
    ROLE = "role"
    DELIVERABLE = "deliverable"
    UNKNOWN = "unknown"


class ExtractionMode(Enum):
    """
    v3.4.0: Extraction mode for controlling accuracy vs recall tradeoff.

    DISCOVERY - Current behavior: Finds potential roles using patterns and NLP.
                May include false positives that need post-extraction verification.
                Best for: Initial document surveys, discovering new roles.

    STRICT    - High-accuracy mode: Only extracts roles that:
                1. Are in the KNOWN_ROLES whitelist
                2. Appear as exact text matches in the document
                3. Each occurrence is linked to its exact character position
                Guarantees 100% verifiable roles.
                Best for: Audits, compliance reviews, production reports.
    """
    DISCOVERY = "discovery"
    STRICT = "strict"


@dataclass
class RoleOccurrence:
    """Represents a single occurrence of a role in the document."""
    role: str
    context: str
    responsibility: str
    action_type: str
    location: str
    confidence: float


class RoleSource(Enum):
    """
    v3.4.0: Indicates where a role came from for verification purposes.

    DICTIONARY  - Role exists in the organization's role dictionary (trusted)
    KNOWN       - Role is in KNOWN_ROLES whitelist (built-in, trusted)
    DISCOVERED  - Role was pattern-matched but not in dictionary (needs review)
    """
    DICTIONARY = "dictionary"
    KNOWN = "known"
    DISCOVERED = "discovered"


@dataclass
class ExtractedRole:
    """Aggregated information about a discovered entity (role or deliverable)."""
    canonical_name: str
    entity_kind: EntityKind = EntityKind.UNKNOWN  # v3.0.12: Added entity classification
    kind_confidence: float = 0.0  # v3.0.12: Confidence in classification
    kind_reason: str = ""  # v3.0.12: Why classified this way
    variants: Set[str] = field(default_factory=set)
    occurrences: List[RoleOccurrence] = field(default_factory=list)
    responsibilities: List[str] = field(default_factory=list)
    action_types: Dict[str, int] = field(default_factory=lambda: defaultdict(int))
    # v3.4.0: Role verification fields
    role_source: RoleSource = RoleSource.DISCOVERED  # Where did this role come from?
    is_verified_in_text: bool = False  # Does this exact text appear in the document?
    text_positions: List[int] = field(default_factory=list)  # Character positions where found
    
    @property
    def frequency(self) -> int:
        return len(self.occurrences)
    
    @property
    def avg_confidence(self) -> float:
        if not self.occurrences:
            return 0.0
        return sum(o.confidence for o in self.occurrences) / len(self.occurrences)
    
    def to_dict(self) -> dict:
        """Convert to dictionary for JSON serialization."""
        return {
            'canonical_name': self.canonical_name,
            'entity_kind': self.entity_kind.value,
            'kind_confidence': self.kind_confidence,
            'kind_reason': self.kind_reason,
            'variants': list(self.variants),
            'frequency': self.frequency,
            'avg_confidence': self.avg_confidence,
            'responsibilities': self.responsibilities,
            'action_types': dict(self.action_types),
            # v3.4.0: Role verification fields
            'role_source': self.role_source.value,
            'is_verified_in_text': self.is_verified_in_text,
            'text_positions': self.text_positions
        }


class RoleExtractor:
    """
    Extracts organizational roles from engineering documents.
    Pure Python implementation for air-gapped networks.
    """
    
    # =========================================================================
    # CONFIGURATION - Customize these for your organization
    # =========================================================================
    
    ROLE_SUFFIXES = [
        'engineer', 'manager', 'lead', 'director', 'officer', 'specialist',
        'analyst', 'coordinator', 'administrator', 'authority', 'chief',
        'supervisor', 'inspector', 'auditor', 'reviewer', 'approver',
        'representative', 'owner', 'custodian', 'architect', 'integrator',
        'technician', 'scientist', 'investigator', 'controller', 'planner',
        'panel', 'board', 'council', 'committee', 'team', 'group'
    ]
    
    ROLE_MODIFIERS = [
        'project', 'program', 'systems', 'system', 'lead', 'chief', 'senior',
        'deputy', 'assistant', 'associate', 'principal', 'technical', 'quality',
        'safety', 'mission', 'flight', 'ground', 'test', 'integration',
        'verification', 'validation', 'configuration', 'data', 'risk',
        'requirements', 'interface', 'software', 'hardware', 'mechanical',
        'electrical', 'structural', 'thermal', 'propulsion', 'avionics',
        'reliability', 'maintainability', 'logistics', 'operations', 'security',
        'environmental', 'human', 'factors', 'design', 'manufacturing', 'production',
        'subsystem', 'component', 'element', 'responsible', 'cognizant',
        'designated', 'authorized', 'certifying', 'contracting', 'procurement',
        'review', 'control', 'change', 'engineering', 'standing', 'independent',
        'working', 'action', 'steering', 'executive', 'advisory'
    ]
    
    KNOWN_ROLES = [
        # Core engineering roles
        'systems engineer', 'project manager', 'program manager', 'chief engineer',
        'lead engineer', 'lead systems engineer', 'technical authority',
        'safety engineer', 'quality assurance', 'quality assurance engineer',
        'configuration manager', 'data manager', 'risk manager', 'test engineer',
        'integration engineer', 'verification engineer', 'validation engineer',
        'software engineer', 'hardware engineer', 'design engineer',
        
        # Specialized engineering
        'system safety engineer', 'reliability engineer', 'maintainability engineer',
        'human factors engineer', 'environmental engineer', 'thermal engineer',
        'structural engineer', 'propulsion engineer', 'avionics engineer',
        'manufacturing engineer', 'production engineer', 'logistics engineer',
        'sustaining engineer', 'operations engineer', 'process engineer',
        
        # Leadership/management
        'mission assurance', 'mission assurance manager', 'flight director', 
        'mission director', 'ground controller', 'principal investigator',
        'co-investigator', 'project scientist', 'technical lead',
        'discipline lead', 'functional lead', 'subsystem lead', 'element lead',
        'software lead', 'hardware lead', 'test lead', 'integration lead',
        'verification lead', 'requirements manager', 'test manager',
        
        # Contract/government roles
        'contracting officer', 'contracting officer representative', 'cor',
        'technical monitor', 'government technical representative',
        'authorizing official', 'designated engineering representative',
        'designated airworthiness representative', 'contractor', 'customer',
        
        # Teams and groups
        'test team', 'project team', 'technical team', 'development team',
        'integration team', 'review team', 'support team',
        'facility operators', 'shift engineers', 'facility operator',
        'shift engineer', 'test personnel', 'technical staff',
        
        # Boards/panels/groups
        'configuration control board', 'change control board',
        'engineering review board', 'technical review board',
        'system safety panel', 'safety panel', 'review panel', 'review board',
        'independent review team', 'standing review board',
        'interface control working group', 'integrated product team',
        'integrated product team lead', 'working group',
        
        # Executive/directorate
        'mission directorate', 'center director', 'deputy director',
        'associate administrator', 'mission directorate associate administrator',
        'office of the chief engineer', 'project planning and control',
        
        # v3.0.91b: Additional roles for broader document support
        # Agile/Scrum roles
        'scrum master', 'scrum team', 'product owner', 'product manager',
        'sprint team', 'agile team', 'agile coach',
        
        # Executive roles
        'chief innovation officer', 'cino', 'deputy cino', 'deputy pgm',
        'chief architect', 'chief technology officer', 'cto',
        'chief information officer', 'cio', 'chief executive officer', 'ceo',
        'chief operations officer', 'coo', 'chief financial officer', 'cfo',
        
        # General organizational roles
        'stakeholder', 'stakeholders', 'subject matter expert', 'sme',
        'consultant', 'consultant team', 'quality auditor', 'business owner',
        'it pm', 'it project manager', 'information technology project manager',
        'project lead', 'city project manager', 'consultant project lead',
        
        # Support roles
        'sponsor', 'project sponsor', 'executive sponsor',
        'administrator', 'coordinator', 'facilitator',
        
        # v3.0.91c: Additional domain-specific roles
        # IT Security
        'chief information security officer', 'ciso', 'security officer',
        'information security officer', 'cybersecurity analyst',
        
        # Healthcare/Clinical
        'medical monitor', 'medical director', 'clinical director',
        'study coordinator', 'clinical research associate', 'cra',
        'data safety monitoring board', 'institutional review board',
        'ethics committee', 'sponsor medical director',

        # v3.2.5: FAA/Aviation-specific roles
        'accountable executive', 'accountable manager', 'safety manager',
        'certificate holder', 'certificate management team', 'flight crew',
        'pilot in command', 'second in command', 'flight engineer',
        'cabin crew', 'dispatcher', 'flight dispatcher', 'load master',
        'maintenance controller', 'director of safety', 'director of operations',
        'director of maintenance', 'chief pilot', 'check airman',
        'training captain', 'standards captain', 'designated examiner',
        'aviation safety inspector', 'principal operations inspector',
        'principal maintenance inspector', 'principal avionics inspector',
        'air carrier', 'operator', 'certificate management office',

        # v3.2.5: OSHA/Safety-specific roles
        'safety committee', 'process safety coordinator', 'process owner',
        'area manager', 'plant manager', 'facility manager',
        'operations manager', 'shift supervisor', 'unit supervisor',
        'emergency coordinator', 'emergency response team', 'fire brigade',
        'hazmat team', 'rescue team', 'first responder', 'safety officer',
        'industrial hygienist', 'environmental health specialist',
        'compliance officer', 'safety inspector', 'loss prevention specialist',

        # v3.3.0: Generic worker/employee roles for OSHA compliance
        'employer', 'employers', 'employee', 'employees',
        'contract employee', 'contract employees', 'host employer',
        'controlling employer', 'staffing agency',
        'temporary worker', 'temporary workers',
        'front-line employee', 'front-line employees',
        'line employee', 'line employees',
        'operating personnel', 'management personnel',
        'maintenance personnel', 'production personnel',
        'technical personnel', 'support personnel',
        'affected employee', 'affected employees',
        'authorized employee', 'authorized employees',
        'competent person', 'qualified person',
        'worker', 'workers', 'staff', 'personnel',

        # v3.3.0: Academic/Research roles for university SOPs
        'graduate student', 'graduate students',
        'undergraduate student', 'undergraduate students',
        'research assistant', 'research assistants',
        'postdoctoral researcher', 'postdoctoral researchers', 'postdoc', 'postdocs',
        'research staff', 'research staff member', 'research staff members',
        'laboratory supervisor', 'lab supervisor',
        'laboratory manager', 'lab manager',
        'research coordinator', 'lab coordinator',
        'thesis advisor', 'faculty advisor', 'faculty member', 'faculty members',
        'procedure author',

        # v3.3.0: Additional FAA/Aviation roles
        'dispatchers', 'ground handling', 'ground handling personnel',
        'ramp personnel', 'flight attendant', 'flight attendants',
        'crew member', 'crew members', 'line pilot', 'line pilots',

        # v3.3.1: Additional roles identified from testing gaps
        # Executive titles
        'president', 'vice president', 'executive director',
        'managing director', 'board member', 'board of directors',

        # Contract roles
        'contract employer', 'contract employers',
        'subcontractor employee', 'subcontractor employees',

        # Analysis/review teams
        'process hazard analysis team', 'hazard analysis team',
        'incident investigation team', 'root cause analysis team',
        'safety review team', 'audit team',

        # Academic administrative roles
        'department chair', 'department head', 'dean',
        'associate dean', 'program director', 'academic advisor',

        # EH&S specific
        'eh&s coordinator', 'ehs coordinator',
        'environmental health and safety coordinator',
        'safety representative', 'safety liaison',

        # v3.3.2: Defense/Military sector roles (MIL-STD compliance)
        # Government acquisition roles (note: 'government' moved to ORGANIZATION_ENTITIES)
        'procuring activity', 'requiring activity',
        'project officer', 'government representative',
        'approving authority', 'approval authority',
        'preparing activity', 'reviewing activity',

        # Contractor roles
        'prime contractor', 'subcontractor', 'subcontractors',
        'vendor', 'vendors', 'supplier', 'suppliers',

        # Technical manual roles
        'technical writer', 'technical writers', 'illustrator', 'illustrators',
        'editor', 'editors', 'author', 'tm author',
        'custodian', 'document custodian',

        # User/operator roles
        'user', 'users', 'end user', 'end users',
        'maintainer', 'maintainers', 'maintenance technician',
        'technician', 'technicians', 'inspector', 'inspectors',

        # Quality roles
        'quality assurance representative', 'quality control',
        'quality inspector', 'qar',

        # Engineering roles
        'engineering personnel', 'technical personnel',
        'logistics engineer', 'logistics personnel',
        'design activity', 'engineering activity',

        # v3.5.0: NASA/Aerospace-specific roles (from NASA SE Handbook analysis)
        'technical fellow', 'independent technical authority',
        'mission systems engineer', 'ground systems engineer',
        'flight systems engineer', 'payload engineer',
        'spacecraft engineer', 'launch vehicle engineer',
        'mission directorate associate administrator',
        'governing program management council',
        'failure review board', 'material review board',
        'anomaly resolution team', 'program management council',
        'insight provider', 'oversight provider',
        'acquisition manager', 'flight operations engineer',
        'ground controller', 'launch director',
        'range safety officer', 'payload integrator',
        'deputy project manager', 'deputy chief engineer',
        'systems engineering lead', 'se&i lead',
        'independent verification and validation team',
        'configuration management board',
    ]
    
    ACRONYM_MAP = {
        'pm': 'project manager', 'se': 'systems engineer',
        'lse': 'lead systems engineer', 'ipt': 'integrated product team',
        'cor': 'contracting officer representative',
        'gtr': 'government technical representative',
        'der': 'designated engineering representative',
        'dar': 'designated airworthiness representative',
        'ccb': 'configuration control board', 'erb': 'engineering review board',
        'trb': 'technical review board', 'irt': 'independent review team',
        'srb': 'standing review board', 'oce': 'office of the chief engineer',
        'mdaa': 'mission directorate associate administrator',
        'qa': 'quality assurance', 'ma': 'mission assurance',
        'sma': 'safety and mission assurance', 'ivv': 'independent verification and validation',
        'pp&c': 'project planning and control', 'icwg': 'interface control working group',
        # v3.0.91c additions
        'ciso': 'chief information security officer', 'cra': 'clinical research associate',
        'irb': 'institutional review board', 'dsmb': 'data safety monitoring board'
    }
    
    ACTION_VERBS = {
        'performs': ['perform', 'performs', 'performing', 'execute', 'executes', 'conduct', 'conducts', 'conducting'],
        'approves': ['approve', 'approves', 'approving', 'authorize', 'authorizes', 'sign', 'signs', 'certify', 'certifies'],
        'reviews': ['review', 'reviews', 'reviewing', 'evaluate', 'evaluates', 'assess', 'assesses', 'examine', 'examines'],
        'coordinates': ['coordinate', 'coordinates', 'coordinating', 'collaborate', 'collaborates', 'liaise', 'liaises'],
        'manages': ['manage', 'manages', 'managing', 'oversee', 'oversees', 'direct', 'directs', 'supervise', 'supervises'],
        'supports': ['support', 'supports', 'supporting', 'assist', 'assists', 'help', 'helps', 'aid', 'aids'],
        'verifies': ['verify', 'verifies', 'verifying', 'validate', 'validates', 'confirm', 'confirms', 'check', 'checks'],
        'develops': ['develop', 'develops', 'developing', 'create', 'creates', 'design', 'designs', 'prepare', 'prepares'],
        'maintains': ['maintain', 'maintains', 'maintaining', 'update', 'updates', 'sustain', 'sustains'],
        'ensures': ['ensure', 'ensures', 'ensuring', 'guarantee', 'guarantees', 'assure', 'assures'],
        'provides': ['provide', 'provides', 'providing', 'supply', 'supplies', 'deliver', 'delivers'],
        'receives': ['receive', 'receives', 'receiving', 'obtain', 'obtains'],
        'reports': ['report', 'reports', 'reporting', 'communicate', 'communicates', 'inform', 'informs'],
        'defines': ['define', 'defines', 'defining', 'specify', 'specifies', 'establish', 'establishes'],
        'implements': ['implement', 'implements', 'implementing', 'deploy', 'deploys'],
        'leads': ['lead', 'leads', 'leading'],
        'monitors': ['monitor', 'monitors', 'monitoring', 'track', 'tracks', 'tracking']
    }
    
    FALSE_POSITIVES = [
        # Generic terms
        'the system', 'the document', 'the customer', 'the contractor', 'the government',
        'the agency', 'the organization', 'the project', 'the program',
        'the process', 'the procedure', 'the requirement', 'the specification',
        'the design', 'the product', 'the hardware', 'the software', 'the element',
        'the component', 'the subsystem', 'the interface', 'the data', 'the information',
        'this document', 'this section', 'this chapter', 'this appendix',
        'system', 'document', 'customer',
        'agency', 'organization', 'project', 'program', 'process', 'procedure',
        # NOTE: 'contractor' and 'government' REMOVED from false positives - they are valid roles in defense docs
        
        # Process/discipline names (not roles)
        'systems engineering', 'project management', 'configuration management',
        'risk management', 'data management', 'technical management',
        'quality assurance process', 'verification process', 'validation process',
        
        # Activity-based
        'integration activities', 'verification activities', 'validation activities',
        'test activities', 'review activities', 'design activities', 'development activities',
        'manufacturing activities', 'production activities', 'operations activities',
        
        # Review/document types
        'technical reviews', 'safety reviews', 'design reviews', 'peer reviews',
        'milestone reviews', 'gate reviews', 'phase reviews',
        'interface specifications', 'design specifications', 'requirements specifications',
        'technical requirements', 'functional requirements', 'performance requirements',
        'system requirements', 'software requirements', 'hardware requirements',
        
        # v3.0.91b: Expanded false positives from testing
        # Generic single words that are not roles
        'progress', 'upcoming', 'distinct', 'others', 'addition', 'manner',
        'work', 'test', 'end', 'task', 'plan', 'phase', 'configuration',
        'property', 'resource', 'attachment', 'reports', 'technical',
        'engineering', 'authority', 'international travel', 'coordinating',
        
        # Events and milestones (not roles)
        'test readiness review', 'test readiness', 'test preparation begins',
        'design review', 'preliminary design review', 'critical design review',
        'system requirements review', 'mission readiness review',
        'operational readiness review', 'flight readiness review',
        'phase transition', 'key decision point', 'milestone',
        
        # Facilities and equipment (not roles)
        'panel test facility', 'flight facility', 'test facility',
        'wind tunnel', 'arc jet', 'vacuum chamber', 'clean room',
        'thermal protection', 'support facility', 'support facilities',
        
        # Processes and methodologies
        'reliability centered maintenance', 'condition based maintenance',
        'preventive maintenance', 'corrective maintenance', 'predictive maintenance',
        'property management', 'resource scheduling', 'contract reporting',
        'safety and environmental compliance', 'thermo physics facilities configuration',
        
        # Document and contract elements
        'statement of work', 'scope of work', 'contract data requirements',
        'task order', 'idiq task orders', 'contract line item',
        'operations and maintenance plan', 'technical project',
        
        # Generic phrases that get extracted incorrectly
        'other disciplines', 'other arc organizations', 'foreign systems',
        'all parties', 'project teams', 'wide variety of skills',
        'monthly financial', 'written status report',
        
        # Verb phrases and fragments
        'facilities perform various', 'mission goes beyond just',
        'develop an approach that optimizes', 'facilitate completion of work',
        'operate facilities', 'ensure that utilities', 'provide property',
        'provide the risk', 'coordinate the demand',
        
        # v3.0.91c: Additional false positives from 5-document validation test
        # Organizational disciplines/functions (not roles unless paired with role suffix)
        'mission assurance', 'mission equipment', 'mission systems',
        'configuration control', 'version control',  # NOTE: 'quality control' removed - it IS a valid role
        'safety and environmental', 'security engineering',
        
        # Truncated or partial role names
        'chief innovation', 'deputy chief', 'assistant deputy',
        
        # IPT/Team fragments
        'staffing integrated product team', 'se ipt lead se',
        'verification engineer',  # Usually too vague without context

        # v3.2.5: Additional false positives from accuracy analysis
        # Location/address-like patterns
        'atlanta federal center', 'curtis center', 'federal center',
        'state office building', 'labor division',
        # Generic organizational references
        'safety objectives', 'safety department',
        'large organization', 'small organization',
        # Run-together word patterns
        'byasafetydepartment', 'bythesafetydepartment', 'thepersonnel',
        'thedifference', 'theaccountable', 'thecoordinate', 'thecertificateholder',

        # v3.3.0: Safety management concepts (processes, not roles)
        'safety management', 'safety management system', 'safety policy',
        'safety promotion', 'safety assurance', 'safety risk management',
        'process safety management', 'process safety', 'safety culture',
        'safety performance', 'safety objective',
        'hazard management', 'risk assessment', 'incident investigation',

        # v3.3.0: Document/regulatory references
        'advisory circular', 'advisory circulars', 'federal register',
        'code of federal regulations', 'regulatory requirements',
        'federal aviation regulations', 'aviation regulations',

        # v3.3.0: Physical hazards (not roles)
        'electric shock', 'physical contact', 'arc flash',
        'radiant heat', 'cold hazard', 'fire hazard', 'explosion hazard',
        'noise hazard', 'ergonomic hazard', 'body impact',

        # v3.3.0: Equipment/materials (not roles)
        'capacitor banks', 'pressure vessels', 'electrical conductors',
        'cobot interaction', 'robotic control software',

        # v3.3.1: Department/organizational unit names (not roles)
        'environmental health and safety', 'human resources',
        'information technology', 'research and development',
        'quality control department', 'safety department'
    ]
    
    # v3.0.91b: Words that should NEVER be roles on their own
    SINGLE_WORD_EXCLUSIONS = {
        # Generic nouns
        'progress', 'upcoming', 'distinct', 'others', 'addition', 'manner',
        'work', 'test', 'end', 'task', 'plan', 'phase', 'reports', 'property',
        'resource', 'attachment', 'travel', 'training', 'construction',
        'calibration', 'demolition', 'repair', 'installation', 'procurement',
        'schedule', 'budget', 'scope', 'status', 'summary', 'overview',
        'configuration', 'process', 'contract', 'mission', 'officer',
        # Adjectives
        'technical', 'functional', 'operational', 'administrative', 'preliminary',
        'final', 'initial', 'current', 'previous', 'various', 'multiple',
        # Verbs/gerunds
        'coordinating', 'managing', 'performing', 'conducting', 'reviewing',
        'approving', 'verifying', 'validating', 'monitoring', 'tracking',
        'planning', 'scheduling', 'reporting', 'documenting', 'processing',
        # v3.2.5: Additional single-word exclusions from accuracy analysis
        'user', 'users', 'owner', 'owners', 'chief', 'team', 'teams',
        'group', 'groups', 'labor', 'section', 'sections', 'suite',
        'center', 'division', 'office', 'department', 'authority',
        'authorities', 'contractor', 'contractors', 'vendor', 'vendors',
        'stakeholder', 'stakeholders', 'evaluator', 'evaluators',
        'operator', 'operators', 'objectives', 'scalability',
        # v3.5.0: Additional exclusions from NASA SE Handbook test (false positives)
        # Common English words that pattern-matching incorrectly extracts as roles
        'there', 'care', 'description', 'same', 'different', 'how', 'what',
        'when', 'where', 'why', 'which', 'often', 'although', 'some',
        'scarcity', 'time', 'metrics', 'review', 'management', 'software',
        'hardware', 'design', 'analysis', 'system', 'systems', 'interface',
        'data', 'information', 'element', 'component', 'subsystem',
        'verification', 'validation', 'integration', 'reliability',
        'safety', 'quality', 'heritage', 'risk', 'requirement',
        'requirements', 'specification', 'standard', 'baseline',
        'product', 'assembly', 'module', 'unit', 'model', 'prototype',
        'document', 'report', 'procedure', 'guideline', 'handbook',
        'program', 'project', 'activity', 'discipline', 'function',
        'approach', 'method', 'technique', 'tool', 'capability',
        # v3.5.0b: Additional false positives from NASA SE Handbook deep analysis
        'framework', 'supplement', 'spacecraft', 'structure', 'request',
        'agreement', 'agreements', 'architecture', 'assessment', 'required',
        'figure', 'prepare', 'areas', 'progression', 'descriptions',
        'representative', 'se', 'ta', 'pi', 'electromechnical',
    }

    # v3.5.0: Organization entities - these are NOT roles (they are orgs/entities)
    # Pattern matching may extract these but they should be filtered out
    ORGANIZATION_ENTITIES = {
        'nasa', 'government', 'agency', 'air force', 'army', 'navy',
        'department of defense', 'dod', 'faa', 'osha', 'epa',
        'european space agency', 'esa', 'jaxa', 'csa',
        'congress', 'senate', 'white house', 'pentagon',
    }
    
    # Deliverable patterns - items that are work products, not roles
    DELIVERABLE_PATTERNS = [
        # Document types
        r'\b(plan|report|specification|analysis|assessment|study)\b',
        r'\b(document|manual|guide|handbook|procedure|instruction)\b',
        r'\b(drawing|schematic|diagram|model|prototype|mockup)\b',
        r'\b(database|repository|archive|library|registry)\b',
        r'\b(schedule|timeline|roadmap|milestone|gantt)\b',
        r'\b(budget|estimate|proposal|quotation|bid)\b',
        r'\b(contract|agreement|memorandum|charter)\b',
        # Acronyms commonly used for deliverables
        r'\bICD\b|\bSRS\b|\bSDD\b|\bCDRL\b|\bDID\b|\bSOW\b|\bWBS\b',
        r'\bSEMP\b|\bSPMP\b|\bCMP\b|\bQAP\b|\bSAFETY\s+PLAN\b',
        # Test/verification outputs
        r'\btest\s+(report|results|data|log|procedure)\b',
        r'\bverification\s+(report|results|matrix)\b',
        r'\bvalidation\s+(report|results)\b',
    ]
    
    BOUNDARY_WORDS = [
        'and', 'or', 'who', 'which', 'that', 'with', 'for', 'in', 'on', 'at', 'as',
        'to', 'from', 'by', 'is', 'are', 'was', 'were', 'has', 'have', 'had',
        'shall', 'will', 'may', 'can', 'should', 'must', 'could', 'would',
        'the', 'a', 'an', 'this', 'that', 'these', 'those', 'their', 'its',
        'prior', 'before', 'after', 'during', 'while', 'when', 'if', 'unless'
    ]

    def __init__(self, custom_roles: List[str] = None, custom_false_positives: List[str] = None,
                 use_dictionary: bool = True, use_nlp: bool = True,
                 extraction_mode: ExtractionMode = None):
        """
        Initialize the role extractor with optional customizations.

        Args:
            custom_roles: Additional role names to recognize
            custom_false_positives: Additional false positives to ignore
            use_dictionary: If True, attempt to load roles from database dictionary
            use_nlp: If True, use NLP enhancement for better accuracy
            extraction_mode: ExtractionMode.DISCOVERY (default) or ExtractionMode.STRICT
                             STRICT mode only extracts whitelisted roles with exact matches
                             for 100% accuracy guarantee.
        """
        # v3.4.0: Extraction mode - STRICT for 100% accuracy, DISCOVERY for recall
        self.extraction_mode = extraction_mode or ExtractionMode.DISCOVERY
        _log(f"Role extractor initialized in {self.extraction_mode.value} mode", level='info')

        self.known_roles = set(r.lower() for r in self.KNOWN_ROLES)
        if custom_roles:
            self.known_roles.update(r.lower() for r in custom_roles)

        # v3.5.0: Organization entities (not roles) - filtered from results
        self._organization_entities = set(self.ORGANIZATION_ENTITIES)

        # Load from dictionary if available
        if use_dictionary:
            dict_roles = self._load_dictionary_roles()
            if dict_roles:
                self.known_roles.update(r.lower() for r in dict_roles)

        self.false_positives = set(fp.lower() for fp in self.FALSE_POSITIVES)
        if custom_false_positives:
            self.false_positives.update(fp.lower() for fp in custom_false_positives)

        # v3.2.0: Load rejected roles from dictionary as false positives
        # This makes the extractor smarter over time as users adjudicate roles
        if use_dictionary:
            rejected_roles = self._load_rejected_roles()
            if rejected_roles:
                self.false_positives.update(r.lower() for r in rejected_roles)
                _log(f"Loaded {len(rejected_roles)} rejected roles as false positives", level='debug')
        
        # Initialize NLP processor for better extraction (v3.1.2 - ENH-008)
        self._nlp_processor = None
        if use_nlp:
            try:
                from nlp_utils import NLPProcessor
                self._nlp_processor = NLPProcessor()
                if self._nlp_processor.is_nlp_available:
                    _log("NLP processor loaded with spaCy for improved role extraction", level='info')
                else:
                    _log("NLP processor loaded (pattern matching mode)", level='debug')
            except ImportError:
                _log("NLP utils not available, using standard extraction", level='debug')
            except Exception as e:
                _log(f"NLP processor init error: {e}, using standard extraction", level='warning')

        # =====================================================================
        # v3.3.0: ENHANCED NLP INTEGRATION
        # =====================================================================
        # Initialize enhanced NLP and adaptive learning for maximum accuracy
        self._enhanced_nlp = None
        self._adaptive_learner = None
        self._technical_dict = None
        if use_nlp:
            try:
                from nlp_integration import (
                    get_enhanced_nlp_integration,
                    get_adaptive_learner_integration
                )
                self._enhanced_nlp = get_enhanced_nlp_integration()
                if self._enhanced_nlp.is_available:
                    _log("Enhanced NLP processor loaded (v3.3.0) with transformer support", level='info')

                self._adaptive_learner = get_adaptive_learner_integration()
                if self._adaptive_learner.is_available:
                    _log("Adaptive learner integration loaded (v3.3.0)", level='info')
            except ImportError:
                _log("v3.3.0 NLP integration not available", level='debug')
            except Exception as e:
                _log(f"v3.3.0 NLP init error: {e}", level='warning')

            # Load technical dictionary for term validation
            try:
                from technical_dictionary import get_technical_dictionary
                self._technical_dict = get_technical_dictionary()
                if self._technical_dict:
                    _log("Technical dictionary loaded for role validation", level='debug')
            except ImportError:
                pass
            except Exception as e:
                _log(f"Technical dictionary init error: {e}", level='warning')

        self._build_patterns()
    
    def _load_dictionary_roles(self) -> List[str]:
        """Load active roles from the database dictionary if available."""
        try:
            import sqlite3
            from pathlib import Path
            from contextlib import contextmanager

            @contextmanager
            def _db_conn(path):
                conn = sqlite3.connect(str(path))
                conn.row_factory = sqlite3.Row
                conn.execute('PRAGMA journal_mode=WAL')
                cur = conn.cursor()
                try:
                    yield (conn, cur)
                    conn.commit()
                except Exception:
                    conn.rollback()
                    raise
                finally:
                    conn.close()

            # Try common database locations
            possible_paths = [
                Path(__file__).parent / 'scan_history.db',
                Path(__file__).parent / 'data' / 'scan_history.db',
                Path.home() / '.twr' / 'scan_history.db'
            ]

            for db_path in possible_paths:
                if db_path.exists():
                    with _db_conn(db_path) as (conn, cursor):
                        # Check if table exists
                        cursor.execute("""
                            SELECT name FROM sqlite_master
                            WHERE type='table' AND name='role_dictionary'
                        """)
                        if cursor.fetchone():
                            cursor.execute("""
                                SELECT role_name, aliases FROM role_dictionary
                                WHERE is_active = 1 AND is_deliverable = 0
                            """)

                            roles = []
                            for row in cursor.fetchall():
                                roles.append(row[0])
                                # Also add aliases
                                if row[1]:
                                    try:
                                        import json
                                        aliases = json.loads(row[1])
                                        roles.extend(aliases)
                                    except Exception:
                                        pass

                            return roles

            return []
        except Exception as e:
            # Silently fail - dictionary is optional
            return []

    def _load_rejected_roles(self) -> List[str]:
        """
        Load rejected roles from the database dictionary to use as false positives.

        v3.2.0: This makes the extractor smarter over time - when users reject
        a detected role through adjudication, it won't be flagged again in
        future extractions.

        Returns:
            List of role names that were rejected during adjudication
        """
        try:
            import sqlite3
            from pathlib import Path
            from contextlib import contextmanager

            @contextmanager
            def _db_conn(path):
                conn = sqlite3.connect(str(path))
                conn.row_factory = sqlite3.Row
                conn.execute('PRAGMA journal_mode=WAL')
                cur = conn.cursor()
                try:
                    yield (conn, cur)
                    conn.commit()
                except Exception:
                    conn.rollback()
                    raise
                finally:
                    conn.close()

            # Try common database locations
            possible_paths = [
                Path(__file__).parent / 'scan_history.db',
                Path(__file__).parent / 'data' / 'scan_history.db',
                Path.home() / '.twr' / 'scan_history.db'
            ]

            for db_path in possible_paths:
                if db_path.exists():
                    with _db_conn(db_path) as (conn, cursor):
                        # Check if table exists
                        cursor.execute("""
                            SELECT name FROM sqlite_master
                            WHERE type='table' AND name='role_dictionary'
                        """)
                        if cursor.fetchone():
                            # Get rejected roles (is_active = 0) that came from adjudication
                            cursor.execute("""
                                SELECT role_name, aliases FROM role_dictionary
                                WHERE is_active = 0 AND source = 'adjudication'
                            """)

                            rejected = []
                            for row in cursor.fetchall():
                                rejected.append(row[0])
                                # Also add aliases so variants are rejected too
                                if row[1]:
                                    try:
                                        import json
                                        aliases = json.loads(row[1])
                                        rejected.extend(aliases)
                                    except Exception:
                                        pass

                            return rejected

            return []
        except Exception as e:
            # Silently fail - this is optional enhancement
            return []

    def _build_patterns(self):
        """Build compiled regex patterns for role detection."""
        
        # Pattern 1: "The [Role] shall/will/is/has/ensures..."
        self.pattern_the_role = re.compile(
            r'\b[Tt]he\s+([A-Z][a-zA-Z]+(?:\s+[A-Z]?[a-zA-Z]+){0,5}?)\s+'
            r'(?:shall|will|has\s+(?:overall\s+)?responsibility|is\s+(?:skilled|responsible)|'
            r'ensures?|provides?|reviews?|approves?|manages?|coordinates?|performs?|'
            r'maintains?|leads?|oversees?|plays|usually|monitors?)',
            re.MULTILINE
        )
        
        # Pattern 2: "[Role] is responsible for..."
        self.pattern_role_is = re.compile(
            r'\b([A-Z][a-zA-Z]+(?:\s+[A-Z]?[a-zA-Z]+){0,5}?)\s+'
            r'(?:is\s+responsible\s+for|shall\s+\w|will\s+(?:be|ensure|provide|review|approve|manage))',
            re.MULTILINE
        )
        
        # Pattern 3: "by/to/from/with the [Role]"
        self.pattern_by_role = re.compile(
            r'(?:by|to|from|with)\s+the\s+([A-Z][a-zA-Z]+(?:\s+[A-Z]?[a-zA-Z]+){0,4}?)'
            r'(?=\s*[,\.\;]|\s+(?:and|or|who|for|in|on|at|as|prior|before|after)\b)',
            re.MULTILINE
        )
        
        # Pattern 4: "approved/reviewed/verified by [Role]"
        self.pattern_action_by = re.compile(
            r'(?:approved|reviewed|coordinated|submitted|verified|validated|signed|'
            r'certified|authorized|prepared|developed|conducted|performed)\s+'
            r'(?:by|with)\s+(?:the\s+)?([A-Z][a-zA-Z]+(?:\s+[A-Z]?[a-zA-Z]+){0,4}?)'
            r'(?=\s*[,\.\;]|\s+(?:and|or|who|for|in|on|at|as|prior|before|after)\b)',
            re.MULTILINE
        )
        
        # Pattern 5: "[Role]'s responsibility/role/duties"
        self.pattern_possessive = re.compile(
            r"([A-Z][a-zA-Z]+(?:\s+[A-Z]?[a-zA-Z]+){0,4}?)'s\s+"
            r"(?:responsibility|role|duties|function|authority|approval)",
            re.MULTILINE
        )
        
        # Pattern 6: "responsibilities/role of the [Role]"
        self.pattern_responsibilities_of = re.compile(
            r'(?:responsibilities?|roles?|duties|functions?|authority)\s+of\s+(?:the\s+)?'
            r'([A-Z][a-zA-Z]+(?:\s+[A-Z]?[a-zA-Z]+){0,4}?)'
            r'(?=\s*[,\.\;]|\s+(?:and|or|include|are|is)\b)',
            re.MULTILINE
        )
        
        # Pattern 7: Acronym in parentheses "Project Manager (PM)"
        # v3.2.5: Extended to capture up to 6 words total for longer role names
        # Also added support for & in role names (e.g., "Safety & Mission Assurance")
        self.pattern_acronym = re.compile(
            r'([A-Z][a-zA-Z]+(?:[\s&]+[A-Z]?[a-zA-Z]+){0,5}?)\s*\(([A-Z][A-Z&/]{1,7})\)',
            re.MULTILINE
        )
        
        # Pattern 8: "as the [Role]" / "serve as [Role]"
        self.pattern_as_role = re.compile(
            r'(?:as\s+(?:the\s+)?|serve\s+as\s+(?:the\s+)?|acting\s+(?:as\s+)?(?:the\s+)?)'
            r'([A-Z][a-zA-Z]+(?:\s+[A-Z]?[a-zA-Z]+){0,4}?)'
            r'(?=\s*[,\.\;]|\s+(?:and|or|who|for|in|on|at)\b)',
            re.MULTILINE
        )
        
        # Pattern 9: "notify/inform/contact/report to the [Role]"
        self.pattern_notify = re.compile(
            r'(?:notify|inform|contact|consult|report\s+to)\s+(?:the\s+)?'
            r'([A-Z][a-zA-Z]+(?:\s+[A-Z]?[a-zA-Z]+){0,4}?)'
            r'(?=\s*[,\.\;]|\s+(?:and|or|who|for|in|on|at|prior|before|after)\b)',
            re.MULTILINE
        )
        
        # Pattern 10: Start of sentence "[Role] shall/will/ensures..."
        self.pattern_sentence_start = re.compile(
            r'(?:^|\.\s+)([A-Z][a-zA-Z]+(?:\s+[A-Z]?[a-zA-Z]+){0,4}?)\s+'
            r'(?:shall|will|must|should|is\s+responsible|ensures?|provides?|'
            r'reviews?|approves?|manages?|coordinates?|performs?|maintains?|leads?|oversees?)',
            re.MULTILINE
        )
        
        # Pattern 11: End of sentence "...by the [Role]."
        self.pattern_end_of_sentence = re.compile(
            r'(?:approved|reviewed|verified|validated|signed|authorized|certified)\s+by\s+(?:the\s+)?'
            r'([A-Z][a-zA-Z]+(?:\s+[A-Z]?[a-zA-Z]+){1,4}?)\s*\.',
            re.MULTILINE
        )
        
        # Pattern 12: Role followed by "and" another role (coordinated with X and Y)
        self.pattern_role_and = re.compile(
            r'(?:by|with|to)\s+(?:the\s+)?([A-Z][a-zA-Z]+(?:\s+[A-Z]?[a-zA-Z]+){0,3}?)\s+and\s+'
            r'(?:the\s+)?([A-Z][a-zA-Z]+(?:\s+[A-Z]?[a-zA-Z]+){0,3}?)(?=\s*[,\.\;]|\s+(?:who|for|in|on|at)\b)',
            re.MULTILINE
        )
        
        # Pattern 13: "from the [Role]" / "require approval from the [Role]"
        self.pattern_from_role = re.compile(
            r'(?:from|require(?:s)?\s+(?:approval\s+)?from)\s+(?:the\s+)?'
            r'([A-Z][a-zA-Z]+(?:\s+[A-Z]?[a-zA-Z]+){0,4}?)(?=\s*[,\.\;]|\s+(?:and|or|who|for|in|on|at)\b)',
            re.MULTILINE
        )
        
        # Pattern 14: "the [Role] assists/works with"
        self.pattern_role_assists = re.compile(
            r'\b[Tt]he\s+([A-Z][a-zA-Z]+(?:\s+[A-Z]?[a-zA-Z]+){0,4}?)\s+'
            r'(?:assists?|works?\s+with|supports?|helps?)',
            re.MULTILINE
        )
        
        # Pattern 15: "conducted by the [multi-word] and approved/reviewed by"
        self.pattern_conducted_by = re.compile(
            r'conducted\s+by\s+(?:the\s+)?([A-Z][a-zA-Z]+(?:\s+[A-Z]?[a-zA-Z]+){0,4}?)\s+and',
            re.MULTILINE
        )
        
        # Pattern 16: "analyzed/assembles/processes" etc - past tense verbs followed by roles
        self.pattern_role_action_object = re.compile(
            r'\b[Tt]he\s+([A-Z][a-zA-Z]+(?:\s+[A-Z]?[a-zA-Z]+){0,3}?)\s+'
            r'(?:analyzes?|assembles?|processes?|develops?|creates?|builds?|designs?|tests?|integrates?)',
            re.MULTILINE
        )
        
        # Pattern 17: "the [Role] in managing/in doing..."
        self.pattern_role_in = re.compile(
            r'\b(?:the\s+)?([A-Z][a-zA-Z]+(?:\s+[A-Z]?[a-zA-Z]+){0,4}?)\s+in\s+managing',
            re.MULTILINE
        )
        
        # Pattern 18: "require[s] approval from the [Role] and [Role]"
        self.pattern_approval_from = re.compile(
            r'(?:require|requires)\s+(?:approval\s+)?from\s+(?:the\s+)?'
            r'([A-Z][a-zA-Z]+(?:\s+[A-Z]?[a-zA-Z]+){0,3}?)\s+and\s+(?:the\s+)?'
            r'([A-Z][a-zA-Z]+(?:\s+[A-Z]?[a-zA-Z]+){0,3}?)(?=\s*[,\.\;])',
            re.MULTILINE
        )
        
        # Pattern 19: Capture "lead [role]" patterns specifically
        self.pattern_lead_role = re.compile(
            r'\b(?:The\s+)?([Ll]ead\s+[A-Z]?[a-zA-Z]+(?:\s+[A-Z]?[a-zA-Z]+){0,2}?)\s+'
            r'(?:shall|will|ensures?|provides?|is\s+responsible|has)',
            re.MULTILINE
        )
        
        # Pattern 20: "performed by the [role] and" - capture compound patterns
        self.pattern_performed_by_and = re.compile(
            r'(?:performed|conducted|reviewed|approved|verified)\s+by\s+(?:the\s+)?'
            r'([A-Z][a-zA-Z]+(?:\s+[A-Z]?[a-zA-Z]+){0,3}?)\s+and\s+(?:the\s+|approved\s+by\s+(?:the\s+)?)?'
            r'([A-Z][a-zA-Z]+(?:\s+[A-Z]?[a-zA-Z]+){0,3}?)(?=\s*[,\.\;])',
            re.MULTILINE
        )
        
        # Pattern 21: "by the [multi-word-role]." at sentence end with multi-word support
        self.pattern_by_multi_word = re.compile(
            r'by\s+(?:the\s+)?([A-Z][a-zA-Z]+(?:\s+[A-Z][a-zA-Z]+){1,4}?)\s*\.',
            re.MULTILINE
        )
        
        # Pattern 22: "as performed by the [role]"
        self.pattern_as_performed = re.compile(
            r'as\s+performed\s+by\s+(?:the\s+)?([A-Z][a-zA-Z]+(?:\s+[A-Z]?[a-zA-Z]+){0,4}?)(?=\s+and)',
            re.MULTILINE
        )
        
        # Pattern 23: "The [role] supports" / "The [role] assists"
        self.pattern_role_supports = re.compile(
            r'\b[Tt]he\s+([a-zA-Z]+(?:\s+[a-zA-Z]+){0,3}?)\s+(?:supports?|assists?)\s+(?:the\s+)?',
            re.MULTILINE
        )
        
        # Pattern 24: "approval from the [Role] and the [Role]" 
        self.pattern_approval_chain = re.compile(
            r'approval\s+from\s+(?:the\s+)?([A-Z][a-zA-Z]+(?:\s+[A-Z]?[a-zA-Z]+){0,3}?)\s+and\s+(?:the\s+)?([A-Z][a-zA-Z]+(?:\s+[A-Z]?[a-zA-Z]+){0,3}?)(?=\s*[,\.\;])',
            re.MULTILINE
        )

    def _clean_candidate(self, candidate: str) -> str:
        """Clean and normalize a candidate role string."""
        candidate = candidate.strip()

        # v3.2.5: PHONE NUMBER AND NUMERIC FILTERING
        # Reject candidates that start with digits (phone numbers, addresses, codes)
        if re.match(r'^\d', candidate):
            return ""

        # v3.2.5: Reject candidates containing phone number patterns
        # Patterns: ###-####, (###), ###.###.####, etc.
        phone_patterns = [
            r'\d{3}[-.\s]?\d{4}',  # ###-#### or ### ####
            r'\(\d{3}\)',          # (###)
            r'\d{3}[-.\s]\d{3}[-.\s]\d{4}',  # ###-###-####
        ]
        for pattern in phone_patterns:
            if re.search(pattern, candidate):
                return ""

        # v3.2.5: Reject candidates that are mostly numeric
        digit_count = sum(1 for c in candidate if c.isdigit())
        if len(candidate) > 0 and digit_count / len(candidate) > 0.3:
            return ""

        # v3.2.5: Strip inline acronyms like "(PM)" from role names
        # "Project Manager (PM)" -> "Project Manager" (acronym stored separately)
        candidate = re.sub(r'\s*\([A-Z][A-Z&/]{1,7}\)\s*', ' ', candidate).strip()

        # v3.2.5: Reject candidates with ZIP code patterns
        if re.search(r'\b\d{5}(?:-\d{4})?\b', candidate):
            return ""

        # v3.2.5: Reject run-together words (PDF extraction artifacts)
        # e.g., "Byasafetydepartment", "Thepersonnel", "Thedifference"
        if len(candidate) > 10 and ' ' not in candidate:
            # Check for lowercase letters followed by uppercase (camelCase pattern)
            if re.search(r'[a-z][A-Z]', candidate):
                return ""
            # Check for common run-together prefixes
            run_together_prefixes = ['bythe', 'bya', 'tothe', 'forthe', 'ofthe',
                                     'inthe', 'onthe', 'atthe', 'asthe', 'isthe',
                                     'thepersonnel', 'thedifference', 'theaccountable',
                                     'sufficientmanagement', 'thecoordinate', 'thecertificate']
            if any(candidate.lower().startswith(p) for p in run_together_prefixes):
                return ""

        # v3.2.5: Reject slash-separated alternatives (not roles)
        # e.g., "Owner/Manager", "Individual/Group"
        if '/' in candidate and not re.search(r'[A-Z]&[A-Z]', candidate):  # Allow IV&V
            return ""

        # v3.2.5: Reject section headers (e.g., "C. Scalability", "1.2 Overview")
        if re.match(r'^[A-Z0-9]+\.\s', candidate):
            return ""

        # v3.2.5: Reject address patterns (e.g., "Suite 670", "1111 Third")
        if re.search(r'\b(suite|floor|building|room)\s*\d+', candidate, re.IGNORECASE):
            return ""
        if re.search(r'\b\d{3,5}\s+(first|second|third|fourth|fifth|main|north|south|east|west)', candidate, re.IGNORECASE):
            return ""

        # Reject candidates with weird spacing patterns (OCR/parsing errors)
        # E.g., "Fa Ci Lities" instead of "Facilities"
        words = candidate.split()
        
        # Check for many single-letter words or 2-letter words in sequence (OCR error)
        short_word_count = sum(1 for w in words if len(w) <= 2 and w.isalpha())
        if short_word_count >= 3 and len(words) >= 4:
            return ""  # Likely an OCR error like "Fa Ci Lities Pro Ject"
        
        # Check for alternating case pattern that suggests broken text
        if len(words) >= 3:
            all_title_case = all(w[0].isupper() and (len(w) == 1 or w[1:].islower()) for w in words if w.isalpha())
            avg_word_len = sum(len(w) for w in words) / len(words)
            if all_title_case and avg_word_len < 4:
                return ""  # Likely "Fa Ci Li Ties" pattern
        
        # Remove trailing boundary words (conjunctions, prepositions, etc.)
        while words and words[-1].lower() in self.BOUNDARY_WORDS:
            words.pop()
        
        # Remove leading boundary words
        while words and words[0].lower() in self.BOUNDARY_WORDS:
            words.pop(0)
        
        # Remove leading articles
        while words and words[0].lower() in ['the', 'a', 'an']:
            words.pop(0)
        
        # Extended list of verbs that shouldn't be part of role names
        # v3.0.91b: Removed 'leads' as it can be part of role names like "Technical Leads"
        trailing_verbs = [
            # Action verbs
            'analyzes', 'assembles', 'processes', 'develops', 'creates',
            'builds', 'designs', 'tests', 'integrates', 'works', 'assists',
            'monitors', 'reviews', 'approves', 'verifies', 'validates',
            'analyzed', 'assembled', 'processed', 'developed', 'created',
            'performs', 'coordinates', 'manages', 'ensures', 'provides',
            'maintains', 'oversees', 'supports', 'implements',
            # Gerunds
            'analyzing', 'assembling', 'processing', 'developing', 'creating',
            'building', 'designing', 'testing', 'integrating', 'working',
            'assisting', 'monitoring', 'reviewing', 'approving', 'verifying',
            # Object words that indicate overextended extraction
            'system', 'reliability', 'components', 'data', 'requirements',
            'specifications', 'documents', 'activities', 'functions',
            'interfaces', 'design', 'analysis', 'report', 'documentation'
        ]
        
        # v3.0.91b: Don't strip 'leads' if preceded by a role modifier (it's part of role name)
        # "Technical Leads" should stay intact, but "The engineer leads" should lose "leads"
        if words and words[-1].lower() == 'leads' and len(words) >= 2:
            # Check if previous word suggests this is a role name
            role_modifiers = {'technical', 'team', 'project', 'program', 'system', 'software',
                             'hardware', 'test', 'qa', 'quality', 'safety', 'design', 'integration'}
            if words[-2].lower() not in role_modifiers:
                words.pop()  # Remove 'leads' if it's likely a verb
        
        while words and words[-1].lower() in trailing_verbs:
            words.pop()
        
        # Also remove leading verbs (shouldn't start a role name)
        leading_verbs = [
            'shall', 'will', 'may', 'can', 'should', 'must',
            'performs', 'coordinates', 'manages', 'ensures', 'provides',
            'reviews', 'approves', 'verifies', 'validates', 'monitors'
        ]
        while words and words[0].lower() in leading_verbs:
            words.pop(0)
        
        # If result is too long or contains verb phrases, try to truncate
        result = ' '.join(words)
        
        # Check for known patterns where extraction went too far
        verb_phrases = ['analyzes system', 'assembles components', 'processes data',
                       'develops requirements', 'creates documents', 'performs analysis',
                       'conducts review', 'provides support', 'manages project']
        for vp in verb_phrases:
            if vp in result.lower():
                # Truncate at the verb
                verb = vp.split()[0]
                idx = result.lower().find(verb)
                if idx > 0:
                    result = result[:idx].strip()
        
        return result

    def _is_valid_role(self, candidate: str) -> Tuple[bool, float]:
        """Determine if a candidate string is likely a valid role."""
        candidate = self._clean_candidate(candidate)
        candidate_lower = candidate.lower().strip()
        
        # v3.0.91b: Check for known acronyms FIRST (before length check)
        # These are short but valid role identifiers
        valid_acronyms = {'cor', 'pm', 'se', 'lse', 'ipt', 'ccb', 'erb', 'ta', 'pi',
                         'qa', 'qae', 'cm', 'sme', 'der', 'dar', 'gtr', 'co', 'cotr',
                         # Executive acronyms
                         'cino', 'cto', 'cio', 'ceo', 'coo', 'cfo', 'pgm', 'dpgm',
                         # IT roles
                         'dba', 'sa', 'sysadmin'}
        if candidate_lower in valid_acronyms:
            return True, 0.92

        # v3.3.0: Check for worker/employee terms EARLY (these are always valid roles)
        # These generic terms are critical for OSHA and safety documents
        worker_terms = {
            'employer', 'employers', 'employee', 'employees',
            'personnel', 'workers', 'worker', 'staff',
            'dispatchers', 'dispatcher', 'operator', 'operators'
        }
        if candidate_lower in worker_terms:
            return True, 0.85

        # v3.3.2: Check for defense/government terms EARLY
        # These are valid roles in MIL-STD and government contract documents
        defense_terms = {
            'government', 'contractor', 'subcontractor', 'vendor', 'supplier',
            'user', 'users', 'maintainer', 'maintainers',
            'technician', 'technicians', 'inspector', 'inspectors',
            'custodian', 'illustrator', 'editor', 'author', 'authors',
            'manager', 'managers', 'owner', 'owners',
            'quality control', 'senior management', 'information owner',
            # v3.3.3: Aerospace/aviation roles
            'lead', 'leads', 'pilot', 'pilots', 'engineer', 'engineers'
        }
        if candidate_lower in defense_terms:
            return True, 0.88

        # v3.3.0: Check for academic roles EARLY
        academic_terms = {
            'graduate student', 'graduate students',
            'postdoctoral researcher', 'postdoctoral researchers',
            'research staff', 'lab supervisor', 'laboratory supervisor',
            'postdoc', 'postdocs', 'faculty member', 'faculty members'
        }
        if candidate_lower in academic_terms:
            return True, 0.88

        # Basic checks
        if len(candidate_lower) < 4 or len(candidate.split()) < 1:
            return False, 0.0

        if len(candidate.split()) > 6:
            return False, 0.0

        # v3.2.5: Reject excessively long role names (likely text fragments)
        if len(candidate) > 60:
            return False, 0.0

        # v3.2.5: Reject if candidate contains phone patterns that slipped through
        if re.search(r'\d{3}[-.\s]?\d{4}', candidate):
            return False, 0.0
        
        # v3.0.91b: Reject single words that are never roles
        words = candidate_lower.split()
        if len(words) == 1 and candidate_lower in self.SINGLE_WORD_EXCLUSIONS:
            return False, 0.0
        
        # v3.0.91b: Reject phrases starting with noise words
        noise_starters = {
            'the', 'a', 'an', 'this', 'that', 'all', 'any', 'each', 'some',
            'contract', 'provide', 'ensure', 'facilitate', 'complete',
            'develop', 'perform', 'operate', 'coordinate', 'manage',
            'services', 'requirements', 'process', 'phase', 'work',
            'idiq', 'task', 'data', 'report', 'plan', 'end', 'wide',
            'construction', 'calibration', 'demolition', 'repair',
            'manner', 'range', 'monthly', 'foreign', 'addition',
            'written', 'final', 'initial', 'current', 'other', 'various',
            # v3.0.91b: Additional noise starters
            'responsible', 'accountable', 'serves', 'serving', 'acts',
            'acting', 'works', 'working', 'reports', 'reporting',
            'directly', 'overall', 'primary', 'secondary', 'main'
        }
        if words and words[0] in noise_starters:
            return False, 0.0
        
        # v3.0.91b: Reject phrases containing connector words in positions 2-4
        connector_words = {'is', 'are', 'was', 'were', 'be', 'being', 'been',
                          'shall', 'will', 'must', 'may', 'can', 'could',
                          'that', 'which', 'who', 'whom', 'whose',
                          'so', 'such', 'very', 'just', 'only', 'also',
                          'including', 'excluding', 'regarding', 'concerning',
                          'for'}  # Added 'for' to catch "Responsible for..."
        if len(words) > 1:
            for word in words[1:4]:  # Check words 2-4
                if word in connector_words:
                    return False, 0.0
        
        # v3.0.91b: Reject phrases with "and" followed by another role word
        # e.g., "COR And Contracting Officer" should be split, not one role
        if 'and' in words:
            and_idx = words.index('and')
            # If there's a role-like word after "and", this is likely two roles combined
            if and_idx > 0 and and_idx < len(words) - 1:
                after_and = ' '.join(words[and_idx+1:])
                # Check if the part after "and" looks like a role
                role_suffixes = {'manager', 'engineer', 'officer', 'director', 'lead', 
                                'analyst', 'team', 'coordinator', 'specialist'}
                if any(s in after_and for s in role_suffixes):
                    return False, 0.0
        
        # v3.0.91b: Reject "X The Y" patterns (sentence fragments)
        if len(words) >= 3 and 'the' in words[1:]:
            the_idx = words.index('the') if 'the' in words else -1
            if the_idx > 0:  # "the" is not at start
                return False, 0.0
        
        # v3.0.91b: Reject phrases ending with noise words
        noise_endings = {'begins', 'ends', 'various', 'just', 'only',
                        'personnel', 'activities', 'services', 'requirements',
                        'orders', 'systems', 'facilities', 'resources'}
        if words and words[-1] in noise_endings:
            return False, 0.0

        # v3.2.5: Reject location/address patterns
        location_words = {'center', 'building', 'suite', 'floor', 'room', 'street',
                         'avenue', 'boulevard', 'road', 'drive', 'lane', 'way',
                         'atlanta', 'chicago', 'washington', 'boston', 'denver',
                         'alaska', 'california', 'connecticut', 'hawaii', 'indiana',
                         'iowa', 'federal'}
        if len(words) >= 2:
            # Check if multiple location words or location + number
            location_count = sum(1 for w in words if w in location_words)
            has_number = any(w.isdigit() or re.match(r'^\d+$', w) for w in words)
            if location_count >= 2 or (location_count >= 1 and has_number):
                # But allow "Federal Aviation Administrator" type roles
                if not any(w in ['administrator', 'director', 'manager', 'officer'] for w in words):
                    return False, 0.0

        # v3.2.5: Reject "Other X" patterns (e.g., "Other Senior Managers")
        if words and words[0] == 'other':
            return False, 0.0

        # v3.2.5: Reject "Own X" patterns (e.g., "Own Data Analysis Group")
        if words and words[0] == 'own':
            return False, 0.0

        # v3.0.91d: Check EXPLICIT false positives FIRST - overrides everything
        # This ensures items like "Verification Engineer" or "Mission Assurance" are filtered
        # even if they appear in known_roles
        if candidate_lower in self.false_positives:
            return False, 0.0
        
        # v3.0.12b: Check for known roles (high confidence)
        # This prevents false-positive rules from rejecting real roles like "Program Manager"
        if candidate_lower in self.known_roles:
            return True, 0.95
        
        # v3.0.12b: Check for role suffix BEFORE generic false-positive patterns
        # Strong role suffixes override generic false-positive patterns
        strong_role_suffixes = ['engineer', 'manager', 'lead', 'analyst', 'specialist', 
                               'coordinator', 'owner', 'reviewer', 'approver', 'author',
                               'director', 'officer', 'supervisor', 'architect']
        for suffix in strong_role_suffixes:
            if candidate_lower.endswith(suffix):
                prefix = candidate_lower[:-len(suffix)].strip()
                if prefix:  # Has a modifier before the suffix
                    return True, 0.90
        
        # Check against partial false positive matches
        for fp in self.false_positives:
            if candidate_lower.startswith(fp + ' ') or candidate_lower.endswith(' ' + fp):
                return False, 0.0
        
        # Check for problematic patterns
        bad_patterns = ['reviewed by', 'approved by', 'conducted by', 'coordinated with',
                       'performed by', 'verified by', 'validated by', 'assists the',
                       'works with the', 'in managing', 'goes beyond', 'perform various',
                       'that optimizes', 'completion of', 'demand for']
        if any(bp in candidate_lower for bp in bad_patterns):
            return False, 0.0
        
        # Check for fragments that don't make sense as roles
        if candidate_lower.startswith('manager assists') or candidate_lower.startswith('engineer analyzes'):
            return False, 0.0
        
        # Single word "Lead" or "Manager" alone is too generic unless it's a known role
        if candidate_lower in ['lead', 'manager', 'engineer', 'director', 'analyst']:
            return False, 0.0
        
        # Check for activity endings
        activity_endings = ['activities', 'reviews', 'specifications', 'requirements', 
                          'procedures', 'processes', 'tasks', 'efforts', 'orders',
                          'travel', 'maintenance', 'protection', 'facility']
        for ending in activity_endings:
            if candidate_lower.endswith(ending):
                # Exception: allow if it's clearly a role
                if not any(s in candidate_lower for s in ['engineer', 'manager', 'lead', 'officer']):
                    return False, 0.0
        
        # Check for role suffixes (general)
        for suffix in self.ROLE_SUFFIXES:
            if candidate_lower.endswith(suffix):
                prefix = candidate_lower[:-len(suffix)].strip()
                if prefix:
                    has_valid_modifier = any(
                        word.lower() in [m.lower() for m in self.ROLE_MODIFIERS]
                        for word in prefix.split()
                    )
                    if has_valid_modifier:
                        return True, 0.90
                    elif len(prefix.split()) <= 2:
                        return True, 0.75
                    else:
                        return True, 0.60
                else:
                    return True, 0.50
        
        # Check for role suffix in middle of phrase
        for word in words:
            if word in self.ROLE_SUFFIXES:
                return True, 0.65
        
        # Check for modifier + something pattern
        if len(words) >= 2 and words[0] in [m.lower() for m in self.ROLE_MODIFIERS]:
            return True, 0.55
        
        # Check for partial match with known roles
        for known in self.known_roles:
            if known in candidate_lower and len(candidate_lower) < len(known) + 15:
                return True, 0.70
        
        return False, 0.0
    
    def _is_deliverable(self, candidate: str) -> bool:
        """Check if a candidate is likely a deliverable/work product rather than a role."""
        candidate_lower = candidate.lower()
        
        # Check against deliverable patterns
        for pattern in self.DELIVERABLE_PATTERNS:
            if re.search(pattern, candidate_lower, re.IGNORECASE):
                return True
        
        # Check for common deliverable naming patterns
        deliverable_indicators = [
            'plan', 'report', 'specification', 'document', 'manual',
            'drawing', 'schematic', 'schedule', 'budget', 'proposal',
            'matrix', 'database', 'archive', 'log', 'checklist'
        ]
        
        words = candidate_lower.split()
        if words and words[-1] in deliverable_indicators:
            return True
        
        return False
    
    def classify_extraction(self, candidate: str) -> dict:
        """
        Classify an extraction as role, deliverable, or unknown.
        Returns dict with type, confidence, and reasoning.
        
        v3.0.12b: Role suffix wins tie-break. If candidate ends with a strong
        role suffix (engineer, manager, lead, etc.), classify as role even if
        deliverable keywords are present.
        
        v3.0.105: BUG-004 FIX - Strong role suffix now ALWAYS wins tiebreak,
        regardless of _is_valid_role result. This ensures "Report Engineer"
        is classified as role, not deliverable.
        """
        candidate_lower = candidate.lower().strip()
        words = candidate_lower.split()
        
        # v3.0.12b: Strong role suffixes that win tie-break against deliverable keywords
        strong_role_suffixes = {'engineer', 'manager', 'lead', 'analyst', 'specialist', 
                               'coordinator', 'owner', 'reviewer', 'approver', 'author',
                               'director', 'officer', 'supervisor', 'architect', 'sme',
                               'integrator', 'administrator', 'technician', 'inspector'}
        
        # Check if ends with strong role suffix - this wins tie-break
        has_strong_role_suffix = words and words[-1] in strong_role_suffixes
        
        # v3.0.105 FIX: If has strong role suffix, classify as role IMMEDIATELY
        # Don't rely on _is_valid_role which may have false negatives
        if has_strong_role_suffix:
            # Still try to get confidence from _is_valid_role if possible
            is_valid, confidence = self._is_valid_role(candidate)
            # Use higher confidence if valid, otherwise use base 0.9 for suffix match
            effective_confidence = max(confidence, 0.9) if is_valid else 0.9
            return {
                'type': 'role',
                'confidence': effective_confidence,
                'reason': f'Strong role suffix ({words[-1]}) wins tie-break'
            }
        
        # Check if it's a deliverable (only if no strong role suffix)
        if self._is_deliverable(candidate):
            return {
                'type': 'deliverable',
                'confidence': 0.85,
                'reason': 'Matches deliverable pattern'
            }
        
        # Check if it's a valid role
        is_valid, confidence = self._is_valid_role(candidate)
        if is_valid:
            return {
                'type': 'role',
                'confidence': confidence,
                'reason': 'Matches role pattern'
            }
        
        # Check if it's in false positives
        if candidate_lower in self.false_positives:
            return {
                'type': 'false_positive',
                'confidence': 0.9,
                'reason': 'Known false positive'
            }
        
        return {
            'type': 'unknown',
            'confidence': 0.3,
            'reason': 'Could not classify'
        }
    
    def _create_extracted_role(self, canonical_name: str, original_text: str = None) -> ExtractedRole:
        """Create an ExtractedRole with entity classification.
        
        v3.0.12: All role creation goes through this method to ensure
        entity_kind is always populated.
        """
        text_to_classify = original_text or canonical_name
        classification = self.classify_extraction(text_to_classify)
        
        # Map classification type to EntityKind enum
        kind_map = {
            'role': EntityKind.ROLE,
            'deliverable': EntityKind.DELIVERABLE,
            'false_positive': EntityKind.UNKNOWN,
            'unknown': EntityKind.UNKNOWN
        }
        entity_kind = kind_map.get(classification['type'], EntityKind.UNKNOWN)
        
        return ExtractedRole(
            canonical_name=canonical_name,
            entity_kind=entity_kind,
            kind_confidence=classification['confidence'],
            kind_reason=classification['reason']
        )
    
    def _extract_responsibility(self, text: str, role_match: str) -> Tuple[str, str]:
        """Extract the responsibility associated with a role from context."""
        text_lower = text.lower()
        role_lower = role_match.lower()
        
        role_pos = text_lower.find(role_lower)
        if role_pos == -1:
            return "", "unknown"
        
        after_role = text[role_pos + len(role_match):].strip()
        action_type = "unknown"
        responsibility = ""
        
        for atype, verbs in self.ACTION_VERBS.items():
            for verb in verbs:
                verb_match = re.search(rf'\b{verb}\b', after_role[:100], re.IGNORECASE)
                if verb_match:
                    action_type = atype
                    verb_pos = verb_match.end()
                    remainder = after_role[verb_pos:].strip()
                    end_match = re.search(r'[\.;]', remainder)
                    if end_match:
                        responsibility = remainder[:end_match.start()].strip()
                    else:
                        responsibility = remainder[:100].strip()
                    break
            if action_type != "unknown":
                break
        
        if action_type == "unknown":
            resp_match = re.search(r'is\s+responsible\s+for\s+(.{10,100}?)(?:\.|;|$)', after_role, re.IGNORECASE)
            if resp_match:
                action_type = "performs"
                responsibility = resp_match.group(1).strip()
        
        return responsibility, action_type
    
    def _normalize_role(self, role: str) -> str:
        """Normalize a role name to a canonical form."""
        role = self._clean_candidate(role)
        normalized = ' '.join(role.split()).title()
        
        # Keep abbreviations uppercase
        abbreviations = {
            'Pm': 'PM', 'Se': 'SE', 'Lse': 'LSE', 'Ipt': 'IPT', 'Cor': 'COR',
            'Gtr': 'GTR', 'Der': 'DER', 'Dar': 'DAR', 'Ccb': 'CCB', 'Erb': 'ERB',
            'Trb': 'TRB', 'Irt': 'IRT', 'Srb': 'SRB', 'Oce': 'OCE', 'Mdaa': 'MDAA',
            'Qa': 'QA', 'Ma': 'MA', 'Sma': 'SMA', 'It': 'IT', 'Iv&V': 'IV&V',
            'Ivv': 'IVV', 'Pp&C': 'PP&C', 'Icwg': 'ICWG'
        }
        
        for abbr, correct in abbreviations.items():
            normalized = re.sub(rf'\b{abbr}\b', correct, normalized)
        
        return normalized
    
    # Role aliases for merging similar roles
    ROLE_ALIASES = {
        # Singular/plural and minor variations
        'Systems Engineer': ['System Engineer', 'Systems Engineers', 'System Engineers', 'Sys Engineer'],
        'Software Engineer': ['Software Engineers', 'SW Engineer', 'SW Engineers'],
        'Hardware Engineer': ['Hardware Engineers', 'HW Engineer', 'HW Engineers'],
        'Quality Assurance': ['Quality Assurance Engineer', 'QA Engineer', 'QA', 'Quality Engineer'],
        'Project Manager': ['Project Managers', 'Proj Manager', 'PM'],
        'Program Manager': ['Program Managers', 'Prog Manager'],
        'Test Engineer': ['Test Engineers', 'Testing Engineer', 'Test Eng'],
        'Safety Engineer': ['Safety Engineers', 'System Safety Engineer', 'Systems Safety Engineer'],
        'Reliability Engineer': ['Reliability Engineers', 'Rel Engineer'],
        'Integration Engineer': ['Integration Engineers', 'Integrator', 'System Integrator'],
        'Configuration Manager': ['Configuration Management', 'CM', 'Config Manager'],
        'Technical Lead': ['Tech Lead', 'Technical Leads', 'Tech Leads'],
        'Chief Engineer': ['Chief Engineers', 'CE', 'Chief Eng'],
        'Lead Engineer': ['Lead Engineers', 'Lead Eng'],
        'Requirements Engineer': ['Requirements Engineers', 'Req Engineer', 'Requirements Analyst'],
        'Design Engineer': ['Design Engineers', 'Designer'],
        'Verification Engineer': ['Verification Engineers', 'V&V Engineer', 'Verification'],
        'Validation Engineer': ['Validation Engineers', 'Validation'],
        'Subcontractor': ['Subcontractors', 'Sub-Contractor', 'Sub Contractor'],
        'Contractor': ['Contractors', 'Prime Contractor'],
        'Customer': ['Customers', 'Client', 'End User'],
        'Government': ['Government Representative', 'Govt', 'Government Customer'],
    }
    
    def _get_canonical_role(self, role_name: str) -> Tuple[str, str]:
        """
        Get the canonical name for a role, merging aliases.
        Returns (canonical_name, original_variant).
        """
        normalized = self._normalize_role(role_name)
        normalized_lower = normalized.lower()
        
        # Check if this role matches any alias
        for canonical, aliases in self.ROLE_ALIASES.items():
            if normalized_lower == canonical.lower():
                return canonical, normalized
            for alias in aliases:
                if normalized_lower == alias.lower():
                    return canonical, normalized
        
        # Check for partial matches (e.g., "System Engineer" should match "Systems Engineer")
        for canonical, aliases in self.ROLE_ALIASES.items():
            # Check similarity - simple approach: compare without 's' suffix
            canon_base = canonical.lower().replace('systems', 'system').replace('engineers', 'engineer')
            norm_base = normalized_lower.replace('systems', 'system').replace('engineers', 'engineer')
            if canon_base == norm_base:
                return canonical, normalized
        
        return normalized, normalized
    
    def _get_sentence_context(self, text: str, match_start: int, match_end: int) -> str:
        """Extract the sentence containing the match."""
        sentence_start = max(0, text.rfind('.', 0, match_start) + 1)
        sentence_end = text.find('.', match_end)
        if sentence_end == -1:
            sentence_end = len(text)
        else:
            sentence_end += 1
        
        context = text[sentence_start:sentence_end].strip()
        return ' '.join(context.split())
    
    def extract_from_text(self, text: str, source_location: str = "unknown") -> Dict[str, ExtractedRole]:
        """Extract roles from plain text.

        v3.4.0: In STRICT mode, uses whitelist-only extraction for 100% accuracy.
        In DISCOVERY mode (default), uses pattern-based extraction for maximum recall.
        """
        # v3.4.0: Use strict extraction if in STRICT mode
        if self.extraction_mode == ExtractionMode.STRICT:
            return self._extract_strict_mode(text, source_location)

        # DISCOVERY mode: Original pattern-based extraction
        extracted_roles: Dict[str, ExtractedRole] = {}
        seen_matches = set()

        patterns = [
            ('the_role', self.pattern_the_role, False),
            ('role_is', self.pattern_role_is, False),
            ('by_role', self.pattern_by_role, False),
            ('action_by', self.pattern_action_by, False),
            ('possessive', self.pattern_possessive, False),
            ('responsibilities_of', self.pattern_responsibilities_of, False),
            ('acronym', self.pattern_acronym, True),
            ('as_role', self.pattern_as_role, False),
            ('notify', self.pattern_notify, False),
            ('sentence_start', self.pattern_sentence_start, False),
            ('end_of_sentence', self.pattern_end_of_sentence, False),
            ('role_and', self.pattern_role_and, True),
            ('from_role', self.pattern_from_role, False),
            ('role_assists', self.pattern_role_assists, False),
            ('conducted_by', self.pattern_conducted_by, False),
            ('role_action_object', self.pattern_role_action_object, False),
            ('role_in', self.pattern_role_in, False),
            ('approval_from', self.pattern_approval_from, True),
            ('lead_role', self.pattern_lead_role, False),
            ('performed_by_and', self.pattern_performed_by_and, True),
            ('by_multi_word', self.pattern_by_multi_word, False),
            ('as_performed', self.pattern_as_performed, False),
            ('role_supports', self.pattern_role_supports, False),
            ('approval_chain', self.pattern_approval_chain, True)
        ]
        
        for pattern_name, pattern, has_multiple_groups in patterns:
            for match in pattern.finditer(text):
                candidates = []
                
                if pattern_name == 'acronym':
                    candidates.append((match.group(1), match.group(2)))
                elif pattern_name == 'role_and':
                    candidates.append((match.group(1), None))
                    candidates.append((match.group(2), None))
                else:
                    candidates.append((match.group(1), None))
                
                for candidate, acronym in candidates:
                    candidate = self._clean_candidate(candidate)
                    
                    if not candidate:
                        continue
                    
                    is_valid, confidence = self._is_valid_role(candidate)
                    
                    if not is_valid:
                        continue
                    
                    # Get canonical name and original variant
                    canonical, variant = self._get_canonical_role(candidate)
                    
                    match_key = (canonical, match.start())
                    if match_key in seen_matches:
                        continue
                    seen_matches.add(match_key)
                    
                    context = self._get_sentence_context(text, match.start(), match.end())
                    responsibility, action_type = self._extract_responsibility(context, candidate)
                    
                    occurrence = RoleOccurrence(
                        role=candidate,
                        context=context,
                        responsibility=responsibility,
                        action_type=action_type,
                        location=source_location,
                        confidence=confidence
                    )
                    
                    if canonical not in extracted_roles:
                        # v3.0.12: Use helper to ensure entity_kind is populated
                        extracted_roles[canonical] = self._create_extracted_role(canonical, candidate)
                    
                    role_entry = extracted_roles[canonical]
                    role_entry.variants.add(variant)
                    if variant != canonical:
                        role_entry.variants.add(candidate)  # Also add original text
                    role_entry.occurrences.append(occurrence)
                    
                    if responsibility:
                        role_entry.responsibilities.append(responsibility)
                    
                    role_entry.action_types[action_type] += 1
                    
                    if acronym:
                        role_entry.variants.add(acronym)
        
        # ADDITIONAL PASS: Direct scan for known roles that might have been missed
        # This catches cases like "Systems Engineer" in unusual sentence structures
        extracted_roles = self._scan_for_known_roles(text, extracted_roles, seen_matches, source_location)
        
        # v2.9.2 A9: Parse formal "Responsibilities" sections for enhanced role extraction
        extracted_roles = self._parse_responsibilities_sections(text, extracted_roles, source_location)
        
        # v3.1.2 ENH-008: NLP Enhancement pass for additional role detection
        if self._nlp_processor:
            extracted_roles = self._apply_nlp_enhancement(text, extracted_roles, seen_matches, source_location)

        # =====================================================================
        # v3.3.0: ENHANCED NLP EXTRACTION (Maximum Accuracy)
        # =====================================================================
        # Apply enhanced NLP with EntityRuler, PhraseMatcher, and transformer models
        # Also apply adaptive learning for confidence boosting
        if self._enhanced_nlp and self._enhanced_nlp.is_available:
            extracted_roles = self._apply_v330_enhancement(text, extracted_roles, seen_matches, source_location)

        # =====================================================================
        # v3.1.10: ROLE VERIFICATION - Filter out hallucinated roles
        # =====================================================================
        # Verify that each role can actually be found in the source text
        # This prevents counting roles that were inferred but don't appear verbatim
        extracted_roles = self._verify_roles_in_text(text, extracted_roles)

        # =====================================================================
        # v3.5.0: POST-VERIFICATION FILTERS
        # =====================================================================

        # Filter 1: Remove organization entities (NASA, Government, etc.)
        org_removed = 0
        for role_name in list(extracted_roles.keys()):
            if role_name.lower() in self._organization_entities:
                del extracted_roles[role_name]
                org_removed += 1
        if org_removed > 0:
            _log(f"Removed {org_removed} organization entities from results", level='info')

        # Filter 2: Remove low-confidence roles (avg_confidence < 0.4)
        low_conf_removed = 0
        for role_name in list(extracted_roles.keys()):
            role = extracted_roles[role_name]
            if role.avg_confidence < 0.4:
                del extracted_roles[role_name]
                low_conf_removed += 1
                _log(f"Removed low-confidence role: {role_name} (conf={role.avg_confidence:.2f})", level='debug')
        if low_conf_removed > 0:
            _log(f"Removed {low_conf_removed} low-confidence roles (< 0.4)", level='info')

        # Filter 3: Remove single-word roles that are in SINGLE_WORD_EXCLUSIONS
        stopword_removed = 0
        for role_name in list(extracted_roles.keys()):
            if len(role_name.split()) == 1 and role_name.lower() in self.SINGLE_WORD_EXCLUSIONS:
                del extracted_roles[role_name]
                stopword_removed += 1
        if stopword_removed > 0:
            _log(f"Removed {stopword_removed} single-word stopword roles", level='info')

        # Filter 4: Remove low-frequency discovered single-word roles
        # In DISCOVERY mode, single-word roles found by NLP (not in KNOWN_ROLES or dictionary)
        # with frequency < 3 are very likely false positives (acronyms, nouns, fragments)
        low_freq_removed = 0
        for role_name in list(extracted_roles.keys()):
            role = extracted_roles[role_name]
            if (len(role_name.split()) == 1
                and role.role_source == RoleSource.DISCOVERED
                and role.frequency < 3):
                del extracted_roles[role_name]
                low_freq_removed += 1
        if low_freq_removed > 0:
            _log(f"Removed {low_freq_removed} low-frequency discovered single-word roles", level='info')

        return extracted_roles

    def _extract_strict_mode(self, text: str, source_location: str = "unknown") -> Dict[str, ExtractedRole]:
        """
        v3.4.0: STRICT MODE EXTRACTION - 100% accuracy guarantee.

        Only extracts roles that:
        1. Are in the KNOWN_ROLES whitelist (or user-added dictionary roles)
        2. Appear as exact text matches in the document
        3. Each occurrence is tracked with exact character positions

        This trades recall for precision - may miss some roles,
        but every role returned is 100% verifiable in the source document.
        """
        _log(f"STRICT mode extraction for {source_location}", level='info')
        extracted_roles: Dict[str, ExtractedRole] = {}
        text_lower = text.lower()

        # Only search for known whitelisted roles
        for known_role in self.known_roles:
            # Skip if in false positives
            if known_role in self.false_positives:
                continue

            # v3.5.0: Skip organization entities
            if known_role in self._organization_entities:
                continue

            # Skip very short entries that might cause false matches
            if len(known_role) < 4:
                continue

            # Find all exact occurrences
            start = 0
            occurrences = []
            while True:
                pos = text_lower.find(known_role, start)
                if pos == -1:
                    break

                # Verify word boundaries (not part of a larger word)
                before_ok = pos == 0 or not text_lower[pos-1].isalnum()
                after_pos = pos + len(known_role)
                after_ok = after_pos >= len(text_lower) or not text_lower[after_pos].isalnum()

                if before_ok and after_ok:
                    # Get the original case version from the text
                    original_text = text[pos:after_pos]

                    # Get surrounding context
                    context_start = max(0, pos - 100)
                    context_end = min(len(text), after_pos + 100)
                    context = text[context_start:context_end].strip()

                    # Try to extract the full sentence for responsibility
                    sentence_start = text.rfind('.', 0, pos)
                    sentence_start = 0 if sentence_start == -1 else sentence_start + 1
                    sentence_end = text.find('.', after_pos)
                    sentence_end = len(text) if sentence_end == -1 else sentence_end + 1
                    responsibility = text[sentence_start:sentence_end].strip()
                    responsibility = ' '.join(responsibility.split())[:300]  # Limit length

                    occurrences.append({
                        'position': pos,
                        'original_text': original_text,
                        'context': context,
                        'responsibility': responsibility
                    })

                start = pos + 1

            # Only add if we found actual occurrences
            if occurrences:
                # Use proper casing from first occurrence
                canonical_name = occurrences[0]['original_text'].title()

                # v3.4.0: Check if role is from dictionary
                dictionary_roles = set()
                try:
                    dict_roles = self._load_dictionary_roles()
                    if dict_roles:
                        dictionary_roles = set(r.lower() for r in dict_roles)
                except Exception:
                    pass

                role_source = RoleSource.DICTIONARY if known_role in dictionary_roles else RoleSource.KNOWN

                if canonical_name not in extracted_roles:
                    extracted_roles[canonical_name] = ExtractedRole(
                        canonical_name=canonical_name,
                        entity_kind=EntityKind.ROLE,
                        kind_confidence=1.0,  # 100% confidence in STRICT mode
                        kind_reason='STRICT mode: exact whitelist match verified in text',
                        role_source=role_source,
                        is_verified_in_text=True,
                        text_positions=[occ['position'] for occ in occurrences]
                    )

                role_data = extracted_roles[canonical_name]
                role_data.variants.add(known_role)

                for occ in occurrences:
                    role_data.occurrences.append(RoleOccurrence(
                        role=canonical_name,
                        context=occ['context'],
                        responsibility=occ['responsibility'],
                        action_type='mentioned',  # Neutral action type for STRICT mode
                        location=f"{source_location}:char_{occ['position']}",
                        confidence=1.0  # 100% confident - we verified the exact text
                    ))

                    # Add responsibility if meaningful
                    if occ['responsibility'] and len(occ['responsibility']) > 20:
                        if occ['responsibility'] not in role_data.responsibilities:
                            role_data.responsibilities.append(occ['responsibility'])

        _log(f"STRICT mode found {len(extracted_roles)} verified roles in {source_location}", level='info')
        return extracted_roles

    def _verify_roles_in_text(self, text: str, roles: Dict[str, ExtractedRole]) -> Dict[str, ExtractedRole]:
        """
        Verify that extracted roles actually appear in the source text.

        v3.1.10: Prevents hallucinated roles from being counted.
        v3.4.0: Enhanced to set role_source, is_verified_in_text, and text_positions.

        A role is considered verified if:
        1. The exact canonical name appears in the text (case-insensitive)
        2. Any of its variants appear in the text
        3. The role words appear adjacently (e.g., "Validation" and "Engineer" near each other)

        Unverified roles are removed. Verified roles are tagged with their source.
        """
        if not text or not roles:
            return roles

        text_lower = text.lower()
        verified_roles = {}
        removed_count = 0

        # v3.4.0: Load dictionary roles for source classification
        dictionary_roles = set()
        try:
            dict_roles = self._load_dictionary_roles()
            if dict_roles:
                dictionary_roles = set(r.lower() for r in dict_roles)
        except Exception:
            pass  # Dictionary not available, continue without it

        for role_name, role_data in roles.items():
            is_verified = False
            verified_count = 0
            found_positions = []

            # Check 1: Exact canonical name match
            role_lower = role_name.lower()
            start = 0
            while True:
                pos = text_lower.find(role_lower, start)
                if pos == -1:
                    break
                # Verify word boundaries
                before_ok = pos == 0 or not text_lower[pos-1].isalnum()
                after_pos = pos + len(role_lower)
                after_ok = after_pos >= len(text_lower) or not text_lower[after_pos].isalnum()
                if before_ok and after_ok:
                    is_verified = True
                    found_positions.append(pos)
                start = pos + 1

            # Check 2: Any variant matches
            # v3.5.0: Skip single-word variants for multi-word roles to prevent
            # false positives (e.g. "Validation" verifying "Validation Engineer")
            if not is_verified:
                role_word_count = len(role_name.split())
                for variant in role_data.variants:
                    variant_lower = variant.lower()
                    variant_word_count = len(variant_lower.split())
                    # Don't let a single-word variant verify a multi-word role
                    if role_word_count >= 2 and variant_word_count == 1:
                        continue
                    # Skip variants that are common English words / stopwords
                    if variant_lower in self.SINGLE_WORD_EXCLUSIONS:
                        continue
                    # Skip variants that are organization entities, not roles
                    if variant_lower in self._organization_entities:
                        continue
                    start = 0
                    while True:
                        pos = text_lower.find(variant_lower, start)
                        if pos == -1:
                            break
                        before_ok = pos == 0 or not text_lower[pos-1].isalnum()
                        after_pos = pos + len(variant_lower)
                        after_ok = after_pos >= len(text_lower) or not text_lower[after_pos].isalnum()
                        if before_ok and after_ok:
                            is_verified = True
                            found_positions.append(pos)
                        start = pos + 1
                    if is_verified:
                        break

            # Check 3: Adjacent words match (for compound roles like "Validation Engineer")
            # v3.4.0: Tightened to max 20 chars gap to reduce false adjacency matches
            if not is_verified and ' ' in role_name:
                words = role_name.lower().split()
                if len(words) >= 2:
                    # Check if words appear within 20 chars of each other (stricter)
                    pattern = r'\b' + r'\b.{0,20}\b'.join(re.escape(w) for w in words) + r'\b'
                    for match in re.finditer(pattern, text_lower, re.IGNORECASE):
                        is_verified = True
                        found_positions.append(match.start())

            if is_verified:
                verified_count = len(found_positions)

                # v3.4.0: Set verification fields
                role_data.is_verified_in_text = True
                role_data.text_positions = found_positions

                # v3.4.0: Determine role source
                if role_lower in dictionary_roles:
                    role_data.role_source = RoleSource.DICTIONARY
                elif role_lower in self.known_roles:
                    role_data.role_source = RoleSource.KNOWN
                else:
                    role_data.role_source = RoleSource.DISCOVERED

                # Update occurrence count to reflect verified matches only
                if verified_count > 0 and verified_count < len(role_data.occurrences):
                    role_data.occurrences = role_data.occurrences[:verified_count]

                verified_roles[role_name] = role_data
            else:
                removed_count += 1
                _log(f"Removed unverified role: {role_name} (not found in text)", level='debug')

        if removed_count > 0:
            _log(f"Verification removed {removed_count} unverified roles", level='info')

        # v3.4.0: Log summary by source
        dict_count = sum(1 for r in verified_roles.values() if r.role_source == RoleSource.DICTIONARY)
        known_count = sum(1 for r in verified_roles.values() if r.role_source == RoleSource.KNOWN)
        discovered_count = sum(1 for r in verified_roles.values() if r.role_source == RoleSource.DISCOVERED)
        _log(f"Verified roles: {dict_count} from dictionary, {known_count} from known list, {discovered_count} discovered", level='info')

        return verified_roles
    
    def _apply_nlp_enhancement(self, text: str, existing_roles: Dict[str, ExtractedRole],
                               seen_matches: set, source_location: str) -> Dict[str, ExtractedRole]:
        """
        Apply NLP enhancement for additional role detection (v3.1.2 ENH-008).

        Uses spaCy NLP for:
        - Named Entity Recognition (NER) for role detection
        - Dependency parsing for role-action relationships
        - Better context extraction and confidence scoring

        v3.2.3: Enhanced integration with NLPProcessor v1.1.0:
        - Better confidence boosting for multi-source detection
        - Uses NLP modifiers for variant tracking
        - Improved duplicate detection using fuzzy matching
        - Lower threshold for NLP-detected roles (more accurate)
        """
        if not self._nlp_processor:
            return existing_roles

        try:
            # Extract roles using NLP processor (spaCy-powered)
            nlp_roles = self._nlp_processor.extract_roles(text)
            added_count = 0
            boosted_count = 0

            # Build a lowercase index for faster lookups
            existing_lower = {k.lower(): k for k in existing_roles.keys()}

            for nlp_role in nlp_roles:
                # Skip if already found with regex patterns
                canonical_lower = nlp_role.normalized_name.lower()

                # v3.2.3: Check for partial/fuzzy matches too
                found_existing = False
                matched_key = None

                for existing_key in existing_lower:
                    # Exact match
                    if canonical_lower == existing_key:
                        found_existing = True
                        matched_key = existing_lower[existing_key]
                        break
                    # Partial match (one contains the other)
                    if canonical_lower in existing_key or existing_key in canonical_lower:
                        if len(canonical_lower) > 5 and len(existing_key) > 5:  # Avoid short matches
                            found_existing = True
                            matched_key = existing_lower[existing_key]
                            break

                if found_existing and matched_key:
                    # Boost confidence of existing role if also found by NLP
                    # v3.2.3: Boost more for high-confidence NLP matches
                    boost = 0.15 if nlp_role.confidence > 0.8 else 0.1
                    for occ in existing_roles[matched_key].occurrences:
                        occ.confidence = min(0.98, occ.confidence + boost)
                    boosted_count += 1

                    # v3.2.3: Add modifiers as variants if available
                    if nlp_role.modifiers:
                        for mod in nlp_role.modifiers:
                            variant_name = f"{mod} {nlp_role.name}"
                            existing_roles[matched_key].variants.add(variant_name)
                    continue

                # Validate with our standard rules
                is_valid, validation_confidence = self._is_valid_role(nlp_role.name)

                # v3.2.5: Only trust high-confidence NLP if it passes basic sanity checks
                # Don't override validation for location patterns, run-together words, etc.
                if not is_valid and nlp_role.confidence >= 0.85:
                    # Check if it contains role-indicative words (manager, engineer, etc.)
                    name_lower = nlp_role.name.lower()
                    role_indicators = ['manager', 'engineer', 'director', 'officer', 'supervisor',
                                      'coordinator', 'administrator', 'inspector', 'auditor',
                                      'analyst', 'specialist', 'lead', 'chief', 'team', 'board',
                                      'committee', 'representative']
                    has_role_indicator = any(ind in name_lower for ind in role_indicators)

                    # Only trust NLP if it has a role indicator
                    if has_role_indicator:
                        is_valid = True
                        validation_confidence = nlp_role.confidence * 0.9

                if not is_valid:
                    continue

                # Use the higher of NLP confidence or validation confidence
                final_confidence = max(nlp_role.confidence, validation_confidence)

                # v3.2.3: Lower threshold for NLP since spaCy-based detection is accurate
                # Especially for verb_association and responsibility sources
                threshold = 0.45 if nlp_role.source in ['verb_association', 'responsibility', 'dependency'] else 0.50

                if final_confidence >= threshold:
                    # v3.2.5: Skip roles longer than 60 characters
                    if len(nlp_role.name) > 60:
                        continue

                    canonical_name, variant = self._get_canonical_role(nlp_role.name)

                    if canonical_name not in existing_roles:
                        existing_roles[canonical_name] = self._create_extracted_role(
                            canonical_name, nlp_role.name
                        )

                    role_entry = existing_roles[canonical_name]
                    role_entry.variants.add(nlp_role.name)

                    # v3.2.5: Extract acronym if present in context (e.g., "Project Manager (PM)")
                    if nlp_role.context:
                        acronym_match = re.search(
                            re.escape(nlp_role.name) + r'\s*\(([A-Z][A-Z&/]{1,7})\)',
                            nlp_role.context, re.IGNORECASE
                        )
                        if acronym_match:
                            acronym = acronym_match.group(1).upper()
                            role_entry.variants.add(acronym)
                            # Also add the acronym to the acronym map if not already there
                            if acronym.lower() not in self.ACRONYM_MAP:
                                self.ACRONYM_MAP[acronym.lower()] = canonical_name.lower()

                    # v3.2.3: Add modifiers as additional variants
                    if nlp_role.modifiers:
                        for mod in nlp_role.modifiers:
                            if mod.lower() not in nlp_role.name.lower():
                                variant_name = f"{mod} {nlp_role.name}"
                                role_entry.variants.add(variant_name)

                    # Determine action type based on NLP source
                    action_type = f"nlp_{nlp_role.source}"  # e.g., nlp_ner, nlp_dependency, nlp_pattern

                    # v3.2.3: Try to extract responsibility from context
                    responsibility = ""
                    if nlp_role.context:
                        responsibility, _ = self._extract_responsibility(nlp_role.context, nlp_role.name)

                    occurrence = RoleOccurrence(
                        role=nlp_role.name,
                        context=nlp_role.context[:250] if nlp_role.context else "",
                        responsibility=responsibility,
                        action_type=action_type,
                        location=source_location,
                        confidence=final_confidence
                    )
                    role_entry.occurrences.append(occurrence)
                    role_entry.action_types[action_type] += 1

                    # v3.2.3: Track responsibilities from NLP context
                    if responsibility:
                        role_entry.responsibilities.append(responsibility)

                    added_count += 1

            if added_count > 0 or boosted_count > 0:
                _log(f"NLP enhancement: added {added_count} roles, boosted {boosted_count} (from {len(nlp_roles)} candidates)", level='debug')

        except Exception as e:
            _log(f"NLP enhancement failed: {e}", level='warning')

        return existing_roles

    def _apply_v330_enhancement(self, text: str, existing_roles: Dict[str, ExtractedRole],
                                seen_matches: set, source_location: str) -> Dict[str, ExtractedRole]:
        """
        v3.3.0: Apply maximum accuracy NLP enhancement.

        Uses:
        - Enhanced NLP processor with transformer models (en_core_web_trf)
        - EntityRuler with 100+ aerospace/defense patterns
        - PhraseMatcher with 150+ role gazetteer entries
        - Adaptive learning for confidence boosting based on user decisions
        - Technical dictionary for role validation
        """
        if not self._enhanced_nlp or not self._enhanced_nlp.is_available:
            return existing_roles

        try:
            # Extract roles using enhanced NLP (EntityRuler + PhraseMatcher + transformer NER)
            enhanced_roles = self._enhanced_nlp.extract_roles(text)
            added_count = 0
            boosted_count = 0

            # Build a lowercase index for faster lookups
            existing_lower = {k.lower(): k for k in existing_roles.keys()}

            for enhanced_role in enhanced_roles:
                role_name = enhanced_role.get('name', '')
                confidence = enhanced_role.get('confidence', 0.5)
                source = enhanced_role.get('source', 'enhanced_nlp')
                context = enhanced_role.get('context', '')
                modifiers = enhanced_role.get('modifiers', [])

                if not role_name:
                    continue

                canonical_lower = role_name.lower()

                # Check for existing matches
                found_existing = False
                matched_key = None

                for existing_key in existing_lower:
                    if canonical_lower == existing_key or canonical_lower in existing_key or existing_key in canonical_lower:
                        if len(canonical_lower) > 4 and len(existing_key) > 4:
                            found_existing = True
                            matched_key = existing_lower[existing_key]
                            break

                # Apply adaptive learning boost
                learning_boost = 0.0
                if self._adaptive_learner and self._adaptive_learner.is_available:
                    learning_boost = self._adaptive_learner.get_role_confidence_boost(role_name, context)

                    # Skip roles that have been consistently rejected
                    if self._adaptive_learner.is_known_invalid_role(role_name):
                        continue

                    # Automatically accept roles that have been consistently confirmed
                    if self._adaptive_learner.is_known_valid_role(role_name):
                        confidence = max(confidence, 0.90)
                        learning_boost = max(learning_boost, 0.15)

                if found_existing and matched_key:
                    # Boost confidence of existing role
                    boost = 0.12 + learning_boost
                    for occ in existing_roles[matched_key].occurrences:
                        occ.confidence = min(0.98, occ.confidence + boost)
                    boosted_count += 1

                    # Add modifiers as variants
                    if modifiers:
                        for mod in modifiers:
                            variant_name = f"{mod} {role_name}"
                            existing_roles[matched_key].variants.add(variant_name)
                    continue

                # Validate with standard rules
                is_valid, validation_confidence = self._is_valid_role(role_name)

                # v3.3.0: Trust high-confidence EntityRuler/PhraseMatcher results more
                if not is_valid and confidence >= 0.80:
                    if source in ['entity_ruler', 'phrase_matcher', 'transformer_ner']:
                        is_valid = True
                        validation_confidence = confidence * 0.92

                if not is_valid:
                    continue

                # Apply learning boost to final confidence
                final_confidence = min(0.98, max(confidence, validation_confidence) + learning_boost)

                # Lower threshold for v3.3.0 enhanced NLP (more accurate than pattern matching)
                threshold = 0.40 if source in ['entity_ruler', 'phrase_matcher'] else 0.45

                if final_confidence >= threshold:
                    # Length validation
                    if len(role_name) > 60:
                        continue

                    canonical_name, variant = self._get_canonical_role(role_name)

                    if canonical_name not in existing_roles:
                        existing_roles[canonical_name] = self._create_extracted_role(
                            canonical_name, role_name
                        )

                    role_entry = existing_roles[canonical_name]
                    role_entry.variants.add(role_name)

                    # Add modifiers as variants
                    if modifiers:
                        for mod in modifiers:
                            if mod.lower() not in role_name.lower():
                                variant_name = f"{mod} {role_name}"
                                role_entry.variants.add(variant_name)

                    # Determine action type
                    action_type = f"v330_{source}"

                    # Extract responsibility
                    responsibility = ""
                    if context:
                        responsibility, _ = self._extract_responsibility(context, role_name)

                    occurrence = RoleOccurrence(
                        role=role_name,
                        context=context[:250] if context else "",
                        responsibility=responsibility,
                        action_type=action_type,
                        location=source_location,
                        confidence=final_confidence
                    )
                    role_entry.occurrences.append(occurrence)
                    role_entry.action_types[action_type] += 1

                    if responsibility:
                        role_entry.responsibilities.append(responsibility)

                    added_count += 1

            if added_count > 0 or boosted_count > 0:
                _log(f"v3.3.0 enhancement: added {added_count} roles, boosted {boosted_count} "
                     f"(from {len(enhanced_roles)} candidates)", level='debug')

        except Exception as e:
            _log(f"v3.3.0 enhancement failed: {e}", level='warning')

        return existing_roles

    def _parse_responsibilities_sections(self, text: str, existing_roles: Dict[str, 'ExtractedRole'],
                                         source_location: str) -> Dict[str, 'ExtractedRole']:
        """
        v2.9.2 A9: Parse formal "Responsibilities" sections that list role duties.
        
        Handles formats like:
        - "RESPONSIBILITIES" / "Roles and Responsibilities" section headers
        - "[Role Name]:" followed by bullet points or numbered items
        - Tabular formats with Role | Responsibility columns
        """
        
        # Pattern to find responsibility section headers
        section_patterns = [
            # Standard section headers
            re.compile(r'(?:^|\n)\s*(?:\d+\.?\s*)?(?:ROLES?\s+AND\s+)?RESPONSIBILITIES\s*(?:\n|$)', re.IGNORECASE | re.MULTILINE),
            re.compile(r'(?:^|\n)\s*(?:\d+\.?\s*)?ORGANIZATIONAL\s+RESPONSIBILITIES\s*(?:\n|$)', re.IGNORECASE | re.MULTILINE),
            re.compile(r'(?:^|\n)\s*(?:\d+\.?\s*)?RESPONSIBILITY\s+MATRIX\s*(?:\n|$)', re.IGNORECASE | re.MULTILINE),
        ]
        
        # Find start of responsibilities section
        section_start = None
        for pattern in section_patterns:
            match = pattern.search(text)
            if match:
                section_start = match.end()
                break
        
        if section_start is None:
            return existing_roles
        
        # Find end of section (next major heading or end of text)
        section_end_pattern = re.compile(r'\n\s*(?:\d+\.?\s*)?[A-Z][A-Z\s]{5,}(?:\n|$)')
        end_match = section_end_pattern.search(text, section_start + 100)  # Skip at least 100 chars
        section_end = end_match.start() if end_match else len(text)
        
        section_text = text[section_start:section_end]
        
        # Pattern for role followed by responsibilities
        # Format: "Role Name:" or "Role Name -" followed by text
        role_duty_pattern = re.compile(
            r'(?:^|\n)\s*(?:•|\*|[-–—]|\d+[.\)]\s*)?'  # Optional bullet/number
            r'([A-Z][A-Za-z]+(?:\s+[A-Z]?[a-zA-Z]+){0,4})'  # Role name (1-5 words, starts with caps)
            r'\s*(?::|[-–—]|shall|will|is responsible)\s*'  # Separator
            r'(.{10,200}?)(?=\n|$)',  # Responsibility text (10-200 chars)
            re.MULTILINE
        )
        
        for match in role_duty_pattern.finditer(section_text):
            role_text = match.group(1).strip()
            responsibility = match.group(2).strip()
            
            # Validate this is actually a role
            role_text = self._clean_candidate(role_text)
            if not role_text:
                continue
            
            is_valid, confidence = self._is_valid_role(role_text)
            if not is_valid:
                continue
            
            canonical, variant = self._get_canonical_role(role_text)
            
            # Create or update role entry
            if canonical not in existing_roles:
                # v3.0.12: Use helper to ensure entity_kind is populated
                existing_roles[canonical] = self._create_extracted_role(canonical, role_text)
            
            role_entry = existing_roles[canonical]
            role_entry.variants.add(variant)
            
            # Add responsibility if not duplicate
            if responsibility and responsibility not in role_entry.responsibilities:
                role_entry.responsibilities.append(responsibility)
            
            # Create occurrence record
            occurrence = RoleOccurrence(
                role=role_text,
                context=match.group(0).strip()[:150],
                responsibility=responsibility,
                action_type='performs',  # Default for responsibility assignments
                location=f"{source_location} (Responsibilities Section)",
                confidence=0.85  # High confidence from formal section
            )
            role_entry.occurrences.append(occurrence)
            role_entry.action_types['performs'] += 1
        
        return existing_roles
    
    def _scan_for_known_roles(self, text: str, existing_roles: Dict[str, 'ExtractedRole'], 
                              seen_matches: set, source_location: str) -> Dict[str, 'ExtractedRole']:
        """
        Directly scan for known roles that may have been missed by pattern matching.
        This ensures high-confidence roles like "Systems Engineer" are always found.
        """
        text_lower = text.lower()
        
        # Combine known roles with their aliases for comprehensive matching
        all_known = set(self.known_roles)
        for canonical, aliases in self.ROLE_ALIASES.items():
            all_known.add(canonical.lower())
            for alias in aliases:
                all_known.add(alias.lower())
        
        for known_role in all_known:
            if len(known_role) < 5:  # Skip short acronyms to avoid false matches
                continue
            
            # v3.0.91d: Skip if this role is in false_positives
            if known_role in self.false_positives:
                continue
            
            # Find all occurrences of this known role
            start = 0
            while True:
                pos = text_lower.find(known_role, start)
                if pos == -1:
                    break
                
                # Check word boundaries
                before_ok = pos == 0 or not text[pos-1].isalnum()
                after_pos = pos + len(known_role)
                after_ok = after_pos >= len(text) or not text[after_pos].isalnum()
                
                if before_ok and after_ok:
                    # Get the actual case from original text
                    actual_role = text[pos:after_pos]
                    canonical, variant = self._get_canonical_role(actual_role)
                    
                    # v3.0.91d: Skip if canonical form is in false_positives
                    if canonical.lower() in self.false_positives:
                        start = pos + 1
                        continue
                    
                    match_key = (canonical, pos)
                    if match_key not in seen_matches:
                        seen_matches.add(match_key)
                        
                        # Get sentence context
                        context = self._get_sentence_context(text, pos, after_pos)
                        responsibility, action_type = self._extract_responsibility(context, actual_role)
                        
                        occurrence = RoleOccurrence(
                            role=actual_role,
                            context=context,
                            responsibility=responsibility,
                            action_type=action_type,
                            location=source_location,
                            confidence=0.90  # High confidence for known roles
                        )
                        
                        if canonical not in existing_roles:
                            # v3.0.12: Use helper to ensure entity_kind is populated
                            # v3.0.12b: Fixed NameError - use actual_role not match.group(0)
                            existing_roles[canonical] = self._create_extracted_role(canonical, actual_role)
                        
                        role_entry = existing_roles[canonical]
                        role_entry.variants.add(variant)
                        role_entry.occurrences.append(occurrence)
                        
                        if responsibility:
                            role_entry.responsibilities.append(responsibility)
                        role_entry.action_types[action_type] += 1
                
                start = pos + 1
        
        return existing_roles
    
    def extract_from_docx(self, filepath: str) -> Dict[str, ExtractedRole]:
        """Extract roles from a Word document."""
        try:
            from docx import Document
        except ImportError:
            raise ImportError("python-docx required. Install: pip install python-docx")
        
        doc = Document(filepath)
        all_roles: Dict[str, ExtractedRole] = {}
        
        for para_num, paragraph in enumerate(doc.paragraphs, 1):
            text = paragraph.text.strip()
            if text:
                para_roles = self.extract_from_text(text, f"Paragraph {para_num}")
                self._merge_roles(all_roles, para_roles)
        
        for table_num, table in enumerate(doc.tables, 1):
            for row_num, row in enumerate(table.rows, 1):
                for cell_num, cell in enumerate(row.cells, 1):
                    text = cell.text.strip()
                    if text:
                        cell_roles = self.extract_from_text(
                            text, f"Table {table_num}, Row {row_num}, Cell {cell_num}"
                        )
                        self._merge_roles(all_roles, cell_roles)
        
        return all_roles
    
    def extract_from_pdf(self, filepath: str) -> Dict[str, ExtractedRole]:
        """Extract roles from a PDF document with enhanced table support."""
        try:
            import pdfplumber
            use_pdfplumber = True
        except ImportError:
            use_pdfplumber = False
            try:
                import PyPDF2
            except ImportError:
                raise ImportError("pdfplumber or PyPDF2 required. Install: pip install pdfplumber")
        
        all_roles: Dict[str, ExtractedRole] = {}
        
        # First, try enhanced table extraction for better accuracy
        table_roles = self._extract_roles_from_pdf_tables(filepath)
        self._merge_roles(all_roles, table_roles)
        _log(f"Found {len(table_roles)} roles from table extraction", level='debug')
        
        # Then extract from text
        if use_pdfplumber:
            with pdfplumber.open(filepath) as pdf:
                for page_num, page in enumerate(pdf.pages, 1):
                    text = page.extract_text() or ""
                    if text:
                        page_roles = self.extract_from_text(text, f"Page {page_num}")
                        self._merge_roles(all_roles, page_roles)
        else:
            with open(filepath, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                for page_num, page in enumerate(reader.pages, 1):
                    text = page.extract_text() or ""
                    if text:
                        page_roles = self.extract_from_text(text, f"Page {page_num}")
                        self._merge_roles(all_roles, page_roles)
        
        return all_roles
    
    def _extract_roles_from_pdf_tables(self, filepath: str) -> Dict[str, ExtractedRole]:
        """
        Extract roles specifically from PDF tables with enhanced confidence.
        
        Tables often contain RACI matrices, responsibility assignments, and
        role definitions with higher reliability than free text.
        """
        roles: Dict[str, ExtractedRole] = {}
        
        try:
            from enhanced_table_extractor import EnhancedTableExtractor
            extractor = EnhancedTableExtractor(prefer_accuracy=True)
            result = extractor.extract_tables(filepath)
            
            for table in result.tables:
                # Check if this looks like a RACI matrix or responsibility table
                is_raci = self._is_raci_table(table.headers, table.rows)
                is_responsibility_table = self._is_responsibility_table(table.headers)
                
                # Confidence boost for table-sourced roles
                confidence_boost = 0.20 if is_raci else (0.15 if is_responsibility_table else 0.10)
                
                # Extract from headers (often contain role names)
                for header in table.headers:
                    header_roles = self.extract_from_text(
                        header, f"Table {table.index} Header"
                    )
                    for canonical, role_data in header_roles.items():
                        # Boost confidence for table headers
                        for occ in role_data.occurrences:
                            occ.confidence = min(1.0, occ.confidence + confidence_boost)
                        self._merge_role(roles, canonical, role_data)
                
                # Extract from table cells
                for row_idx, row in enumerate(table.rows):
                    for col_idx, cell in enumerate(row):
                        if cell and len(cell.strip()) > 2:
                            cell_roles = self.extract_from_text(
                                cell, f"Table {table.index}, Row {row_idx+1}"
                            )
                            for canonical, role_data in cell_roles.items():
                                for occ in role_data.occurrences:
                                    occ.confidence = min(1.0, occ.confidence + confidence_boost * 0.75)
                                self._merge_role(roles, canonical, role_data)
                                
        except ImportError:
            _log("Enhanced table extractor not available, using basic extraction", level='debug')
        except Exception as e:
            _log(f"Table role extraction failed: {e}", level='debug')
        
        return roles
    
    def _is_raci_table(self, headers: List[str], rows: List[List[str]]) -> bool:
        """Detect if a table is a RACI matrix."""
        # Check headers for RACI indicators
        raci_headers = {'r', 'a', 'c', 'i', 'responsible', 'accountable', 'consulted', 'informed'}
        header_text = ' '.join(h.lower() for h in headers)
        
        if any(raci in header_text for raci in ['raci', 'rasci', 'responsibility matrix']):
            return True
        
        # Check if cells contain mostly R/A/C/I values
        raci_values = {'r', 'a', 'c', 'i', 'x', '●', '✓', ''}
        raci_count = 0
        total_cells = 0
        
        for row in rows[:5]:  # Check first few rows
            for cell in row[1:]:  # Skip first column (usually role names)
                total_cells += 1
                if cell.strip().lower() in raci_values or len(cell.strip()) <= 2:
                    raci_count += 1
        
        return total_cells > 0 and raci_count / total_cells > 0.6
    
    def _is_responsibility_table(self, headers: List[str]) -> bool:
        """Detect if a table contains responsibility assignments."""
        responsibility_keywords = {
            'role', 'responsibility', 'function', 'task', 'activity',
            'action', 'owner', 'lead', 'support', 'department', 'organization'
        }
        header_text = ' '.join(h.lower() for h in headers)
        return any(kw in header_text for kw in responsibility_keywords)
    
    def _merge_role(self, target: Dict[str, ExtractedRole], 
                    canonical: str, role_data: ExtractedRole):
        """Merge a single role into target dictionary."""
        if canonical not in target:
            target[canonical] = role_data
        else:
            existing = target[canonical]
            existing.variants.update(role_data.variants)
            existing.occurrences.extend(role_data.occurrences)
            existing.responsibilities.extend(role_data.responsibilities)
    
    def _merge_roles(self, target: Dict[str, ExtractedRole], source: Dict[str, ExtractedRole]):
        """Merge source roles into target."""
        for canonical, role_data in source.items():
            # v3.2.5: Skip excessively long roles (>60 chars) as final safeguard
            if len(canonical) > 60:
                continue
            if canonical not in target:
                target[canonical] = role_data
            else:
                existing = target[canonical]
                existing.variants.update(role_data.variants)
                existing.occurrences.extend(role_data.occurrences)
                existing.responsibilities.extend(role_data.responsibilities)
                for action, count in role_data.action_types.items():
                    existing.action_types[action] += count

    def extract_deliverables(self, text: str, source_location: str = "unknown") -> List[Dict]:
        """
        Extract deliverables from text using NLP (v3.1.2 ENH-008).

        Args:
            text: Document text to analyze
            source_location: Location identifier for tracking

        Returns:
            List of deliverable dictionaries with name, type, confidence, context
        """
        deliverables = []

        # Use NLP processor if available
        if self._nlp_processor:
            try:
                nlp_deliverables = self._nlp_processor.extract_deliverables(text)
                for d in nlp_deliverables:
                    deliverables.append({
                        'name': d.name,
                        'normalized_name': d.normalized_name,
                        'type': d.deliverable_type,
                        'confidence': d.confidence,
                        'source': d.source,
                        'context': d.context[:200] if d.context else "",
                        'location': source_location
                    })
                _log(f"NLP found {len(deliverables)} deliverables", level='debug')
            except Exception as e:
                _log(f"NLP deliverable extraction failed: {e}", level='warning')

        # Fallback/supplement with pattern-based extraction
        deliverables.extend(self._extract_deliverables_patterns(text, source_location))

        # Deduplicate
        seen = set()
        unique = []
        for d in deliverables:
            key = d['normalized_name'].lower()
            if key not in seen:
                seen.add(key)
                unique.append(d)

        return sorted(unique, key=lambda x: -x['confidence'])

    def _extract_deliverables_patterns(self, text: str, source_location: str) -> List[Dict]:
        """Extract deliverables using regex patterns."""
        deliverables = []

        # Pattern: shall deliver/provide/submit [deliverable]
        pattern = r'\b(?:shall|will|must)\s+(?:deliver|provide|submit|prepare|develop|create)\s+(?:the\s+)?([A-Za-z][A-Za-z\s]+?(?:document|report|plan|specification|design|procedure|list|data|package|matrix)s?)\b'

        for match in re.finditer(pattern, text, re.IGNORECASE):
            name = match.group(1).strip()
            deliverables.append({
                'name': name,
                'normalized_name': name.title(),
                'type': self._classify_deliverable_type(name),
                'confidence': 0.7,
                'source': 'pattern',
                'context': text[max(0, match.start()-50):match.end()+50],
                'location': source_location
            })

        # Pattern: CDRL/DID references
        cdrl_pattern = r'\b((?:CDRL|DID|DI-)[\s-]?[A-Z0-9-]+)\b'
        for match in re.finditer(cdrl_pattern, text):
            deliverables.append({
                'name': match.group(1),
                'normalized_name': match.group(1).upper(),
                'type': 'artifact',
                'confidence': 0.95,
                'source': 'pattern',
                'context': text[max(0, match.start()-30):match.end()+30],
                'location': source_location
            })

        return deliverables

    def _classify_deliverable_type(self, name: str) -> str:
        """Classify a deliverable by type."""
        name_lower = name.lower()
        type_map = {
            'document': ['document', 'doc', 'manual', 'guide', 'handbook'],
            'report': ['report', 'summary', 'assessment', 'analysis', 'review'],
            'plan': ['plan', 'schedule', 'roadmap', 'strategy'],
            'specification': ['specification', 'spec', 'requirement', 'standard'],
            'design': ['design', 'architecture', 'drawing', 'diagram'],
            'procedure': ['procedure', 'process', 'instruction', 'protocol'],
            'data': ['data', 'dataset', 'package', 'file'],
            'list': ['list', 'inventory', 'catalog', 'matrix']
        }
        for dtype, keywords in type_map.items():
            for kw in keywords:
                if kw in name_lower:
                    return dtype
        return 'artifact'

    def extract_acronyms(self, text: str) -> List[Dict]:
        """
        Extract acronyms from text using NLP (v3.1.2 ENH-008).

        Args:
            text: Document text to analyze

        Returns:
            List of acronym dictionaries with acronym, expansion, is_defined, usage_count
        """
        acronyms = []

        # Use NLP processor if available
        if self._nlp_processor:
            try:
                nlp_acronyms = self._nlp_processor.extract_acronyms(text)
                for a in nlp_acronyms:
                    acronyms.append({
                        'acronym': a.acronym,
                        'expansion': a.expansion,
                        'is_defined': a.is_defined,
                        'confidence': a.confidence,
                        'usage_count': len(a.usage_locations),
                        'definition_location': a.definition_location
                    })
                _log(f"NLP found {len(acronyms)} acronyms", level='debug')
            except Exception as e:
                _log(f"NLP acronym extraction failed: {e}", level='warning')

        return acronyms

    def generate_report(self, roles: Dict[str, ExtractedRole], 
                       min_confidence: float = 0.5,
                       min_occurrences: int = 1,
                       sort_by: str = 'frequency') -> str:
        """Generate a formatted report of extracted roles."""
        lines = []
        lines.append("=" * 80)
        lines.append("ROLE EXTRACTION REPORT")
        lines.append("=" * 80)
        lines.append("")
        
        filtered_roles = [
            (name, data) for name, data in roles.items()
            if data.avg_confidence >= min_confidence and data.frequency >= min_occurrences
        ]
        
        if sort_by == 'frequency':
            filtered_roles.sort(key=lambda x: (-x[1].frequency, -x[1].avg_confidence, x[0]))
        elif sort_by == 'confidence':
            filtered_roles.sort(key=lambda x: (-x[1].avg_confidence, -x[1].frequency, x[0]))
        else:
            filtered_roles.sort(key=lambda x: x[0])
        
        lines.append(f"Total unique roles found: {len(filtered_roles)}")
        lines.append(f"Filters: min_confidence={min_confidence}, min_occurrences={min_occurrences}")
        lines.append("")
        lines.append("-" * 80)
        
        for name, data in filtered_roles:
            lines.append(f"\n■ {name}")
            lines.append(f"  Frequency: {data.frequency} | Confidence: {data.avg_confidence:.0%}")
            
            if data.variants and len(data.variants) > 1:
                variants = sorted(v for v in data.variants if v.lower() != name.lower())
                if variants:
                    lines.append(f"  Also appears as: {', '.join(variants[:5])}")
            
            actions = {k: v for k, v in data.action_types.items() if k != 'unknown'}
            if actions:
                action_str = ', '.join(f"{k}({v})" for k, v in sorted(actions.items(), key=lambda x: -x[1])[:5])
                lines.append(f"  Actions: {action_str}")
            
            if data.responsibilities:
                unique_resp = list(set(data.responsibilities))[:3]
                lines.append(f"  Responsibilities:")
                for resp in unique_resp:
                    lines.append(f"    • {resp[:80]}{'...' if len(resp) > 80 else ''}")
            
            if data.occurrences:
                occ = data.occurrences[0]
                ctx = occ.context[:120] + ('...' if len(occ.context) > 120 else '')
                lines.append(f"  Example: \"{ctx}\"")
                lines.append(f"           [{occ.location}]")
        
        lines.append("")
        lines.append("=" * 80)
        lines.append("END OF REPORT")
        lines.append("=" * 80)
        
        return '\n'.join(lines)
    
    def export_to_csv(self, roles: Dict[str, ExtractedRole], filepath: str,
                     min_confidence: float = 0.5, min_occurrences: int = 1):
        """Export extracted roles to CSV format."""
        filtered_roles = [
            (name, data) for name, data in roles.items()
            if data.avg_confidence >= min_confidence and data.frequency >= min_occurrences
        ]
        filtered_roles.sort(key=lambda x: (-x[1].frequency, -x[1].avg_confidence))
        
        with open(filepath, 'w', newline='', encoding='utf-8') as f:
            writer = csv.writer(f)
            writer.writerow([
                'Role', 'Frequency', 'Confidence', 'Variants', 
                'Primary_Actions', 'Sample_Responsibilities', 'Source_Locations'
            ])
            
            for name, data in filtered_roles:
                actions = {k: v for k, v in data.action_types.items() if k != 'unknown'}
                writer.writerow([
                    name,
                    data.frequency,
                    f"{data.avg_confidence:.0%}",
                    '; '.join(sorted(data.variants)),
                    '; '.join(f"{k}({v})" for k, v in sorted(actions.items(), key=lambda x: -x[1])[:3]),
                    '; '.join(list(set(data.responsibilities))[:3]),
                    '; '.join(sorted(set(occ.location for occ in data.occurrences[:5])))
                ])
        
        _log(f"Exported {len(filtered_roles)} roles to {filepath}")


# =============================================================================
# COMPREHENSIVE TESTING
# =============================================================================

def run_tests():
    """Run comprehensive tests against NASA document text."""
    
    # Extended test corpus from NASA Systems Engineering Handbook
    test_corpus = """
    The systems engineer is skilled in the art and science of balancing organizational, cost, and 
    technical interactions in complex systems. The systems engineer and supporting organization are 
    vital to supporting program and Project Planning and Control (PP&C) with accurate and timely cost 
    and schedule information for the technical activities.
    
    The lead systems engineer ensures that the system technically fulfills the defined needs and 
    requirements and that a proper systems engineering approach is being followed. The systems engineer 
    oversees the project's systems engineering activities as performed by the technical team and directs, 
    communicates, monitors, and coordinates tasks. The systems engineer reviews and evaluates the 
    technical aspects of the project.
    
    The project manager has overall responsibility for managing the project team and ensuring that 
    the project delivers a technically correct system within cost and schedule. The project manager 
    may sometimes perform these practices for small projects.
    
    Systems engineering plays a key role in the project organization. Managing a project consists of 
    three main objectives: managing the technical aspects of the project, managing the project team, 
    and managing the cost and schedule.
    
    The exact role and responsibility of the systems engineer may change from project to project 
    depending on the size and complexity of the project.
    
    For large projects, there may be one or more systems engineers. The systems engineer usually 
    plays the key role in leading the development of the concept of operations (ConOps).
    
    The Configuration Manager maintains the integrity and traceability of product configuration 
    throughout the system life cycle.
    
    Technical reviews are conducted by the technical team and reviewed by the Chief Engineer.
    
    The Mission Directorate Associate Administrator (MDAA) approves major program decisions.
    
    Integration activities are coordinated with the Integration Lead and verified by Quality Assurance.
    
    The Principal Investigator is responsible for the scientific objectives of the mission.
    
    Safety reviews are conducted by the Safety Engineer and approved by the System Safety Panel.
    
    The Contracting Officer Representative (COR) monitors contractor performance and ensures 
    compliance with contract requirements.
    
    The Risk Manager is responsible for identifying, analyzing, and mitigating project risks.
    
    Interface specifications are developed by the Interface Control Working Group (ICWG) and approved 
    by the Configuration Control Board (CCB).
    
    The Software Lead coordinates with the Hardware Lead to ensure system integration.
    
    Verification activities are performed by the Test Engineer and validated by Quality Assurance.
    
    The Technical Authority provides independent technical oversight of the project.
    
    Requirements are traced by the Requirements Manager and verified by the Verification Lead.
    
    The Flight Director is responsible for all real-time mission operations.
    
    Program-level decisions require approval from the Program Manager and the Chief Engineer.
    
    The Integration Engineer assembles components according to the integration plan and works with 
    the Test Engineer to verify proper assembly.
    
    Mission assurance is provided by the Mission Assurance Manager who ensures compliance with 
    quality standards.
    
    The Deputy Project Manager assists the Project Manager in managing day-to-day operations.
    
    The Reliability Engineer analyzes system reliability and provides recommendations for 
    improvement.
    
    Configuration audits are conducted by the Configuration Manager with support from Quality 
    Assurance.
    
    The technical team supports the systems engineer in performing technical assessments.
    """
    
    # Expected roles that should be found
    expected_roles = [
        'Systems Engineer', 'Lead Systems Engineer', 'Project Manager', 'Chief Engineer',
        'Configuration Manager', 'Principal Investigator', 'Safety Engineer',
        'Risk Manager', 'Software Lead', 'Hardware Lead', 'Test Engineer',
        'Technical Authority', 'Requirements Manager', 'Verification Lead',
        'Flight Director', 'Program Manager', 'Integration Lead', 'Integration Engineer',
        'Mission Directorate Associate Administrator', 'Contracting Officer Representative',
        'Interface Control Working Group', 'Configuration Control Board', 'System Safety Panel',
        'Quality Assurance', 'Technical Team', 'Mission Assurance Manager',
        'Deputy Project Manager', 'Reliability Engineer'
    ]
    
    print("=" * 80)
    print("ROLE EXTRACTOR v3.0 - COMPREHENSIVE TEST")
    print("=" * 80)
    print()
    
    extractor = RoleExtractor()
    roles = extractor.extract_from_text(test_corpus, "Test Corpus")
    
    # Generate report
    report = extractor.generate_report(roles, min_confidence=0.5, min_occurrences=1)
    print(report)
    
    # Validation
    print("\n" + "=" * 80)
    print("VALIDATION RESULTS")
    print("=" * 80)
    
    found_roles = set(roles.keys())
    expected_set = set(expected_roles)
    
    # Check coverage
    matched = found_roles & expected_set
    missed = expected_set - found_roles
    extra = found_roles - expected_set
    
    print(f"\nExpected roles: {len(expected_set)}")
    print(f"Found roles: {len(found_roles)}")
    print(f"Matched: {len(matched)} ({len(matched)/len(expected_set)*100:.1f}%)")
    
    if missed:
        print(f"\nMissed ({len(missed)}):")
        for role in sorted(missed):
            print(f"  - {role}")
    
    if extra:
        print(f"\nAdditional roles found ({len(extra)}):")
        for role in sorted(extra):
            data = roles[role]
            print(f"  + {role} (freq={data.frequency}, conf={data.avg_confidence:.0%})")
    
    # Calculate accuracy metrics
    precision = len(matched) / len(found_roles) if found_roles else 0
    recall = len(matched) / len(expected_set) if expected_set else 0
    f1 = 2 * precision * recall / (precision + recall) if (precision + recall) > 0 else 0
    
    print(f"\n--- Metrics ---")
    print(f"Precision: {precision:.1%}")
    print(f"Recall: {recall:.1%}")
    print(f"F1 Score: {f1:.1%}")
    
    return roles


if __name__ == "__main__":
    run_tests()
