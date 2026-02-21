"""
AEGIS Proposal Parser — Extracts financial data from DOCX, PDF, and XLSX files.

Pure extraction — NO AI/LLM involvement. Uses:
- openpyxl for Excel spreadsheets
- python-docx + lxml for DOCX tables
- EnhancedTableExtractor (camelot → tabula → pdfplumber) for PDF tables (primary)
- pymupdf4llm / fitz for PDF tables (fallback)
- Data-pattern column inference for headerless tables
- Regex patterns for dollar amounts, percentages, dates

Returns structured ProposalData objects with company info, tables, and financial items.
"""

import re
import os
import logging
from dataclasses import dataclass, field
from typing import List, Dict, Optional, Any, Tuple

logger = logging.getLogger(__name__)

# ──────────────────────────────────────────────
# Data classes
# ──────────────────────────────────────────────

@dataclass
class LineItem:
    """A single financial line item extracted from a proposal."""
    description: str = ''
    amount: Optional[float] = None
    amount_raw: str = ''        # Original string, e.g. "$1,234.56"
    quantity: Optional[float] = None
    unit_price: Optional[float] = None
    unit: str = ''
    category: str = ''          # Labor, Material, Travel, ODC, etc.
    row_index: int = 0          # Source row in the table
    table_index: int = 0        # Which table it came from
    source_sheet: str = ''      # Sheet name (Excel) or page (PDF)
    confidence: float = 1.0     # Extraction confidence (1.0 = certain)

@dataclass
class ExtractedTable:
    """A table extracted from a document."""
    headers: List[str] = field(default_factory=list)
    rows: List[List[str]] = field(default_factory=list)
    source: str = ''           # Sheet name, page number, etc.
    table_index: int = 0
    has_financial_data: bool = False
    total_row_index: Optional[int] = None  # Index of "Total" row if found

@dataclass
class ProposalData:
    """Complete extracted data from a single proposal document."""
    filename: str = ''
    filepath: str = ''
    file_type: str = ''        # 'xlsx', 'docx', 'pdf'
    company_name: str = ''
    proposal_title: str = ''
    date: str = ''

    # Extracted data
    tables: List[ExtractedTable] = field(default_factory=list)
    line_items: List[LineItem] = field(default_factory=list)

    # Financial summary
    total_amount: Optional[float] = None
    total_raw: str = ''
    currency: str = 'USD'

    # Contract term
    contract_term: str = ''    # e.g., "3 Year", "Base + 4 Options", "36 Months"

    # Metadata
    page_count: int = 0
    extraction_notes: List[str] = field(default_factory=list)
    extraction_text: str = ''   # Full extracted text for doc viewer fallback

    def to_dict(self) -> Dict[str, Any]:
        """Convert to JSON-serializable dict."""
        return {
            'filename': self.filename,
            'filepath': self.filepath,
            'file_type': self.file_type,
            'company_name': self.company_name,
            'proposal_title': self.proposal_title,
            'date': self.date,
            'tables': [
                {
                    'headers': t.headers,
                    'rows': t.rows,
                    'source': t.source,
                    'table_index': t.table_index,
                    'has_financial_data': t.has_financial_data,
                    'total_row_index': t.total_row_index,
                }
                for t in self.tables
            ],
            'line_items': [
                {
                    'description': li.description,
                    'amount': li.amount,
                    'amount_raw': li.amount_raw,
                    'quantity': li.quantity,
                    'unit_price': li.unit_price,
                    'unit': li.unit,
                    'category': li.category,
                    'row_index': li.row_index,
                    'table_index': li.table_index,
                    'source_sheet': li.source_sheet,
                    'confidence': li.confidence,
                }
                for li in self.line_items
            ],
            'total_amount': self.total_amount,
            'total_raw': self.total_raw,
            'currency': self.currency,
            'contract_term': self.contract_term,
            'page_count': self.page_count,
            'extraction_notes': self.extraction_notes,
            'extraction_text': self.extraction_text,
        }


# ──────────────────────────────────────────────
# Regex patterns for financial extraction
# ──────────────────────────────────────────────

# Dollar amounts: $1,234.56, $1234, $ 1,234.56, $1,234,567.89
DOLLAR_PATTERN = re.compile(
    r'\$\s*[\d,]+(?:\.\d{1,2})?'
)

# Percentage: 10%, 10.5%, 0.5%
PERCENT_PATTERN = re.compile(
    r'\d+(?:\.\d+)?\s*%'
)

# Date patterns: MM/DD/YYYY, YYYY-MM-DD, Month DD YYYY, etc.
DATE_PATTERN = re.compile(
    r'(?:\d{1,2}/\d{1,2}/\d{2,4})|'
    r'(?:\d{4}-\d{2}-\d{2})|'
    r'(?:(?:January|February|March|April|May|June|July|August|September|October|November|December)\s+\d{1,2},?\s+\d{4})|'
    r'(?:(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)\.?\s+\d{1,2},?\s+\d{4})',
    re.IGNORECASE
)

# Numeric with optional commas: 1,234.56 or 1234.56
NUMERIC_PATTERN = re.compile(
    r'^[\d,]+(?:\.\d{1,4})?$'
)

# Total / subtotal / grand total row indicators
TOTAL_KEYWORDS = re.compile(
    r'\b(?:total|subtotal|sub-total|sub\s*total|grand\s*total|sum|net|gross|overall|'
    r'year\s*total|annual\s*total|contract\s*total|program\s*total|'
    r'total\s*price|total\s*cost|total\s*amount|total\s*value)\b',
    re.IGNORECASE
)

# Category keywords for classification
CATEGORY_PATTERNS = {
    'Labor': re.compile(
        r'\b(?:labor|labour|personnel|staffing|hours?|FTE|man-?hours?|salaries?|wages?|'
        r'engineering|engineer|development|developer|software\s+dev|systems?\s+eng|'
        r'program\s*management|project\s*management|management\s+labor|'
        r'technical\s+(?:support|staff|writer|lead)|analyst|architect|designer|'
        r'SME|subject\s+matter|test(?:ing|er)?|QA|quality\s+assurance)\b', re.IGNORECASE),
    'Material': re.compile(
        r'\b(?:material|supplies|equipment|hardware|procurement|'
        r'tool(?:ing|s)?|parts?|component|inventory|COTS|GOTS)\b', re.IGNORECASE),
    'Software': re.compile(
        r'\b(?:software\s+license|COTS\s+software|software\s+(?:tool|package|suite|product)|'
        r'SaaS|cloud\s+service|platform\s+license)\b', re.IGNORECASE),
    'License': re.compile(
        r'\b(?:licens(?:e|ing)|subscription|annual\s+(?:fee|renewal)|seat\s+(?:fee|license)|'
        r'user\s+licens|per\s+(?:user|seat)|maintenance\s+(?:agreement|fee|renewal)|'
        r'support\s+(?:agreement|contract|renewal)|SLA|service\s+level)\b', re.IGNORECASE),
    'Travel': re.compile(
        r'\b(?:travel|per\s*diem|airfare|lodging|mileage|transportation|trip)\b', re.IGNORECASE),
    'Training': re.compile(
        r'\b(?:training|course|certification|workshop|seminar|instruction)\b', re.IGNORECASE),
    'ODC': re.compile(
        r'\b(?:ODC|other\s*direct|subcontract|consultant|vendor|third.?party)\b', re.IGNORECASE),
    'Overhead': re.compile(
        r'\b(?:overhead|indirect|G&A|general\s*(?:and|&)\s*admin|fringe|burden|wrap\s*rate)\b', re.IGNORECASE),
    'Fee': re.compile(
        r'\b(?:fee|profit|margin|markup|award)\b', re.IGNORECASE),
}

# Company name patterns (found in headers/footers/cover)
COMPANY_INDICATORS = re.compile(
    r'\b(?:LLC|Inc\.?|Corp\.?|Corporation|Company|Co\.|Ltd\.?|Group|Associates|Partners|Consulting|Technologies|Solutions|Systems|Services|International|Aerospace|Defense|Defence)\b',
    re.IGNORECASE
)


# ──────────────────────────────────────────────
# Utility functions
# ──────────────────────────────────────────────

def parse_dollar_amount(text: str) -> Optional[float]:
    """Parse a dollar string into a float. Returns None if not parseable."""
    if not text:
        return None
    # Remove $ and spaces, keep digits, commas, periods, minus
    cleaned = re.sub(r'[^\d,.\-()]', '', text)
    if not cleaned:
        return None
    # Handle parenthetical negatives: (1,234.56) -> -1234.56
    is_negative = '(' in text and ')' in text
    cleaned = cleaned.replace('(', '').replace(')', '').replace(',', '')
    try:
        val = float(cleaned)
        return -val if is_negative else val
    except ValueError:
        return None


def classify_line_item(description: str) -> str:
    """Classify a line item into a category based on description text."""
    for category, pattern in CATEGORY_PATTERNS.items():
        if pattern.search(description):
            return category
    return 'Other'


def extract_company_from_text(text: str, max_chars: int = 2000, filename: str = '') -> str:
    """Try to extract company name from document text (first N chars).

    Uses multiple strategies in priority order:
    1. Lines containing company legal suffixes (LLC, Inc, Corp, etc.)
    2. "Prepared by" / "Submitted by" attribution lines
    3. First prominent heading-like line (short, no numbers, title case)
    4. Company name extracted from the filename
    """
    search_text = text[:max_chars]

    # Strategy 1: Look for lines with company indicators
    for line in search_text.split('\n'):
        line = line.strip()
        if not line or len(line) > 120:
            continue
        # Skip lines that look like table data (contain lots of numbers/dollar signs)
        if DOLLAR_PATTERN.search(line):
            continue
        # Skip lines with too many numbers (likely data rows, not company names)
        digit_ratio = sum(c.isdigit() for c in line) / max(len(line), 1)
        if digit_ratio > 0.3:
            continue
        if COMPANY_INDICATORS.search(line):
            # Clean up the line
            name = re.sub(r'^\W+|\W+$', '', line)
            # Remove common prefixes
            name = re.sub(r'^(?:prepared\s+by|submitted\s+by|from|proposal\s+by)\s*:?\s*', '', name, flags=re.IGNORECASE)
            if 5 < len(name) < 80:
                return name.strip()

    # Strategy 2: Look for "Prepared by: XXX" or "Submitted by: XXX"
    match = re.search(
        r'(?:prepared|submitted|proposed|offered|provided)\s+by\s*:?\s*(.+?)(?:\n|$)',
        search_text, re.IGNORECASE
    )
    if match:
        name = match.group(1).strip()
        if 3 < len(name) < 80:
            return name

    # Strategy 3: Look for "Vendor:" or "Supplier:" or "Contractor:" labels
    match = re.search(
        r'(?:vendor|supplier|contractor|company|firm|offeror)\s*:?\s*(.+?)(?:\n|$)',
        search_text, re.IGNORECASE
    )
    if match:
        name = match.group(1).strip()
        # Skip if it looks like a generic label or instruction
        if 3 < len(name) < 80 and not re.search(r'^\s*(?:name|enter|fill|n/a)', name, re.IGNORECASE):
            return name

    # Strategy 4: First prominent heading-like line (title case, short, not generic)
    skip_words = {'proposal', 'pricing', 'quote', 'quotation', 'bid', 'offer', 'response',
                  'statement', 'work', 'scope', 'table', 'contents', 'page', 'date',
                  'section', 'attachment', 'appendix', 'exhibit', 'schedule', 'cost',
                  'price', 'summary', 'executive', 'technical', 'management', 'volume'}
    for line in search_text.split('\n')[:30]:
        line = line.strip()
        if not line or len(line) < 3 or len(line) > 60:
            continue
        if DOLLAR_PATTERN.search(line):
            continue
        digit_ratio = sum(c.isdigit() for c in line) / max(len(line), 1)
        if digit_ratio > 0.15:
            continue
        words = line.split()
        if len(words) > 6 or len(words) < 1:
            continue
        # Check if all words start capitalized and none are skip words
        lower_words = {w.lower().rstrip('.,;:') for w in words}
        if lower_words & skip_words:
            continue
        if all(w[0].isupper() for w in words if w):
            return line

    # Strategy 5: Extract from filename (e.g., "Dell_Proposal_3yr.pdf" → "Dell")
    if filename:
        # Remove extension and common suffixes
        base = re.sub(r'\.[^.]+$', '', filename)
        # Replace separators with spaces
        base = re.sub(r'[_\-]+', ' ', base)
        # Remove common terms
        base = re.sub(
            r'\b(?:proposal|quote|pricing|bid|offer|3\s*yr|5\s*yr|3\s*year|5\s*year|draft|final|v\d+|rev\s*\d*)\b',
            '', base, flags=re.IGNORECASE
        ).strip()
        # Clean up
        base = re.sub(r'\s+', ' ', base).strip()
        if 2 < len(base) < 60:
            return base

    return ''


def extract_dates_from_text(text: str, max_chars: int = 3000) -> str:
    """Extract the most prominent date from document text."""
    search_text = text[:max_chars]
    matches = DATE_PATTERN.findall(search_text)
    if matches:
        return matches[0]
    return ''


def extract_contract_term(text: str, sheet_names: List[str] = None) -> str:
    """Extract contract term/period from proposal text and/or sheet names.

    Detects patterns like:
    - "3 year contract", "5-year period"
    - "Base Year plus 4 Option Years"
    - "Period of Performance: 36 months"
    - Sheet tabs named "Year 1", "BY", "OY1", etc.
    """
    search_text = text[:5000].lower() if text else ''

    # Pattern 1: X-year contract/period/term
    m = re.search(r'(\d+)\s*[-‐]?\s*year\s*(contract|period|term|agreement|effort)', search_text)
    if m:
        return f'{m.group(1)} Year'

    # Pattern 2: Base Year + Option Years
    m = re.search(r'base\s*year\s*(?:plus|with|\+|and)\s*(\d+)\s*option', search_text)
    if m:
        return f'Base + {m.group(1)} Options'

    # Pattern 3: Period of performance in months
    m = re.search(r'period\s*of\s*performance\s*:?\s*(\d+)\s*months?', search_text)
    if m:
        months = int(m.group(1))
        if months % 12 == 0:
            return f'{months // 12} Year'
        return f'{months} Months'

    # Pattern 4: BY/OY abbreviations in text
    oy_matches = re.findall(r'\bOY\s*(\d+)\b', search_text, re.IGNORECASE)
    if oy_matches:
        max_oy = max(int(n) for n in oy_matches)
        return f'Base + {max_oy} Options'

    # Pattern 5: Detect from Excel sheet tab names
    if sheet_names:
        year_tabs = []
        option_tabs = []
        for name in sheet_names:
            nl = name.lower().strip()
            if re.match(r'year\s*\d+', nl) or re.match(r'yr\s*\d+', nl):
                year_tabs.append(nl)
            elif re.match(r'oy\s*\d+', nl) or 'option' in nl:
                option_tabs.append(nl)
            elif nl in ('by', 'base year', 'base'):
                option_tabs.append(nl)  # Base year = part of option structure

        if year_tabs:
            return f'{len(year_tabs)} Year'
        if option_tabs:
            return f'Base + {len(option_tabs) - 1} Options' if len(option_tabs) > 1 else 'Base Year'

    return ''


def is_financial_table(headers: List[str], rows: List[List[str]]) -> bool:
    """Determine if a table likely contains financial data."""
    header_text = ' '.join(h.lower() for h in headers if h)

    # Check headers for financial indicators
    financial_headers = ['cost', 'price', 'amount', 'total', 'rate', 'fee',
                         'budget', 'estimate', 'value', 'subtotal', 'charge',
                         'extended', 'unit price', 'labor', 'material', 'travel']
    if any(kw in header_text for kw in financial_headers):
        return True

    # Check if rows contain dollar amounts OR large numeric values
    financial_cell_count = 0
    for row in rows[:10]:  # Check first 10 rows
        for cell in row:
            cell_str = str(cell).strip()
            # Dollar-formatted strings
            if DOLLAR_PATTERN.search(cell_str):
                financial_cell_count += 1
            # Plain numeric values > 100 (likely financial, not just IDs/quantities)
            elif _cell_to_float(cell) is not None:
                val = _cell_to_float(cell)
                if val is not None and val > 100:
                    financial_cell_count += 1

    # If more than 15% of checked cells have financial values, it's financial
    total_cells = sum(len(row) for row in rows[:10])
    if total_cells > 0 and financial_cell_count / total_cells > 0.15:
        return True

    return False


def _cell_to_float(cell_val) -> Optional[float]:
    """Convert a cell value to float — handles both dollar-formatted strings and raw numbers.

    openpyxl data_only=True returns raw floats/ints for number cells,
    while DOCX/PDF tables return formatted strings like '$250,000.00'.
    """
    if cell_val is None:
        return None
    if isinstance(cell_val, (int, float)):
        return float(cell_val) if cell_val != 0 else None
    s = str(cell_val).strip()
    if not s:
        return None
    # Try dollar pattern first
    dollar_match = DOLLAR_PATTERN.search(s)
    if dollar_match:
        return parse_dollar_amount(dollar_match.group())
    # Try plain numeric (with optional commas)
    cleaned = s.replace(',', '').replace(' ', '')
    if NUMERIC_PATTERN.match(cleaned):
        try:
            val = float(cleaned)
            return val if val != 0 else None
        except ValueError:
            return None
    return None


def find_total_row(rows: List[List[str]]) -> Tuple[Optional[int], Optional[float]]:
    """Find the "Total" row in a table and extract its amount."""
    for i, row in enumerate(rows):
        for j, cell in enumerate(row):
            if TOTAL_KEYWORDS.search(str(cell)):
                # Look for the largest numeric value in the same row
                # (the total is usually the rightmost/largest number)
                best_amount = None
                for k, val in enumerate(row):
                    if k == j:
                        continue  # Skip the cell with "Total" text
                    amount = _cell_to_float(val)
                    if amount is not None and amount > 0:
                        if best_amount is None or amount > best_amount:
                            best_amount = amount
                if best_amount is not None:
                    return i, best_amount
    return None, None


def _infer_columns_from_data(headers: List[str], rows: List[List[str]]) -> Dict[str, Optional[int]]:
    """Infer column roles from data patterns when headers are generic or missing.

    Analyzes the first N data rows to determine which columns contain:
    - Dollar amounts (description vs unit price vs extended/total)
    - Quantities (small integers)
    - Descriptions (long text strings)
    - Item numbers / IDs (short numeric or alphanumeric)

    Returns dict with keys: desc_col, amount_col, qty_col, unit_price_col
    """
    if not rows:
        return {'desc_col': None, 'amount_col': None, 'qty_col': None, 'unit_price_col': None}

    num_cols = max(len(r) for r in rows) if rows else 0
    if num_cols == 0:
        return {'desc_col': None, 'amount_col': None, 'qty_col': None, 'unit_price_col': None}

    sample_rows = rows[:min(10, len(rows))]

    # Per-column stats
    col_stats = []
    for col_idx in range(num_cols):
        stats = {
            'dollar_count': 0,      # Cells with $ prefix
            'numeric_count': 0,     # Pure numeric cells (no $)
            'text_count': 0,        # Cells with alphabetic content
            'small_int_count': 0,   # Integers 1-999 (likely qty)
            'large_num_count': 0,   # Numbers > 100 (likely amounts)
            'avg_text_len': 0,      # Average text length
            'total_cells': 0,
            'sample_values': [],
            'numeric_sum': 0,       # Sum of numeric values (larger = more likely total col)
        }

        for row in sample_rows:
            if col_idx >= len(row):
                continue
            cell = str(row[col_idx]).strip()
            if not cell:
                continue

            stats['total_cells'] += 1
            stats['sample_values'].append(cell)

            # Check for dollar amount
            if DOLLAR_PATTERN.search(cell):
                stats['dollar_count'] += 1
                val = parse_dollar_amount(DOLLAR_PATTERN.search(cell).group())
                if val is not None:
                    stats['numeric_sum'] += val
                    if val > 100:
                        stats['large_num_count'] += 1

            # Check for plain numeric
            cleaned = cell.replace(',', '').replace(' ', '')
            try:
                num_val = float(cleaned)
                stats['numeric_count'] += 1
                stats['numeric_sum'] += num_val
                if 1 <= num_val <= 999 and num_val == int(num_val):
                    stats['small_int_count'] += 1
                if num_val > 100:
                    stats['large_num_count'] += 1
            except ValueError:
                pass

            # Check for text content
            alpha_chars = sum(1 for c in cell if c.isalpha())
            if alpha_chars > 3:
                stats['text_count'] += 1

            stats['avg_text_len'] += len(cell)

        if stats['total_cells'] > 0:
            stats['avg_text_len'] /= stats['total_cells']

        col_stats.append(stats)

    total_sample = len(sample_rows)
    if total_sample == 0:
        return {'desc_col': None, 'amount_col': None, 'qty_col': None, 'unit_price_col': None}

    # Identify column roles
    desc_col = None
    amount_col = None
    qty_col = None
    unit_price_col = None

    # Step 1: Find description column (longest average text, mostly alpha)
    best_desc_score = -1
    for i, s in enumerate(col_stats):
        if s['total_cells'] == 0:
            continue
        text_ratio = s['text_count'] / s['total_cells']
        # Description columns have high text ratio AND long average text
        score = text_ratio * s['avg_text_len']
        if text_ratio > 0.5 and s['avg_text_len'] > 5 and score > best_desc_score:
            best_desc_score = score
            desc_col = i

    # Step 2: Find dollar/amount columns (cells with $ or large numbers)
    dollar_cols = []
    for i, s in enumerate(col_stats):
        if i == desc_col:
            continue
        if s['total_cells'] == 0:
            continue
        dollar_ratio = s['dollar_count'] / s['total_cells']
        large_num_ratio = s['large_num_count'] / s['total_cells']
        if dollar_ratio > 0.4 or large_num_ratio > 0.5:
            dollar_cols.append((i, s['numeric_sum'], dollar_ratio + large_num_ratio))

    # Sort dollar columns by position (rightmost is usually total/extended)
    dollar_cols.sort(key=lambda x: x[0])

    if len(dollar_cols) >= 3:
        # 3+ dollar cols: likely retail, your price, total
        # Last = total/extended (amount), second-to-last = your price, first = retail/list
        amount_col = dollar_cols[-1][0]
        unit_price_col = dollar_cols[0][0]
    elif len(dollar_cols) == 2:
        # 2 dollar cols: unit price + extended/total
        # The one with larger values is the total
        if dollar_cols[1][1] >= dollar_cols[0][1]:
            amount_col = dollar_cols[1][0]
            unit_price_col = dollar_cols[0][0]
        else:
            amount_col = dollar_cols[0][0]
            unit_price_col = dollar_cols[1][0]
    elif len(dollar_cols) == 1:
        amount_col = dollar_cols[0][0]

    # Step 3: Find quantity column (small integers)
    for i, s in enumerate(col_stats):
        if i == desc_col or i == amount_col or i == unit_price_col:
            continue
        if s['total_cells'] == 0:
            continue
        small_int_ratio = s['small_int_count'] / s['total_cells']
        if small_int_ratio > 0.5:
            qty_col = i
            break

    logger.debug(f'[AEGIS ProposalParser] Column inference: desc={desc_col}, amount={amount_col}, '
                 f'qty={qty_col}, unit_price={unit_price_col} '
                 f'(from {num_cols} cols, {total_sample} sample rows)')

    return {
        'desc_col': desc_col,
        'amount_col': amount_col,
        'qty_col': qty_col,
        'unit_price_col': unit_price_col,
    }


def _headers_are_generic(headers: List[str]) -> bool:
    """Check if headers are generic auto-generated names (Col1, Col2, etc.)."""
    if not headers:
        return True
    generic_patterns = re.compile(r'^(col\s*\d+|column\s*\d+|\d+|unnamed.*|)$', re.IGNORECASE)
    generic_count = sum(1 for h in headers if generic_patterns.match(h.strip()))
    return generic_count > len(headers) * 0.6


def extract_line_items_from_table(table: ExtractedTable) -> List[LineItem]:
    """Extract financial line items from a table."""
    items = []
    if not table.rows:
        return items
    if not table.headers:
        # No headers at all — try data-pattern inference
        inferred = _infer_columns_from_data([], table.rows)
        if inferred['amount_col'] is not None:
            return _extract_with_inferred_columns(table, inferred)
        return items

    # Check if headers are generic (Col1, Col2...) — use data inference instead
    if _headers_are_generic(table.headers):
        inferred = _infer_columns_from_data(table.headers, table.rows)
        if inferred['amount_col'] is not None:
            logger.debug(f'[AEGIS ProposalParser] Using data-pattern inference for generic headers')
            return _extract_with_inferred_columns(table, inferred)

    # Identify column roles using tiered priority
    headers_lower = [h.lower().strip() for h in table.headers]

    desc_col = None
    amount_col = None
    qty_col = None
    unit_price_col = None

    # -- Description column: prioritize actual description keywords over identifiers --
    # Tier 1: Strong description indicators
    desc_priority_1 = ['description', 'task description', 'line item description',
                        'product description', 'description of services',
                        'service', 'deliverable', 'scope', 'product']
    # Tier 2: Medium description indicators
    desc_priority_2 = ['task', 'name', 'line item', 'wbs', 'wbs element', 'work',
                        'part description', 'item description']
    # Tier 3: Weak — identifiers (only use if nothing better found)
    desc_priority_3 = ['item', 'clin']

    for tier_keywords in [desc_priority_1, desc_priority_2, desc_priority_3]:
        if desc_col is not None:
            break
        for i, h in enumerate(headers_lower):
            if not h:
                continue
            if any(kw in h for kw in tier_keywords):
                desc_col = i
                break

    # -- Amount column: prefer 'total'/'extended' over generic 'cost'/'price' --
    # This prevents picking 'Unit Price' when 'Total Cost' exists
    amount_priority_1 = ['total cost', 'extended price', 'extended cost', 'total price',
                          'total amount', 'extended amount', 'total value', 'subtotal',
                          'your price', 'our price', 'proposed price', 'bid price',
                          'your cost', 'net price', 'net cost']
    amount_priority_2 = ['total', 'amount', 'extended', 'sum', 'value']
    amount_priority_3 = ['cost', 'price', 'charge']

    for tier_keywords in [amount_priority_1, amount_priority_2, amount_priority_3]:
        if amount_col is not None:
            break
        for i, h in enumerate(headers_lower):
            if not h:
                continue
            # Skip columns already assigned to other roles
            if i == desc_col:
                continue
            if any(kw in h for kw in tier_keywords):
                # Don't pick 'unit price' / 'unit cost' as the total amount column
                if any(kw in h for kw in ['unit price', 'unit cost', 'rate', 'per unit', 'hourly']):
                    continue
                amount_col = i
                break

    # -- Quantity column --
    for i, h in enumerate(headers_lower):
        if not h or i == desc_col or i == amount_col:
            continue
        if any(kw in h for kw in ['qty', 'quantity', 'hours', 'units', 'count', 'each', 'ea']):
            qty_col = i
            break

    # -- Unit price column --
    for i, h in enumerate(headers_lower):
        if not h or i == desc_col or i == amount_col or i == qty_col:
            continue
        if any(kw in h for kw in ['rate', 'unit price', 'unit cost', 'per unit', 'hourly',
                                   'retail price', 'retail', 'list price', 'list', 'msrp',
                                   'each price']):
            unit_price_col = i
            break

    # -- Fallback: if desc_col is an identifier column (item #, CLIN), try to find a better one --
    if desc_col is not None:
        desc_header = headers_lower[desc_col]
        is_id_col = any(kw == desc_header or (kw in desc_header and len(desc_header) < 12)
                        for kw in ['item', 'clin', '#', 'no', 'num', 'id', 'line'])
        if is_id_col:
            # Check if the NEXT column looks more like a description
            # (longer text values, more words)
            candidate = desc_col + 1
            if candidate < len(headers_lower) and candidate != amount_col and candidate != qty_col and candidate != unit_price_col:
                # Check if candidate column has longer text values
                sample_lengths = []
                for row in table.rows[:5]:
                    if candidate < len(row):
                        sample_lengths.append(len(str(row[candidate]).strip()))
                avg_len = sum(sample_lengths) / len(sample_lengths) if sample_lengths else 0
                if avg_len > 3:
                    desc_col = candidate

    # -- Fallback: if no description column found, use first non-empty text column --
    if desc_col is None:
        for i, h in enumerate(headers_lower):
            if h and i != amount_col and i != qty_col and i != unit_price_col:
                if not any(kw in h for kw in ['#', 'no', 'num', 'id']):
                    desc_col = i
                    break

    # -- Fallback: if no amount column, find column with most numeric values --
    if amount_col is None:
        col_numeric_counts = {}
        for row in table.rows:
            for j, cell in enumerate(row):
                if j == desc_col or j == qty_col:
                    continue
                val = _cell_to_float(cell)
                if val is not None and val > 0:
                    col_numeric_counts[j] = col_numeric_counts.get(j, 0) + 1
        if col_numeric_counts:
            amount_col = max(col_numeric_counts, key=col_numeric_counts.get)

    # -- Ultimate fallback: if keyword matching found nothing, try data-pattern inference --
    if amount_col is None:
        inferred = _infer_columns_from_data(table.headers, table.rows)
        if inferred['amount_col'] is not None:
            logger.debug('[AEGIS ProposalParser] Keyword matching failed, falling back to data-pattern inference')
            return _extract_with_inferred_columns(table, inferred)
        return items

    for row_idx, row in enumerate(table.rows):
        # Skip if this is a total/subtotal row
        row_text = ' '.join(str(c) for c in row)
        if TOTAL_KEYWORDS.search(row_text) and row_idx > len(table.rows) * 0.5:
            continue  # Skip totals in the bottom half

        # Get description
        desc = str(row[desc_col]).strip() if desc_col is not None and desc_col < len(row) else ''
        if not desc or len(desc) < 2:
            continue

        # Get amount — handle both dollar-formatted strings AND raw numbers
        amount_val = row[amount_col] if amount_col < len(row) else None
        amount = _cell_to_float(amount_val)
        if amount is None or amount == 0:
            continue

        # Build raw string for display
        amount_str = str(amount_val).strip() if amount_val is not None else ''
        if DOLLAR_PATTERN.search(amount_str):
            amount_raw = DOLLAR_PATTERN.search(amount_str).group()
        else:
            amount_raw = f'${amount:,.2f}'

        # Get optional fields
        qty = None
        if qty_col is not None and qty_col < len(row):
            qty = _cell_to_float(row[qty_col])

        unit_price = None
        if unit_price_col is not None and unit_price_col < len(row):
            unit_price = _cell_to_float(row[unit_price_col])

        item = LineItem(
            description=desc,
            amount=amount,
            amount_raw=amount_raw,
            quantity=qty,
            unit_price=unit_price,
            category=classify_line_item(desc),
            row_index=row_idx,
            table_index=table.table_index,
            source_sheet=table.source,
        )
        items.append(item)

    return items


def _extract_with_inferred_columns(table: ExtractedTable, inferred: Dict[str, Optional[int]]) -> List[LineItem]:
    """Extract line items using data-pattern inferred column roles."""
    items = []
    desc_col = inferred.get('desc_col')
    amount_col = inferred.get('amount_col')
    qty_col = inferred.get('qty_col')
    unit_price_col = inferred.get('unit_price_col')

    if amount_col is None:
        return items

    for row_idx, row in enumerate(table.rows):
        # Skip total/subtotal rows in bottom half
        row_text = ' '.join(str(c) for c in row)
        if TOTAL_KEYWORDS.search(row_text) and row_idx > len(table.rows) * 0.5:
            continue

        # Skip rows that are all empty or all dashes (separator rows)
        non_empty = [c for c in row if str(c).strip() and str(c).strip() not in ('-', '—', '–', '---')]
        if len(non_empty) < 2:
            continue

        # Get description
        desc = ''
        if desc_col is not None and desc_col < len(row):
            desc = str(row[desc_col]).strip()
        if not desc or len(desc) < 2:
            # Try to find the longest text cell in the row as fallback
            best_text = ''
            for ci, cell in enumerate(row):
                cell_str = str(cell).strip()
                if ci != amount_col and ci != qty_col and ci != unit_price_col:
                    if len(cell_str) > len(best_text) and sum(1 for c in cell_str if c.isalpha()) > 2:
                        best_text = cell_str
            desc = best_text
        if not desc or len(desc) < 2:
            continue

        # Get amount
        amount_val = row[amount_col] if amount_col < len(row) else None
        amount = _cell_to_float(amount_val)
        if amount is None or amount == 0:
            continue

        # Build raw string
        amount_str = str(amount_val).strip() if amount_val is not None else ''
        if DOLLAR_PATTERN.search(amount_str):
            amount_raw = DOLLAR_PATTERN.search(amount_str).group()
        else:
            amount_raw = f'${amount:,.2f}'

        # Optional fields
        qty = None
        if qty_col is not None and qty_col < len(row):
            qty = _cell_to_float(row[qty_col])

        unit_price = None
        if unit_price_col is not None and unit_price_col < len(row):
            unit_price = _cell_to_float(row[unit_price_col])

        item = LineItem(
            description=desc,
            amount=amount,
            amount_raw=amount_raw,
            quantity=qty,
            unit_price=unit_price,
            category=classify_line_item(desc),
            row_index=row_idx,
            table_index=table.table_index,
            source_sheet=table.source,
            confidence=0.8,  # Lower confidence for inferred columns
        )
        items.append(item)

    return items


# ──────────────────────────────────────────────
# Auto-Calculation Helper
# ──────────────────────────────────────────────

def _auto_calculate_line_items(items: list) -> int:
    """Auto-calculate missing financial fields when 2 of 3 (qty, unit_price, amount) are present.

    Returns the number of items that were auto-calculated.
    """
    calculated = 0
    for item in items:
        qty = item.quantity
        unit = item.unit_price
        amt = item.amount

        if qty and unit and not amt:
            item.amount = round(qty * unit, 2)
            item.amount_raw = f'${item.amount:,.2f}'
            calculated += 1
        elif amt and qty and not unit:
            item.unit_price = round(amt / qty, 2)
            calculated += 1
        elif amt and unit and not qty:
            item.quantity = round(amt / unit, 2)
            calculated += 1

    return calculated


# ──────────────────────────────────────────────
# Excel Parser
# ──────────────────────────────────────────────

def parse_excel(filepath: str) -> ProposalData:
    """Parse an Excel file (.xlsx/.xls) for proposal financial data."""
    import openpyxl

    proposal = ProposalData(
        filename=os.path.basename(filepath),
        filepath=filepath,
        file_type='xlsx',
    )

    try:
        wb = openpyxl.load_workbook(filepath, data_only=True, read_only=True)
    except Exception as e:
        proposal.extraction_notes.append(f'Failed to open Excel file: {e}')
        return proposal

    all_text_parts = []

    for sheet_idx, sheet_name in enumerate(wb.sheetnames):
        ws = wb[sheet_name]

        # Read all rows
        all_rows = []
        for row in ws.iter_rows(values_only=True):
            str_row = [str(cell) if cell is not None else '' for cell in row]
            all_rows.append(str_row)
            all_text_parts.append(' '.join(str_row))

        if not all_rows:
            continue

        # Find header row (first row with 3+ non-empty cells)
        header_idx = 0
        for i, row in enumerate(all_rows[:10]):
            non_empty = sum(1 for c in row if c.strip())
            if non_empty >= 3:
                header_idx = i
                break

        headers = all_rows[header_idx] if header_idx < len(all_rows) else []
        data_rows = all_rows[header_idx + 1:]

        # Filter out completely empty rows
        data_rows = [r for r in data_rows if any(c.strip() for c in r)]

        if not data_rows:
            continue

        table = ExtractedTable(
            headers=headers,
            rows=data_rows,
            source=sheet_name,
            table_index=sheet_idx,
        )

        # Check if financial
        table.has_financial_data = is_financial_table(headers, data_rows)

        # Find total row
        total_idx, total_amount = find_total_row(data_rows)
        table.total_row_index = total_idx

        proposal.tables.append(table)

        # Extract line items from financial tables
        if table.has_financial_data:
            items = extract_line_items_from_table(table)
            proposal.line_items.extend(items)

            if total_amount is not None and (proposal.total_amount is None or total_amount > (proposal.total_amount or 0)):
                proposal.total_amount = total_amount
                proposal.total_raw = f'${total_amount:,.2f}'

    wb.close()

    # Extract company name and date from all text
    full_text = '\n'.join(all_text_parts)
    if not proposal.company_name:
        proposal.company_name = extract_company_from_text(full_text, filename=proposal.filename)
    if not proposal.date:
        proposal.date = extract_dates_from_text(full_text)

    # Contract term detection (from text + sheet tab names)
    proposal.contract_term = extract_contract_term(full_text, sheet_names=list(wb.sheetnames))
    if proposal.contract_term:
        proposal.extraction_notes.append(f'Contract term: {proposal.contract_term}')

    # Auto-calculate missing fields (qty * unit_price = amount, etc.)
    auto_calced = _auto_calculate_line_items(proposal.line_items)
    if auto_calced:
        proposal.extraction_notes.append(f'Auto-calculated {auto_calced} missing financial fields')

    # If no total found from total rows, sum line items
    if proposal.total_amount is None and proposal.line_items:
        proposal.total_amount = sum(li.amount for li in proposal.line_items if li.amount)
        proposal.total_raw = f'${proposal.total_amount:,.2f}'

    proposal.extraction_text = full_text
    proposal.extraction_notes.append(f'Found {len(proposal.tables)} sheets, {len(proposal.line_items)} line items')

    return proposal


# ──────────────────────────────────────────────
# DOCX Parser
# ──────────────────────────────────────────────

def parse_docx(filepath: str) -> ProposalData:
    """Parse a DOCX file for proposal financial data."""
    from docx import Document as DocxDocument

    proposal = ProposalData(
        filename=os.path.basename(filepath),
        filepath=filepath,
        file_type='docx',
    )

    try:
        doc = DocxDocument(filepath)
    except Exception as e:
        proposal.extraction_notes.append(f'Failed to open DOCX file: {e}')
        return proposal

    # Extract full text for company/date detection
    full_text_parts = []
    for para in doc.paragraphs:
        full_text_parts.append(para.text)

    # Extract title from first heading or large text
    for para in doc.paragraphs[:50]:
        if para.style and 'heading' in (para.style.name or '').lower():
            if not proposal.proposal_title and len(para.text.strip()) > 5:
                proposal.proposal_title = para.text.strip()
                break

    # Extract tables
    for tbl_idx, table in enumerate(doc.tables):
        headers = []
        rows = []

        for row_idx, row in enumerate(table.rows):
            cells = [cell.text.strip() for cell in row.cells]
            if row_idx == 0:
                headers = cells
            else:
                rows.append(cells)

        if not headers and not rows:
            continue

        # Filter empty rows
        rows = [r for r in rows if any(c.strip() for c in r)]

        ext_table = ExtractedTable(
            headers=headers,
            rows=rows,
            source=f'Table {tbl_idx + 1}',
            table_index=tbl_idx,
        )

        ext_table.has_financial_data = is_financial_table(headers, rows)
        total_idx, total_amount = find_total_row(rows)
        ext_table.total_row_index = total_idx

        proposal.tables.append(ext_table)

        if ext_table.has_financial_data:
            items = extract_line_items_from_table(ext_table)
            proposal.line_items.extend(items)

            if total_amount is not None and (proposal.total_amount is None or total_amount > (proposal.total_amount or 0)):
                proposal.total_amount = total_amount
                proposal.total_raw = f'${total_amount:,.2f}'

    # Also scan paragraphs for inline dollar amounts (not in tables)
    inline_amounts = []
    full_text = '\n'.join(full_text_parts)
    for match in DOLLAR_PATTERN.finditer(full_text):
        amount = parse_dollar_amount(match.group())
        if amount and amount > 50:  # Capture amounts $50+
            # Get surrounding context
            start = max(0, match.start() - 80)
            end = min(len(full_text), match.end() + 40)
            context = full_text[start:end].replace('\n', ' ').strip()
            inline_amounts.append((amount, match.group(), context))

    if inline_amounts:
        existing_amounts = {li.amount for li in proposal.line_items if li.amount}
        added_inline = 0
        for idx, (amount, raw, context) in enumerate(inline_amounts[:50]):
            if amount in existing_amounts:
                continue
            proposal.line_items.append(LineItem(
                description=context,
                amount=amount,
                amount_raw=raw,
                category=classify_line_item(context),
                row_index=idx,
                source_sheet='Body Text' if not proposal.tables else 'Body Text Supplementary',
                confidence=0.7 if not proposal.tables else 0.5,
            ))
            added_inline += 1
        if added_inline:
            proposal.extraction_notes.append(f'Found {added_inline} additional dollar amounts in body text')

    # Company and date
    if not proposal.company_name:
        proposal.company_name = extract_company_from_text(full_text, filename=proposal.filename)
    if not proposal.date:
        proposal.date = extract_dates_from_text(full_text)

    # Contract term detection
    proposal.contract_term = extract_contract_term(full_text)
    if proposal.contract_term:
        proposal.extraction_notes.append(f'Contract term: {proposal.contract_term}')

    # Auto-calculate missing fields (qty * unit_price = amount, etc.)
    auto_calced = _auto_calculate_line_items(proposal.line_items)
    if auto_calced:
        proposal.extraction_notes.append(f'Auto-calculated {auto_calced} missing financial fields')

    # Total
    if proposal.total_amount is None and proposal.line_items:
        proposal.total_amount = sum(li.amount for li in proposal.line_items if li.amount)
        proposal.total_raw = f'${proposal.total_amount:,.2f}'

    proposal.page_count = len(doc.paragraphs) // 30  # Rough estimate
    proposal.extraction_text = full_text
    proposal.extraction_notes.append(f'Found {len(doc.tables)} tables, {len(proposal.line_items)} line items')

    return proposal


# ──────────────────────────────────────────────
# PDF Parser
# ──────────────────────────────────────────────

def parse_pdf(filepath: str) -> ProposalData:
    """Parse a PDF file for proposal financial data.

    Extraction strategy (best to worst):
    1. EnhancedTableExtractor (camelot lattice → camelot stream → tabula → pdfplumber)
    2. pymupdf4llm Markdown tables + fitz find_tables() (legacy fallback)
    3. pdfplumber text + inline dollar extraction (final fallback)
    """
    proposal = ProposalData(
        filename=os.path.basename(filepath),
        filepath=filepath,
        file_type='pdf',
    )

    full_text = ''  # For company/date extraction

    # ── Always extract full text for company name/date (independent of table extraction) ──
    try:
        import pdfplumber
        with pdfplumber.open(filepath) as pdf:
            if not proposal.page_count:
                proposal.page_count = len(pdf.pages)
            text_parts = []
            for page in pdf.pages[:50]:
                text = page.extract_text()
                if text:
                    text_parts.append(text)
            full_text = '\n'.join(text_parts)
    except ImportError:
        pass
    except Exception as e:
        logger.debug(f'[AEGIS ProposalParser] pdfplumber text pre-extraction failed: {e}')

    if not full_text:
        try:
            import fitz
            doc = fitz.open(filepath)
            text_parts = []
            for page in doc[:50]:
                text = page.get_text()
                if text:
                    text_parts.append(text)
            full_text = '\n'.join(text_parts)
            doc.close()
        except ImportError:
            pass
        except Exception as e:
            logger.debug(f'[AEGIS ProposalParser] fitz text pre-extraction failed: {e}')

    # Extract company name, date, and contract term from text BEFORE table extraction
    if full_text:
        proposal.company_name = extract_company_from_text(full_text, filename=proposal.filename)
        proposal.date = extract_dates_from_text(full_text)
        proposal.contract_term = extract_contract_term(full_text)
        if proposal.company_name:
            proposal.extraction_notes.append(f'Company: {proposal.company_name}')
        if proposal.contract_term:
            proposal.extraction_notes.append(f'Contract term: {proposal.contract_term}')
    elif proposal.filename:
        # No text extracted yet — try filename as last resort for company name
        proposal.company_name = extract_company_from_text('', filename=proposal.filename)

    # ── Strategy 1: EnhancedTableExtractor (primary — best accuracy) ──
    enhanced_success = False
    try:
        import sys
        # enhanced_table_extractor.py lives in the project root
        project_root = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
        if project_root not in sys.path:
            sys.path.insert(0, project_root)

        from enhanced_table_extractor import EnhancedTableExtractor, ExtractionResult as ETResult

        logger.debug(f'[AEGIS ProposalParser] Trying EnhancedTableExtractor for {os.path.basename(filepath)}')
        extractor = EnhancedTableExtractor(prefer_accuracy=True)
        et_result = extractor.extract_tables(filepath)

        proposal.page_count = et_result.total_pages

        if et_result.tables:
            enhanced_success = True
            extraction_method = et_result.extraction_method
            proposal.extraction_notes.append(
                f'EnhancedTableExtractor: {len(et_result.tables)} tables via {extraction_method} '
                f'({et_result.extraction_time_ms:.0f}ms)'
            )

            for et in et_result.tables:
                ext_table = ExtractedTable(
                    headers=et.headers,
                    rows=et.rows,
                    source=f'Page {et.page} ({et.source})',
                    table_index=et.index,
                )

                # Check if financial
                ext_table.has_financial_data = is_financial_table(et.headers, et.rows)

                # Also check via data patterns for headerless tables
                if not ext_table.has_financial_data and _headers_are_generic(et.headers):
                    # Check if rows contain dollar amounts
                    dollar_cell_count = 0
                    total_cells = 0
                    for row in et.rows[:10]:
                        for cell in row:
                            total_cells += 1
                            if DOLLAR_PATTERN.search(str(cell)):
                                dollar_cell_count += 1
                    if total_cells > 0 and dollar_cell_count / total_cells > 0.1:
                        ext_table.has_financial_data = True

                # Find total row
                total_idx, total_amount = find_total_row(et.rows)
                ext_table.total_row_index = total_idx

                proposal.tables.append(ext_table)

                if ext_table.has_financial_data:
                    items = extract_line_items_from_table(ext_table)
                    proposal.line_items.extend(items)
                    logger.debug(f'[AEGIS ProposalParser] Table {et.index} ({et.source}): '
                                 f'{len(items)} line items, headers_generic={_headers_are_generic(et.headers)}')

                    if total_amount is not None and (proposal.total_amount is None or total_amount > (proposal.total_amount or 0)):
                        proposal.total_amount = total_amount
                        proposal.total_raw = f'${total_amount:,.2f}'

            if et_result.warnings:
                for w in et_result.warnings:
                    proposal.extraction_notes.append(f'Warning: {w}')

        else:
            proposal.extraction_notes.append('EnhancedTableExtractor found no tables')

    except ImportError as ie:
        proposal.extraction_notes.append(f'EnhancedTableExtractor not available: {ie}')
        logger.debug(f'[AEGIS ProposalParser] EnhancedTableExtractor import failed: {ie}')
    except Exception as e:
        proposal.extraction_notes.append(f'EnhancedTableExtractor error: {e}')
        logger.warning(f'[AEGIS ProposalParser] EnhancedTableExtractor failed: {e}', exc_info=True)

    # ── Strategy 2: pymupdf4llm + fitz (fallback if enhanced found nothing) ──
    if not proposal.line_items:
        try:
            import pymupdf4llm
            import fitz

            doc = fitz.open(filepath)
            if not proposal.page_count:
                proposal.page_count = len(doc)

            # Get full text via pymupdf4llm (Markdown with tables)
            md_text = pymupdf4llm.to_markdown(filepath)
            full_text = re.sub(r'[#*|_\-]+', ' ', md_text)

            # Extract tables from Markdown
            tables = _extract_tables_from_markdown(md_text)

            if tables:
                proposal.extraction_notes.append(f'pymupdf4llm fallback: {len(tables)} Markdown tables')

            for tbl_idx, (headers, rows) in enumerate(tables):
                tbl_index = len(proposal.tables)
                ext_table = ExtractedTable(
                    headers=headers,
                    rows=rows,
                    source=f'Page (Markdown)',
                    table_index=tbl_index,
                )
                ext_table.has_financial_data = is_financial_table(headers, rows)
                total_idx, total_amount = find_total_row(rows)
                ext_table.total_row_index = total_idx

                proposal.tables.append(ext_table)

                if ext_table.has_financial_data:
                    items = extract_line_items_from_table(ext_table)
                    proposal.line_items.extend(items)

                    if total_amount is not None and (proposal.total_amount is None or total_amount > (proposal.total_amount or 0)):
                        proposal.total_amount = total_amount
                        proposal.total_raw = f'${total_amount:,.2f}'

            # Also try fitz table extraction
            for page_num in range(len(doc)):
                page = doc[page_num]
                page_tables = page.find_tables()
                if page_tables and page_tables.tables:
                    for ft_idx, ft in enumerate(page_tables.tables):
                        try:
                            extracted = ft.extract()
                            if not extracted or len(extracted) < 2:
                                continue

                            headers = [str(c) if c else '' for c in extracted[0]]
                            rows = [[str(c) if c else '' for c in r] for r in extracted[1:]]
                            rows = [r for r in rows if any(c.strip() for c in r)]

                            if not rows:
                                continue

                            # Dedup check
                            h_key = '|'.join(h.strip().lower() for h in headers if h.strip())
                            existing_keys = set()
                            for et in proposal.tables:
                                existing_keys.add('|'.join(h.strip().lower() for h in et.headers if h.strip()))
                            if h_key in existing_keys:
                                continue

                            tbl_index = len(proposal.tables)
                            ext_table = ExtractedTable(
                                headers=headers,
                                rows=rows,
                                source=f'Page {page_num + 1}',
                                table_index=tbl_index,
                            )
                            ext_table.has_financial_data = is_financial_table(headers, rows)
                            total_idx, total_amount = find_total_row(rows)
                            ext_table.total_row_index = total_idx

                            proposal.tables.append(ext_table)

                            if ext_table.has_financial_data:
                                items = extract_line_items_from_table(ext_table)
                                proposal.line_items.extend(items)

                                if total_amount is not None and (proposal.total_amount is None or total_amount > (proposal.total_amount or 0)):
                                    proposal.total_amount = total_amount
                                    proposal.total_raw = f'${total_amount:,.2f}'
                        except Exception:
                            continue

            doc.close()

        except ImportError:
            proposal.extraction_notes.append('pymupdf4llm not available')
        except Exception as e:
            proposal.extraction_notes.append(f'pymupdf4llm fallback error: {e}')
            logger.warning(f'[AEGIS ProposalParser] pymupdf4llm fallback failed: {e}')

    # ── Inline dollar amounts from text (always capture, dedup against table amounts) ──
    if full_text:
        inline_amounts = []
        for match in DOLLAR_PATTERN.finditer(full_text):
            amount = parse_dollar_amount(match.group())
            if amount and amount > 10:
                start = max(0, match.start() - 80)
                end = min(len(full_text), match.end() + 40)
                context = full_text[start:end].replace('\n', ' ').strip()
                inline_amounts.append((amount, match.group(), context))

        if inline_amounts:
            existing_amounts = {li.amount for li in proposal.line_items if li.amount}
            added_inline = 0
            for idx, (amount, raw, context) in enumerate(inline_amounts[:50]):
                if amount in existing_amounts:
                    continue
                proposal.line_items.append(LineItem(
                    description=context,
                    amount=amount,
                    amount_raw=raw,
                    category=classify_line_item(context),
                    row_index=idx,
                    source_sheet='PDF Text' if not proposal.tables else 'PDF Text Supplementary',
                    confidence=0.5 if not proposal.tables else 0.4,
                ))
                added_inline += 1
            if added_inline:
                proposal.extraction_notes.append(f'Found {added_inline} additional dollar amounts in PDF text')

    # Auto-calculate missing fields (qty * unit_price = amount, etc.)
    auto_calced = _auto_calculate_line_items(proposal.line_items)
    if auto_calced:
        proposal.extraction_notes.append(f'Auto-calculated {auto_calced} missing financial fields')

    # ── Compute total ──
    if proposal.total_amount is None and proposal.line_items:
        proposal.total_amount = sum(li.amount for li in proposal.line_items if li.amount)
        proposal.total_raw = f'${proposal.total_amount:,.2f}'

    proposal.extraction_text = full_text
    proposal.extraction_notes.append(f'Found {len(proposal.tables)} tables, {len(proposal.line_items)} line items')

    return proposal


def _extract_tables_from_markdown(md_text: str) -> List[Tuple[List[str], List[List[str]]]]:
    """Extract tables from Markdown text (pymupdf4llm output)."""
    tables = []
    lines = md_text.split('\n')

    i = 0
    while i < len(lines):
        line = lines[i].strip()

        # Detect Markdown table: line with | separators
        if '|' in line and line.startswith('|'):
            # Potential table start
            header_cells = [c.strip() for c in line.split('|')]
            header_cells = [c for c in header_cells if c]  # Remove empty from split

            # Check for separator row (---|---|---)
            if i + 1 < len(lines) and re.match(r'\|[\s\-:|]+\|', lines[i + 1].strip()):
                i += 2  # Skip header + separator

                rows = []
                while i < len(lines) and '|' in lines[i] and lines[i].strip().startswith('|'):
                    row_cells = [c.strip() for c in lines[i].strip().split('|')]
                    row_cells = [c for c in row_cells if c or len(row_cells) > 3]
                    # Clean: remove first/last empty from | split
                    if row_cells and not row_cells[0]:
                        row_cells = row_cells[1:]
                    if row_cells and not row_cells[-1]:
                        row_cells = row_cells[:-1]
                    if row_cells:
                        rows.append(row_cells)
                    i += 1

                if rows:
                    tables.append((header_cells, rows))
                continue

        i += 1

    return tables


# ──────────────────────────────────────────────
# Main entry point
# ──────────────────────────────────────────────

SUPPORTED_EXTENSIONS = {'.xlsx', '.xls', '.docx', '.pdf'}

def parse_proposal(filepath: str) -> ProposalData:
    """Parse a proposal document and extract financial data.

    Supports: .xlsx, .xls, .docx, .pdf
    Returns a ProposalData object with extracted tables, line items, and totals.
    """
    ext = os.path.splitext(filepath)[1].lower()

    if ext in ('.xlsx', '.xls'):
        return parse_excel(filepath)
    elif ext == '.docx':
        return parse_docx(filepath)
    elif ext == '.pdf':
        return parse_pdf(filepath)
    else:
        data = ProposalData(
            filename=os.path.basename(filepath),
            filepath=filepath,
            file_type=ext.lstrip('.'),
        )
        data.extraction_notes.append(f'Unsupported file type: {ext}')
        return data
