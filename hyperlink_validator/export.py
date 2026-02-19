"""
Hyperlink Validator Export Functions
====================================
Export validation results to various formats: CSV, JSON, HTML, and highlighted
copies of original documents (DOCX, Excel).

This module is designed to be independent and can be tested separately.

Exclusion Handling:
- URLs marked as excluded with treat_as_valid=True are shown as WORKING/OK
- The exclusion_reason field indicates why the URL was excluded

Highlighted Export (v3.0.110):
- export_highlighted_docx: Creates DOCX with broken links highlighted in red
- export_highlighted_excel: Creates Excel with rows containing broken links in red
"""

import csv
import json
import io
import os
import re
import zipfile
import copy
import html as html_lib  # For HTML escaping
import xml.etree.ElementTree as ET
from datetime import datetime
from typing import List, Optional, Any, Dict, Set, Tuple

from .models import ValidationResult, ValidationSummary, ValidationRun

# Check for python-docx availability
try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_COLOR_INDEX
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# Check for openpyxl availability
try:
    from openpyxl import load_workbook
    from openpyxl.styles import PatternFill, Font, Border, Side, Alignment
    from openpyxl.comments import Comment
    from openpyxl.utils import get_column_letter
    OPENPYXL_AVAILABLE = True
except ImportError:
    OPENPYXL_AVAILABLE = False


def _apply_exclusion_display(result: ValidationResult) -> ValidationResult:
    """
    Apply exclusion display rules: excluded URLs with treat_as_valid=True
    should show as WORKING in exports.

    Returns a copy of the result with adjusted status if applicable.
    """
    if result.excluded and getattr(result, 'treat_as_valid', True):
        # Create a modified copy for display
        modified = ValidationResult(
            url=result.url,
            status='WORKING',  # Show as OK
            status_code=200,  # Synthetic OK code
            message=f'Excluded: {result.exclusion_reason}' if result.exclusion_reason else 'Excluded (treated as valid)',
            redirect_url=result.redirect_url,
            redirect_count=result.redirect_count,
            response_time_ms=0,
            dns_resolved=True,
            ssl_valid=True,
            auth_used=result.auth_used,
            attempts=0,
            checked_at=result.checked_at,
            # Extended fields
            dns_ip_addresses=result.dns_ip_addresses,
            ssl_issuer=result.ssl_issuer,
            ssl_days_until_expiry=result.ssl_days_until_expiry,
            is_soft_404=False,
            is_suspicious=False,
            domain_category=result.domain_category,
            excluded=True,
            exclusion_reason=result.exclusion_reason
        )
        return modified
    return result


def export_csv(
    results: List[ValidationResult],
    summary: Optional[ValidationSummary] = None,
    apply_exclusion_rules: bool = True
) -> str:
    """
    Export results to CSV format.

    Args:
        results: List of validation results
        summary: Optional summary statistics
        apply_exclusion_rules: If True, excluded URLs show as WORKING

    Returns:
        CSV content as string
    """
    output = io.StringIO()

    # Use UTF-8 BOM for Excel compatibility
    output.write('\ufeff')

    writer = csv.writer(output, quoting=csv.QUOTE_MINIMAL)

    # Header row
    writer.writerow([
        'URL',
        'Status',
        'Status Code',
        'Message',
        'Redirect URL',
        'Redirect Count',
        'Response Time (ms)',
        'DNS Resolved',
        'SSL Valid',
        'Auth Used',
        'Attempts',
        'Excluded',
        'Checked At'
    ])

    # Data rows
    for result in results:
        # Apply exclusion display rules
        display_result = _apply_exclusion_display(result) if apply_exclusion_rules else result

        writer.writerow([
            display_result.url,
            display_result.status,
            display_result.status_code or '',
            display_result.message,
            display_result.redirect_url or '',
            display_result.redirect_count,
            f'{display_result.response_time_ms:.1f}' if display_result.response_time_ms else '',
            'Yes' if display_result.dns_resolved else 'No',
            'Yes' if display_result.ssl_valid else 'No',
            display_result.auth_used,
            display_result.attempts,
            'Yes' if display_result.excluded else 'No',
            display_result.checked_at
        ])

    # Add summary section
    if summary:
        writer.writerow([])  # Empty row
        writer.writerow(['SUMMARY'])
        writer.writerow(['Total URLs', summary.total])
        writer.writerow(['Working', summary.working])
        writer.writerow(['Broken', summary.broken])
        writer.writerow(['Redirect', summary.redirect])
        writer.writerow(['Timeout', summary.timeout])
        writer.writerow(['Blocked', summary.blocked])
        writer.writerow(['DNS Failed', summary.dns_failed])
        writer.writerow(['SSL Error', summary.ssl_error])
        writer.writerow(['Invalid', summary.invalid])
        writer.writerow(['Unknown', summary.unknown])
        writer.writerow(['Success Rate', f'{summary.success_rate:.1f}%'])
        writer.writerow(['Average Response', f'{summary.average_response_ms:.1f}ms'])
        writer.writerow(['Total Time', f'{summary.total_time_seconds:.1f}s'])

    return output.getvalue()


def export_json(
    results: List[ValidationResult],
    summary: Optional[ValidationSummary] = None,
    run: Optional[ValidationRun] = None,
    apply_exclusion_rules: bool = True
) -> str:
    """
    Export results to JSON format.

    Args:
        results: List of validation results
        summary: Optional summary statistics
        run: Optional run metadata
        apply_exclusion_rules: If True, excluded URLs show as WORKING

    Returns:
        JSON content as string
    """
    # Apply exclusion display rules
    display_results = [
        _apply_exclusion_display(r) if apply_exclusion_rules else r
        for r in results
    ]

    data = {
        'exported_at': datetime.utcnow().isoformat() + 'Z',
        'exporter': 'AEGIS HyperlinkValidator v1.0.0',
        'results': [r.to_dict() for r in display_results]
    }

    if summary:
        data['summary'] = summary.to_dict()

    if run:
        data['run'] = {
            'run_id': run.run_id,
            'job_id': run.job_id,
            'created_at': run.created_at,
            'completed_at': run.completed_at,
            'mode': run.mode,
            'status': run.status
        }

    return json.dumps(data, indent=2)


def export_html(
    results: List[ValidationResult],
    summary: Optional[ValidationSummary] = None,
    run: Optional[ValidationRun] = None,
    apply_exclusion_rules: bool = True
) -> str:
    """
    Export results to HTML report format.

    Args:
        results: List of validation results
        summary: Optional summary statistics
        run: Optional run metadata
        apply_exclusion_rules: If True, excluded URLs show as WORKING

    Returns:
        HTML content as string
    """
    # Apply exclusion display rules
    display_results = [
        _apply_exclusion_display(r) if apply_exclusion_rules else r
        for r in results
    ]

    # Get status color class
    def status_class(status: str, excluded: bool = False) -> str:
        if excluded:
            return 'status-excluded'  # Special styling for excluded items
        status_map = {
            'WORKING': 'status-working',
            'REDIRECT': 'status-redirect',
            'BROKEN': 'status-broken',
            'TIMEOUT': 'status-timeout',
            'BLOCKED': 'status-blocked',
            'DNSFAILED': 'status-dns',
            'SSLERROR': 'status-ssl',
            'INVALID': 'status-invalid',
            'UNKNOWN': 'status-unknown',
            'PENDING': 'status-pending',
            'SKIPPED': 'status-skipped'
        }
        return status_map.get(status.upper(), 'status-unknown')

    # Build HTML
    html = f'''<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Hyperlink Validation Report</title>
    <style>
        :root {{
            --color-working: #22c55e;
            --color-redirect: #3b82f6;
            --color-broken: #ef4444;
            --color-timeout: #f59e0b;
            --color-blocked: #8b5cf6;
            --color-dns: #ec4899;
            --color-ssl: #f97316;
            --color-invalid: #6b7280;
            --color-unknown: #9ca3af;
        }}

        * {{
            box-sizing: border-box;
            margin: 0;
            padding: 0;
        }}

        body {{
            font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, 'Helvetica Neue', Arial, sans-serif;
            line-height: 1.6;
            color: #1f2937;
            background: #f9fafb;
            padding: 2rem;
        }}

        .container {{
            max-width: 1400px;
            margin: 0 auto;
        }}

        h1 {{
            font-size: 1.875rem;
            font-weight: 700;
            margin-bottom: 0.5rem;
            color: #111827;
        }}

        .subtitle {{
            color: #6b7280;
            margin-bottom: 2rem;
        }}

        .summary-grid {{
            display: grid;
            grid-template-columns: repeat(auto-fit, minmax(150px, 1fr));
            gap: 1rem;
            margin-bottom: 2rem;
        }}

        .summary-card {{
            background: white;
            border-radius: 8px;
            padding: 1rem;
            box-shadow: 0 1px 3px rgba(0, 0, 0, 0.1);
            text-align: center;
        }}

        .summary-card .value {{
            font-size: 2rem;
            font-weight: 700;
            margin-bottom: 0.25rem;
        }}

        .summary-card .label {{
            font-size: 0.875rem;
            color: #6b7280;
        }}

        .summary-card.working .value {{ color: var(--color-working); }}
        .summary-card.broken .value {{ color: var(--color-broken); }}
        .summary-card.redirect .value {{ color: var(--color-redirect); }}
        .summary-card.timeout .value {{ color: var(--color-timeout); }}
        .summary-card.blocked .value {{ color: var(--color-blocked); }}

        .results-table {{
            width: 100%;
            background: white;
            border-radius: 8px;
            box-shadow: 0 1px 3px rgba(0, 0, 0, 0.1);
            overflow: hidden;
        }}

        .results-table table {{
            width: 100%;
            border-collapse: collapse;
        }}

        .results-table th {{
            background: #f3f4f6;
            padding: 0.75rem 1rem;
            text-align: left;
            font-weight: 600;
            font-size: 0.875rem;
            color: #374151;
            border-bottom: 1px solid #e5e7eb;
        }}

        .results-table td {{
            padding: 0.75rem 1rem;
            border-bottom: 1px solid #f3f4f6;
            font-size: 0.875rem;
        }}

        .results-table tr:hover {{
            background: #f9fafb;
        }}

        .results-table .url {{
            max-width: 400px;
            word-break: break-all;
        }}

        .results-table .url a {{
            color: #2563eb;
            text-decoration: none;
        }}

        .results-table .url a:hover {{
            text-decoration: underline;
        }}

        .status-badge {{
            display: inline-flex;
            align-items: center;
            padding: 0.25rem 0.75rem;
            border-radius: 9999px;
            font-size: 0.75rem;
            font-weight: 600;
            text-transform: uppercase;
        }}

        .status-working {{ background: #dcfce7; color: #166534; }}
        .status-redirect {{ background: #dbeafe; color: #1e40af; }}
        .status-broken {{ background: #fee2e2; color: #991b1b; }}
        .status-timeout {{ background: #fef3c7; color: #92400e; }}
        .status-blocked {{ background: #ede9fe; color: #5b21b6; }}
        .status-dns {{ background: #fce7f3; color: #9d174d; }}
        .status-ssl {{ background: #ffedd5; color: #c2410c; }}
        .status-invalid {{ background: #f3f4f6; color: #374151; }}
        .status-unknown {{ background: #f3f4f6; color: #6b7280; }}
        .status-pending {{ background: #e0e7ff; color: #3730a3; }}
        .status-skipped {{ background: #f3f4f6; color: #9ca3af; }}
        .status-excluded {{ background: #d1fae5; color: #065f46; border: 1px dashed #10b981; }}

        .meta-info {{
            margin-top: 2rem;
            padding: 1rem;
            background: #f3f4f6;
            border-radius: 8px;
            font-size: 0.75rem;
            color: #6b7280;
        }}

        @media print {{
            body {{
                background: white;
                padding: 0;
            }}

            .results-table {{
                box-shadow: none;
                border: 1px solid #e5e7eb;
            }}
        }}
    </style>
</head>
<body>
    <div class="container">
        <h1>Hyperlink Validation Report</h1>
        <p class="subtitle">Generated: {datetime.utcnow().strftime('%Y-%m-%d %H:%M:%S')} UTC</p>
'''

    # v5.0.5: Always calculate summary from display_results to ensure consistency
    # between the summary cards and the results table. The original summary object
    # uses pre-exclusion statuses, but display_results has exclusion rules applied,
    # so the counts would mismatch if we used the original summary directly.
    total = len(display_results)
    working = sum(1 for r in display_results if r.status and r.status.upper() == 'WORKING')
    broken = sum(1 for r in display_results if r.status and r.status.upper() == 'BROKEN')
    redirect = sum(1 for r in display_results if r.status and r.status.upper() == 'REDIRECT')
    timeout = sum(1 for r in display_results if r.status and r.status.upper() == 'TIMEOUT')
    auth_req = sum(1 for r in display_results if r.status and r.status.upper() == 'AUTH_REQUIRED')
    other = total - working - broken - redirect - timeout - auth_req

    html += f'''
        <div class="summary-grid">
            <div class="summary-card">
                <div class="value">{total}</div>
                <div class="label">Total URLs</div>
            </div>
            <div class="summary-card working">
                <div class="value">{working}</div>
                <div class="label">Working</div>
            </div>
            <div class="summary-card broken">
                <div class="value">{broken}</div>
                <div class="label">Broken</div>
            </div>
            <div class="summary-card redirect">
                <div class="value">{redirect}</div>
                <div class="label">Redirect</div>
            </div>
            <div class="summary-card timeout">
                <div class="value">{timeout}</div>
                <div class="label">Timeout</div>
            </div>
            <div class="summary-card blocked">
                <div class="value">{auth_req}</div>
                <div class="label">Auth Required</div>
            </div>
            <div class="summary-card">
                <div class="value">{other}</div>
                <div class="label">Other</div>
            </div>
        </div>
'''

    # Results table
    html += '''
        <div class="results-table">
            <table>
                <thead>
                    <tr>
                        <th>Status</th>
                        <th>URL</th>
                        <th>Code</th>
                        <th>Message</th>
                        <th>Time</th>
                    </tr>
                </thead>
                <tbody>
'''

    for result in display_results:
        # BUG-M09 fix: Properly handle None values and escape HTML
        time_str = f'{result.response_time_ms:.0f}ms' if result.response_time_ms else '-'
        code_str = str(result.status_code) if result.status_code else '-'
        is_excluded = getattr(result, 'excluded', False)

        # Safely escape URL and message for HTML
        safe_url = html_lib.escape(result.url or '', quote=True) if result.url else ''
        safe_message = html_lib.escape(result.message or '-', quote=True) if result.message else '-'
        safe_status = html_lib.escape(str(result.status) if result.status else 'UNKNOWN', quote=True)

        html += f'''
                    <tr>
                        <td><span class="status-badge {status_class(result.status or 'UNKNOWN', is_excluded)}">{safe_status}{' (Excluded)' if is_excluded else ''}</span></td>
                        <td class="url"><a href="{safe_url}" target="_blank" rel="noopener">{safe_url}</a></td>
                        <td>{code_str}</td>
                        <td>{safe_message}</td>
                        <td>{time_str}</td>
                    </tr>
'''

    html += '''
                </tbody>
            </table>
        </div>
'''

    # Meta info
    if run:
        html += f'''
        <div class="meta-info">
            <strong>Run ID:</strong> {run.run_id} |
            <strong>Mode:</strong> {run.mode} |
            <strong>Started:</strong> {run.created_at} |
            <strong>Completed:</strong> {run.completed_at or 'N/A'}
'''
        if summary:
            # v5.0.5: Calculate success rate from display_results for consistency
            display_success_rate = ((working + redirect) / total * 100) if total > 0 else 0
            html += f''' |
            <strong>Total Time:</strong> {summary.total_time_seconds:.1f}s |
            <strong>Success Rate:</strong> {display_success_rate:.1f}%
'''
        html += '''
        </div>
'''

    html += '''
    </div>
</body>
</html>
'''

    return html


# =============================================================================
# HIGHLIGHTED DOCUMENT EXPORTS (v3.0.110)
# =============================================================================

def _get_broken_urls(results: List[ValidationResult]) -> Set[str]:
    """
    Get set of URLs that are broken/failed validation.

    Includes: BROKEN, TIMEOUT, DNSFAILED, SSLERROR, INVALID, BLOCKED statuses.
    Excludes: WORKING, REDIRECT, AUTH_REQUIRED (these are not broken), excluded URLs treated as valid.
    v5.9.29: Removed AUTH_REQUIRED — link exists but needs credentials, not broken.
    """
    broken_statuses = {'BROKEN', 'TIMEOUT', 'DNSFAILED', 'SSLERROR', 'INVALID', 'BLOCKED'}
    broken_urls = set()

    for result in results:
        # Skip excluded URLs that are treated as valid
        if result.excluded and getattr(result, 'treat_as_valid', True):
            continue

        if result.status.upper() in broken_statuses:
            broken_urls.add(result.url)
            # Also add normalized versions (without trailing slash, etc.)
            normalized = result.url.rstrip('/')
            broken_urls.add(normalized)
            if result.url.startswith('http://'):
                broken_urls.add(result.url.replace('http://', 'https://'))
            elif result.url.startswith('https://'):
                broken_urls.add(result.url.replace('https://', 'http://'))

    return broken_urls


def _get_result_for_url(url: str, results: List[ValidationResult]) -> Optional[ValidationResult]:
    """Get the validation result for a specific URL."""
    for result in results:
        if result.url == url or result.url.rstrip('/') == url.rstrip('/'):
            return result
    return None


def export_highlighted_docx(
    source_path: str,
    results: List[ValidationResult],
    output_path: Optional[str] = None
) -> Tuple[bool, str, bytes]:
    """
    Create a copy of a DOCX file with broken hyperlinks highlighted.

    Broken links are marked with:
    - Red text color
    - Yellow highlight background
    - Strikethrough formatting
    - A comment indicating the error

    Args:
        source_path: Path to the original DOCX file
        results: List of validation results
        output_path: Optional path for output file (if None, returns bytes)

    Returns:
        Tuple of (success: bool, message: str, file_bytes: bytes)
    """
    if not DOCX_AVAILABLE:
        return False, "python-docx library not installed. Cannot create highlighted DOCX.", b''

    if not os.path.exists(source_path):
        return False, f"Source file not found: {source_path}", b''

    broken_urls = _get_broken_urls(results)

    if not broken_urls:
        return False, "No broken links found to highlight.", b''

    try:
        # Load the document
        doc = Document(source_path)
        highlighted_count = 0

        # Process all paragraphs
        for para in doc.paragraphs:
            highlighted_count += _highlight_broken_links_in_paragraph(para, broken_urls, results)

        # Process tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        highlighted_count += _highlight_broken_links_in_paragraph(para, broken_urls, results)

        # Process headers and footers
        for section in doc.sections:
            for header in [section.header, section.first_page_header, section.even_page_header]:
                if header:
                    for para in header.paragraphs:
                        highlighted_count += _highlight_broken_links_in_paragraph(para, broken_urls, results)
            for footer in [section.footer, section.first_page_footer, section.even_page_footer]:
                if footer:
                    for para in footer.paragraphs:
                        highlighted_count += _highlight_broken_links_in_paragraph(para, broken_urls, results)

        # Save to bytes buffer
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        file_bytes = buffer.getvalue()

        # Optionally save to file
        if output_path:
            with open(output_path, 'wb') as f:
                f.write(file_bytes)

        return True, f"Highlighted {highlighted_count} broken link(s) in document.", file_bytes

    except Exception as e:
        return False, f"Error processing DOCX: {str(e)}", b''


def _highlight_broken_links_in_paragraph(
    para,
    broken_urls: Set[str],
    results: List[ValidationResult]
) -> int:
    """
    Highlight broken hyperlinks in a paragraph.

    Returns the number of links highlighted.
    """
    highlighted = 0

    # Access the underlying XML to find hyperlinks
    try:
        # Get hyperlink elements from paragraph XML
        para_xml = para._element

        # Find all hyperlink elements
        hyperlinks = para_xml.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}hyperlink')

        for hyperlink in hyperlinks:
            # Get the relationship ID
            r_id = hyperlink.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id')

            if r_id:
                # Try to get the URL from relationships
                try:
                    # Access document part to get relationships
                    part = para.part
                    if hasattr(part, 'rels') and r_id in part.rels:
                        rel = part.rels[r_id]
                        url = rel.target_ref if hasattr(rel, 'target_ref') else str(rel._target)

                        # Check if this URL is broken
                        if url in broken_urls or url.rstrip('/') in broken_urls:
                            # Highlight all runs within this hyperlink
                            runs = hyperlink.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}r')
                            for run in runs:
                                _apply_broken_link_formatting(run)

                            highlighted += 1
                except Exception:
                    pass

        # Also check for URLs in plain text (not hyperlinked)
        for run in para.runs:
            if run.text:
                for url in broken_urls:
                    if url in run.text:
                        _apply_broken_link_formatting_to_run(run)
                        highlighted += 1
                        break

    except Exception:
        pass

    return highlighted


def _apply_broken_link_formatting(run_element):
    """Apply red/strikethrough formatting to a run XML element."""
    try:
        # Get or create run properties
        rPr = run_element.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rPr')
        if rPr is None:
            rPr = OxmlElement('w:rPr')
            run_element.insert(0, rPr)

        # Add red color
        color = OxmlElement('w:color')
        color.set(qn('w:val'), 'FF0000')  # Red
        rPr.append(color)

        # Add yellow highlight
        highlight = OxmlElement('w:highlight')
        highlight.set(qn('w:val'), 'yellow')
        rPr.append(highlight)

        # Add strikethrough
        strike = OxmlElement('w:strike')
        strike.set(qn('w:val'), 'true')
        rPr.append(strike)

    except Exception:
        pass


def _apply_broken_link_formatting_to_run(run):
    """Apply broken link formatting to a python-docx Run object."""
    try:
        run.font.color.rgb = RGBColor(255, 0, 0)  # Red
        run.font.highlight_color = WD_COLOR_INDEX.YELLOW
        run.font.strike = True
    except Exception:
        pass


def export_highlighted_excel(
    source_path: str,
    results: List[ValidationResult],
    output_path: Optional[str] = None,
    link_column: Optional[int] = None
) -> Tuple[bool, str, bytes]:
    """
    Create a copy of an Excel file with rows containing broken links highlighted.

    Broken link rows are marked with:
    - Red background fill on the entire row
    - Bold red text on the URL cell
    - A comment on the URL cell with the error details

    Args:
        source_path: Path to the original Excel file
        results: List of validation results
        output_path: Optional path for output file (if None, returns bytes)
        link_column: Optional column index (1-based) containing URLs.
                     If None, will auto-detect.

    Returns:
        Tuple of (success: bool, message: str, file_bytes: bytes)
    """
    if not OPENPYXL_AVAILABLE:
        return False, "openpyxl library not installed. Cannot create highlighted Excel.", b''

    if not os.path.exists(source_path):
        return False, f"Source file not found: {source_path}", b''

    broken_urls = _get_broken_urls(results)

    if not broken_urls:
        return False, "No broken links found to highlight.", b''

    try:
        # Load workbook
        wb = load_workbook(source_path)
        highlighted_count = 0

        # Define styles for broken links
        red_fill = PatternFill(start_color='FFCCCC', end_color='FFCCCC', fill_type='solid')
        red_font = Font(color='CC0000', bold=True)
        error_fill = PatternFill(start_color='FF6666', end_color='FF6666', fill_type='solid')

        # Process each sheet
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]

            # Find URL columns if not specified
            url_columns = _find_url_columns(ws) if link_column is None else [link_column]

            if not url_columns:
                continue

            # Process each row
            for row_idx, row in enumerate(ws.iter_rows(min_row=2), start=2):  # Skip header
                row_has_broken_link = False
                broken_cells = []

                for col_idx in url_columns:
                    if col_idx <= len(row):
                        cell = row[col_idx - 1]
                        cell_value = str(cell.value) if cell.value else ''

                        # Check if cell contains a broken URL
                        for url in broken_urls:
                            if url in cell_value or cell_value in broken_urls:
                                row_has_broken_link = True
                                broken_cells.append((cell, url))
                                break

                        # Also check hyperlink target
                        if cell.hyperlink and cell.hyperlink.target:
                            target = cell.hyperlink.target
                            if target in broken_urls or target.rstrip('/') in broken_urls:
                                row_has_broken_link = True
                                broken_cells.append((cell, target))

                # Highlight the entire row if it has a broken link
                if row_has_broken_link:
                    highlighted_count += 1

                    # Apply red fill to entire row
                    for cell in row:
                        cell.fill = red_fill

                    # Apply stronger formatting to the URL cells
                    for cell, url in broken_cells:
                        cell.fill = error_fill
                        cell.font = red_font

                        # Add comment with error details
                        result = _get_result_for_url(url, results)
                        if result:
                            comment_text = f"BROKEN LINK\nStatus: {result.status}\nMessage: {result.message}"
                            if result.status_code:
                                comment_text += f"\nHTTP Code: {result.status_code}"
                            cell.comment = Comment(comment_text, "Hyperlink Validator")

        # Save to bytes buffer
        buffer = io.BytesIO()
        wb.save(buffer)
        buffer.seek(0)
        file_bytes = buffer.getvalue()

        # Optionally save to file
        if output_path:
            with open(output_path, 'wb') as f:
                f.write(file_bytes)

        wb.close()

        return True, f"Highlighted {highlighted_count} row(s) with broken links.", file_bytes

    except Exception as e:
        return False, f"Error processing Excel: {str(e)}", b''


def _find_url_columns(ws) -> List[int]:
    """
    Auto-detect columns that likely contain URLs.

    Looks for:
    - Header row with 'URL', 'Link', 'Hyperlink', 'Website', 'Web' in name
    - Columns with cells containing http://, https://, www., or mailto:

    Returns list of 1-based column indices.
    """
    url_columns = []
    url_patterns = ['url', 'link', 'hyperlink', 'website', 'web', 'href']
    url_regex = re.compile(r'https?://|www\.|mailto:', re.IGNORECASE)

    # Check header row
    header_row = list(ws.iter_rows(min_row=1, max_row=1))[0] if ws.max_row > 0 else []

    for col_idx, cell in enumerate(header_row, start=1):
        if cell.value:
            header_text = str(cell.value).lower()
            if any(pattern in header_text for pattern in url_patterns):
                url_columns.append(col_idx)

    # If no headers matched, scan first few rows for URL-like content
    if not url_columns:
        url_col_candidates = {}

        for row in ws.iter_rows(min_row=1, max_row=min(20, ws.max_row)):
            for col_idx, cell in enumerate(row, start=1):
                if cell.value:
                    cell_text = str(cell.value)
                    if url_regex.search(cell_text) or (cell.hyperlink and cell.hyperlink.target):
                        url_col_candidates[col_idx] = url_col_candidates.get(col_idx, 0) + 1

        # Select columns with at least 2 URL-like values
        url_columns = [col for col, count in url_col_candidates.items() if count >= 2]

    # If still nothing, check all columns with hyperlinks
    if not url_columns:
        for col_idx in range(1, ws.max_column + 1):
            for row in ws.iter_rows(min_row=2, max_row=min(100, ws.max_row)):
                cell = row[col_idx - 1] if col_idx <= len(row) else None
                if cell and cell.hyperlink:
                    url_columns.append(col_idx)
                    break

    return url_columns


# =============================================================================
# MULTI-COLOR STATUS HIGHLIGHTING (v5.9.33)
# =============================================================================

# Status → color mapping for row highlighting
_STATUS_COLORS = {
    # Green shades — link verified working
    'WORKING':       {'fill': 'C6EFCE', 'font': '006100', 'label': 'Working'},
    'OK':            {'fill': 'C6EFCE', 'font': '006100', 'label': 'Working'},

    # Light green — redirect (link works but redirects)
    'REDIRECT':      {'fill': 'D5F5E3', 'font': '1E8449', 'label': 'Redirect'},

    # Yellow — warnings (link works but has issues)
    'SSL_WARNING':   {'fill': 'FFF2CC', 'font': '7D6608', 'label': 'SSL Warning'},
    'REDIRECT_LOOP': {'fill': 'FFF2CC', 'font': '7D6608', 'label': 'Redirect Loop'},
    'REDIRECT_ERROR':{'fill': 'FFF2CC', 'font': '7D6608', 'label': 'Redirect Error'},

    # Orange — auth/access issues (link exists but can't verify)
    'AUTH_REQUIRED': {'fill': 'FCE4D6', 'font': '974706', 'label': 'Auth Required'},
    'BLOCKED':       {'fill': 'FCE4D6', 'font': '974706', 'label': 'Blocked'},

    # Red — broken/failed
    'BROKEN':        {'fill': 'FFC7CE', 'font': 'C00000', 'label': 'Broken'},
    'INVALID':       {'fill': 'FFC7CE', 'font': 'C00000', 'label': 'Invalid URL'},
    'TIMEOUT':       {'fill': 'FFC7CE', 'font': 'C00000', 'label': 'Timeout'},
    'DNSFAILED':     {'fill': 'FFC7CE', 'font': 'C00000', 'label': 'DNS Failed'},
    'SSLERROR':      {'fill': 'FFC7CE', 'font': 'C00000', 'label': 'SSL Error'},

    # Grey — no URL / not tested
    'NO_URL':        {'fill': 'F2F2F2', 'font': '808080', 'label': 'No URL'},
    'EXCLUDED':      {'fill': 'E2EFDA', 'font': '548235', 'label': 'Excluded (OK)'},
}

# Status categories for summary counts
_STATUS_CATEGORY = {
    'WORKING': 'good', 'OK': 'good', 'REDIRECT': 'good', 'EXCLUDED': 'good',
    'SSL_WARNING': 'warning', 'REDIRECT_LOOP': 'warning', 'REDIRECT_ERROR': 'warning',
    'AUTH_REQUIRED': 'caution', 'BLOCKED': 'caution',
    'BROKEN': 'broken', 'INVALID': 'broken', 'TIMEOUT': 'broken',
    'DNSFAILED': 'broken', 'SSLERROR': 'broken',
    'NO_URL': 'no_url',
}


def _build_url_status_map(results: List[ValidationResult]) -> Dict[str, ValidationResult]:
    """
    Build a map of URL → ValidationResult for quick lookup.
    Includes normalized variants (no trailing slash, http/https swap).
    """
    url_map = {}
    for result in results:
        # Apply exclusion display rules
        display_result = _apply_exclusion_display(result)
        url = display_result.url
        url_map[url] = display_result
        url_map[url.rstrip('/')] = display_result
        # Add protocol-swapped variant
        if url.startswith('http://'):
            url_map[url.replace('http://', 'https://')] = display_result
        elif url.startswith('https://'):
            url_map[url.replace('https://', 'http://')] = display_result
    return url_map


def export_highlighted_excel_multicolor(
    source_path: str,
    results: List[ValidationResult],
    output_path: Optional[str] = None,
    link_column: Optional[int] = None
) -> Tuple[bool, str, bytes]:
    """
    Create a copy of an Excel file with ALL rows color-coded by link validation status.

    Color scheme:
    - Green:  WORKING, OK, REDIRECT — link verified working
    - Yellow: SSL_WARNING, REDIRECT_LOOP — link works but has issues
    - Orange: AUTH_REQUIRED, BLOCKED — link exists but can't fully verify
    - Red:    BROKEN, INVALID, TIMEOUT, DNSFAILED, SSLERROR — link is broken
    - Grey:   No URL in row — neutral, not tested

    Every row with a URL gets colored. Rows without URLs get light grey.
    A status column is added at the end showing the validation status.
    A summary sheet is added with counts by status category.

    Args:
        source_path: Path to the original Excel file
        results: List of ALL validation results (working + broken + everything)
        output_path: Optional path for output file (if None, returns bytes)
        link_column: Optional column index (1-based) containing URLs.
                     If None, will auto-detect.

    Returns:
        Tuple of (success: bool, message: str, file_bytes: bytes)
    """
    if not OPENPYXL_AVAILABLE:
        return False, "openpyxl library not installed. Cannot create highlighted Excel.", b''

    if not os.path.exists(source_path):
        return False, f"Source file not found: {source_path}", b''

    if not results:
        return False, "No validation results provided.", b''

    try:
        # Load workbook
        wb = load_workbook(source_path)

        # Build URL → result lookup map
        url_status_map = _build_url_status_map(results)

        # Pre-create PatternFill and Font objects for each status
        status_fills = {}
        status_fonts = {}
        for status_key, colors in _STATUS_COLORS.items():
            status_fills[status_key] = PatternFill(
                start_color=colors['fill'], end_color=colors['fill'], fill_type='solid'
            )
            status_fonts[status_key] = Font(color=colors['font'])

        # Bold font variants for URL cells
        status_fonts_bold = {}
        for status_key, colors in _STATUS_COLORS.items():
            status_fonts_bold[status_key] = Font(color=colors['font'], bold=True)

        # Header style
        header_fill = PatternFill(start_color='4472C4', end_color='4472C4', fill_type='solid')
        header_font = Font(color='FFFFFF', bold=True)

        # Tracking counts
        total_highlighted = 0
        category_counts = {'good': 0, 'warning': 0, 'caution': 0, 'broken': 0, 'no_url': 0}

        # Process each sheet
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]

            if ws.max_row is None or ws.max_row < 2:
                continue

            # Find URL columns if not specified
            url_columns = _find_url_columns(ws) if link_column is None else [link_column]

            # Add "Link Status" column header
            status_col = ws.max_column + 1
            status_header_cell = ws.cell(row=1, column=status_col)
            status_header_cell.value = 'Link Status'
            status_header_cell.fill = header_fill
            status_header_cell.font = header_font
            status_header_cell.alignment = Alignment(horizontal='center')

            # Also add "Link Details" column
            details_col = status_col + 1
            details_header_cell = ws.cell(row=1, column=details_col)
            details_header_cell.value = 'Link Details'
            details_header_cell.fill = header_fill
            details_header_cell.font = header_font

            # Set column widths
            ws.column_dimensions[get_column_letter(status_col)].width = 16
            ws.column_dimensions[get_column_letter(details_col)].width = 50

            # Process each data row
            for row_idx in range(2, ws.max_row + 1):
                row_url = None
                row_result = None

                # Find URL in this row
                for col_idx in url_columns:
                    cell = ws.cell(row=row_idx, column=col_idx)
                    cell_value = str(cell.value).strip() if cell.value else ''

                    if cell_value and ('http://' in cell_value.lower() or 'https://' in cell_value.lower()):
                        row_url = cell_value
                        # Look up in our status map
                        row_result = url_status_map.get(cell_value) or \
                                     url_status_map.get(cell_value.rstrip('/'))
                        if row_result:
                            break

                    # Also check hyperlink target
                    if cell.hyperlink and cell.hyperlink.target:
                        target = cell.hyperlink.target
                        row_result = url_status_map.get(target) or \
                                     url_status_map.get(target.rstrip('/'))
                        if row_result:
                            row_url = target
                            break

                # Determine status for coloring
                if row_result:
                    status_key = row_result.status.upper()
                    # Handle excluded URLs
                    if row_result.excluded and getattr(row_result, 'treat_as_valid', True):
                        status_key = 'EXCLUDED'
                elif row_url:
                    # URL exists but wasn't in our results (wasn't validated)
                    status_key = 'NO_URL'
                else:
                    # No URL in this row
                    status_key = 'NO_URL'

                # Get colors (fall back to NO_URL grey if unknown status)
                if status_key not in _STATUS_COLORS:
                    status_key = 'BROKEN' if row_result else 'NO_URL'

                fill = status_fills.get(status_key, status_fills['NO_URL'])
                font = status_fonts.get(status_key, status_fonts['NO_URL'])
                font_bold = status_fonts_bold.get(status_key, status_fonts_bold['NO_URL'])
                color_info = _STATUS_COLORS.get(status_key, _STATUS_COLORS['NO_URL'])

                # Apply fill to entire row
                for col_idx in range(1, status_col):  # Don't include new status columns yet
                    cell = ws.cell(row=row_idx, column=col_idx)
                    cell.fill = fill

                # Apply bold font to the URL cell specifically
                for col_idx in url_columns:
                    cell = ws.cell(row=row_idx, column=col_idx)
                    if cell.value:
                        cell.font = font_bold

                # Write status column
                status_cell = ws.cell(row=row_idx, column=status_col)
                status_cell.value = color_info['label']
                status_cell.fill = fill
                status_cell.font = font_bold
                status_cell.alignment = Alignment(horizontal='center')

                # Write details column
                details_cell = ws.cell(row=row_idx, column=details_col)
                if row_result:
                    detail_parts = []
                    if row_result.status_code:
                        detail_parts.append(f'HTTP {row_result.status_code}')
                    if row_result.message:
                        detail_parts.append(row_result.message[:200])
                    details_cell.value = ' — '.join(detail_parts) if detail_parts else ''
                elif row_url:
                    details_cell.value = 'URL present but not validated'
                else:
                    details_cell.value = ''
                details_cell.fill = fill
                details_cell.font = font

                # Track counts
                cat = _STATUS_CATEGORY.get(status_key, 'no_url')
                category_counts[cat] = category_counts.get(cat, 0) + 1
                if status_key != 'NO_URL':
                    total_highlighted += 1

        # Add summary sheet
        _add_summary_sheet(wb, results, category_counts, total_highlighted)

        # Save to bytes buffer
        buffer = io.BytesIO()
        wb.save(buffer)
        buffer.seek(0)
        file_bytes = buffer.getvalue()

        # Optionally save to file
        if output_path:
            with open(output_path, 'wb') as f:
                f.write(file_bytes)

        wb.close()

        good = category_counts.get('good', 0)
        warn = category_counts.get('warning', 0)
        caution = category_counts.get('caution', 0)
        broken = category_counts.get('broken', 0)

        message = (f"Highlighted {total_highlighted} rows: "
                   f"{good} working (green), {warn} warning (yellow), "
                   f"{caution} auth/blocked (orange), {broken} broken (red)")

        return True, message, file_bytes

    except Exception as e:
        import traceback
        traceback.print_exc()
        return False, f"Error processing Excel: {str(e)}", b''


def _add_summary_sheet(wb, results: List[ValidationResult],
                       category_counts: Dict[str, int],
                       total_highlighted: int):
    """Add a summary sheet to the workbook with validation statistics."""
    try:
        # Create or get summary sheet
        summary_name = 'Link Validation Summary'
        if summary_name in wb.sheetnames:
            del wb[summary_name]
        ws = wb.create_sheet(summary_name)

        # Styles
        title_font = Font(size=16, bold=True, color='1F4E79')
        subtitle_font = Font(size=12, bold=True, color='2E75B6')
        header_fill = PatternFill(start_color='4472C4', end_color='4472C4', fill_type='solid')
        header_font = Font(color='FFFFFF', bold=True, size=11)
        normal_font = Font(size=11)
        bold_font = Font(size=11, bold=True)

        # Title
        ws.cell(row=1, column=1).value = 'Hyperlink Validation Report'
        ws.cell(row=1, column=1).font = title_font
        ws.merge_cells('A1:D1')

        ws.cell(row=2, column=1).value = f'Generated: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}'
        ws.cell(row=2, column=1).font = Font(size=10, color='808080')

        # Summary counts section
        ws.cell(row=4, column=1).value = 'Status Summary'
        ws.cell(row=4, column=1).font = subtitle_font

        # Headers
        for col, header_text in enumerate(['Category', 'Count', 'Color'], start=1):
            cell = ws.cell(row=5, column=col)
            cell.value = header_text
            cell.fill = header_fill
            cell.font = header_font

        # Category rows
        summary_rows = [
            ('Working (verified)', category_counts.get('good', 0), 'C6EFCE', '006100'),
            ('Warning (SSL/redirect)', category_counts.get('warning', 0), 'FFF2CC', '7D6608'),
            ('Auth/Blocked', category_counts.get('caution', 0), 'FCE4D6', '974706'),
            ('Broken', category_counts.get('broken', 0), 'FFC7CE', 'C00000'),
            ('No URL', category_counts.get('no_url', 0), 'F2F2F2', '808080'),
        ]

        for i, (label, count, fill_color, font_color) in enumerate(summary_rows, start=6):
            ws.cell(row=i, column=1).value = label
            ws.cell(row=i, column=1).font = bold_font
            ws.cell(row=i, column=2).value = count
            ws.cell(row=i, column=2).font = normal_font
            ws.cell(row=i, column=2).alignment = Alignment(horizontal='center')
            # Color swatch
            ws.cell(row=i, column=3).fill = PatternFill(
                start_color=fill_color, end_color=fill_color, fill_type='solid'
            )

        total_row = 6 + len(summary_rows)
        ws.cell(row=total_row, column=1).value = 'Total URLs Validated'
        ws.cell(row=total_row, column=1).font = bold_font
        ws.cell(row=total_row, column=2).value = total_highlighted
        ws.cell(row=total_row, column=2).font = bold_font
        ws.cell(row=total_row, column=2).alignment = Alignment(horizontal='center')

        # Detailed status breakdown
        detail_start = total_row + 2
        ws.cell(row=detail_start, column=1).value = 'Detailed Status Breakdown'
        ws.cell(row=detail_start, column=1).font = subtitle_font

        for col, header_text in enumerate(['Status', 'Count', 'Color', 'Description'], start=1):
            cell = ws.cell(row=detail_start + 1, column=col)
            cell.value = header_text
            cell.fill = header_fill
            cell.font = header_font

        # Count each specific status
        status_detail_counts = {}
        for result in results:
            s = result.status.upper()
            status_detail_counts[s] = status_detail_counts.get(s, 0) + 1

        detail_row = detail_start + 2
        status_descriptions = {
            'WORKING': 'Link verified working (HTTP 200)',
            'REDIRECT': 'Link redirects to another URL but works',
            'SSL_WARNING': 'Link works but SSL certificate not trusted',
            'AUTH_REQUIRED': 'Link requires authentication credentials',
            'BLOCKED': 'Link blocked by bot protection or firewall',
            'BROKEN': 'Link returns an error (404, 500, etc.)',
            'INVALID': 'URL format is invalid',
            'TIMEOUT': 'Connection timed out before response',
            'DNSFAILED': 'Domain name could not be resolved',
            'SSLERROR': 'SSL/TLS connection error',
            'REDIRECT_LOOP': 'Too many redirects detected',
            'REDIRECT_ERROR': 'Redirect chain encountered an error',
        }

        for status_key in ['WORKING', 'REDIRECT', 'SSL_WARNING', 'AUTH_REQUIRED',
                           'BLOCKED', 'BROKEN', 'INVALID', 'TIMEOUT', 'DNSFAILED',
                           'SSLERROR', 'REDIRECT_LOOP', 'REDIRECT_ERROR']:
            count = status_detail_counts.get(status_key, 0)
            if count > 0:
                colors = _STATUS_COLORS.get(status_key, _STATUS_COLORS['NO_URL'])
                ws.cell(row=detail_row, column=1).value = colors['label']
                ws.cell(row=detail_row, column=1).font = bold_font
                ws.cell(row=detail_row, column=2).value = count
                ws.cell(row=detail_row, column=2).alignment = Alignment(horizontal='center')
                ws.cell(row=detail_row, column=3).fill = PatternFill(
                    start_color=colors['fill'], end_color=colors['fill'], fill_type='solid'
                )
                ws.cell(row=detail_row, column=4).value = status_descriptions.get(status_key, '')
                detail_row += 1

        # Column widths
        ws.column_dimensions['A'].width = 30
        ws.column_dimensions['B'].width = 12
        ws.column_dimensions['C'].width = 10
        ws.column_dimensions['D'].width = 50

        # Legend section
        legend_start = detail_row + 2
        ws.cell(row=legend_start, column=1).value = 'Color Legend'
        ws.cell(row=legend_start, column=1).font = subtitle_font

        legend_items = [
            ('Green', 'C6EFCE', 'Link verified working — safe'),
            ('Yellow', 'FFF2CC', 'Link works but has SSL or redirect warning'),
            ('Orange', 'FCE4D6', 'Link requires authentication or is blocked by firewall'),
            ('Red', 'FFC7CE', 'Link is broken, timed out, or unreachable'),
            ('Grey', 'F2F2F2', 'No URL in this row — not tested'),
        ]
        for i, (label, fill_color, desc) in enumerate(legend_items, start=legend_start + 1):
            ws.cell(row=i, column=1).value = label
            ws.cell(row=i, column=1).font = bold_font
            ws.cell(row=i, column=1).fill = PatternFill(
                start_color=fill_color, end_color=fill_color, fill_type='solid'
            )
            ws.cell(row=i, column=2).value = desc
            ws.merge_cells(start_row=i, start_column=2, end_row=i, end_column=4)

    except Exception:
        pass  # Summary sheet is optional — don't fail the export


def is_highlighted_export_available() -> Dict[str, bool]:
    """Check which highlighted export formats are available."""
    return {
        'docx': DOCX_AVAILABLE,
        'excel': OPENPYXL_AVAILABLE
    }
