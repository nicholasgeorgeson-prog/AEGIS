"""
AEGIS - Document Structure Analyzer
Version: 1.0.0
Created: February 3, 2026

Provides deep document structure analysis for DOCX files.
Extracts headings, table of contents, cross-references, and document hierarchy.

Features:
- Heading hierarchy extraction and validation
- Table of Contents analysis and verification
- Cross-reference detection (figures, tables, sections)
- Document outline generation
- Structure consistency checking
- Numbering scheme validation
- Section depth analysis
- Air-gap compatible (no external API calls)

Usage:
    from structure_analyzer import StructureAnalyzer

    analyzer = StructureAnalyzer()
    structure = analyzer.analyze_docx('/path/to/document.docx')

    # Get heading tree
    headings = analyzer.get_heading_tree()

    # Validate structure
    issues = analyzer.validate_structure()
"""

import re
import os
from typing import Dict, List, Any, Optional, Tuple, Set
from collections import defaultdict
from dataclasses import dataclass, field, asdict
from enum import Enum
import xml.etree.ElementTree as ET
from zipfile import ZipFile
import json

# Optional python-docx for enhanced extraction
try:
    from docx import Document
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


class StructureIssueType(Enum):
    """Types of structure issues."""
    MISSING_HEADING = "missing_heading"
    SKIPPED_LEVEL = "skipped_level"
    INCONSISTENT_NUMBERING = "inconsistent_numbering"
    ORPHAN_SECTION = "orphan_section"
    DUPLICATE_HEADING = "duplicate_heading"
    TOC_MISMATCH = "toc_mismatch"
    BROKEN_REFERENCE = "broken_reference"
    MISSING_FIGURE_CAPTION = "missing_figure_caption"
    MISSING_TABLE_CAPTION = "missing_table_caption"
    INCONSISTENT_STYLE = "inconsistent_style"


@dataclass
class Heading:
    """Represents a document heading."""
    text: str
    level: int
    number: Optional[str] = None
    page: Optional[int] = None
    paragraph_index: int = 0
    style_name: str = ""
    children: List['Heading'] = field(default_factory=list)

    def to_dict(self) -> Dict[str, Any]:
        return {
            'text': self.text,
            'level': self.level,
            'number': self.number,
            'page': self.page,
            'paragraph_index': self.paragraph_index,
            'style_name': self.style_name,
            'children': [c.to_dict() for c in self.children]
        }


@dataclass
class CrossReference:
    """Represents a cross-reference in the document."""
    ref_type: str  # 'figure', 'table', 'section', 'equation', 'bookmark'
    ref_text: str
    target: Optional[str] = None
    paragraph_index: int = 0
    is_valid: bool = True
    context: str = ""

    def to_dict(self) -> Dict[str, Any]:
        return asdict(self)


@dataclass
class StructureIssue:
    """Represents a document structure issue."""
    issue_type: StructureIssueType
    message: str
    location: str = ""
    severity: str = "warning"  # error, warning, info
    suggestion: Optional[str] = None

    def to_dict(self) -> Dict[str, Any]:
        return {
            'issue_type': self.issue_type.value,
            'message': self.message,
            'location': self.location,
            'severity': self.severity,
            'suggestion': self.suggestion
        }


class StructureAnalyzer:
    """
    Analyzes document structure including headings, TOC, and cross-references.
    """

    VERSION = '1.0.0'

    # Standard heading style names in Word
    HEADING_STYLES = [
        'Heading 1', 'Heading 2', 'Heading 3', 'Heading 4',
        'Heading 5', 'Heading 6', 'Heading 7', 'Heading 8', 'Heading 9',
        'Title', 'Subtitle', 'TOC Heading'
    ]

    # Alternative heading style patterns
    HEADING_PATTERNS = [
        r'^Heading\s*(\d+)$',
        r'^H(\d+)$',
        r'^Level\s*(\d+)$',
        r'^Section\s*(\d+)$',
    ]

    # Numbering patterns
    NUMBERING_PATTERNS = {
        'decimal': r'^(\d+\.)+\d*\s*',           # 1.2.3
        'alpha_lower': r'^[a-z]+\.\s*',           # a. b. c.
        'alpha_upper': r'^[A-Z]+\.\s*',           # A. B. C.
        'roman_lower': r'^[ivxlcdm]+\.\s*',       # i. ii. iii.
        'roman_upper': r'^[IVXLCDM]+\.\s*',       # I. II. III.
        'parenthetical': r'^\([a-z0-9]+\)\s*',    # (a) (1)
        'bullet': r'^[•●○◦▪▫-]\s*',              # Bullets
    }

    # Cross-reference patterns
    XREF_PATTERNS = {
        'figure': [
            r'(?:Figure|Fig\.?|Exhibit)\s+(\d+(?:\.\d+)*(?:-\d+)?)',
            r'(?:see|refer to|shown in)\s+(?:Figure|Fig\.?)\s+(\d+(?:\.\d+)*)',
        ],
        'table': [
            r'(?:Table|Tbl\.?)\s+(\d+(?:\.\d+)*(?:-\d+)?)',
            r'(?:see|refer to|shown in)\s+(?:Table|Tbl\.?)\s+(\d+(?:\.\d+)*)',
        ],
        'section': [
            r'(?:Section|Sect\.?|§)\s+(\d+(?:\.\d+)*)',
            r'(?:see|refer to)\s+(?:Section|Sect\.?)\s+(\d+(?:\.\d+)*)',
            r'(?:paragraph|para\.?)\s+(\d+(?:\.\d+)*)',
        ],
        'appendix': [
            r'(?:Appendix|App\.?|Annex)\s+([A-Z](?:\.\d+)*)',
        ],
        'equation': [
            r'(?:Equation|Eq\.?)\s+\(?(\d+(?:\.\d+)*)\)?',
        ],
    }

    # Caption patterns
    CAPTION_PATTERNS = {
        'figure': r'^(?:Figure|Fig\.?|Exhibit)\s+(\d+(?:\.\d+)*(?:-\d+)?)[:\.\s]+(.+)',
        'table': r'^(?:Table|Tbl\.?)\s+(\d+(?:\.\d+)*(?:-\d+)?)[:\.\s]+(.+)',
    }

    def __init__(self):
        """Initialize the structure analyzer."""
        self.headings: List[Heading] = []
        self.cross_references: List[CrossReference] = []
        self.figures: List[Dict[str, Any]] = []
        self.tables: List[Dict[str, Any]] = []
        self.toc_entries: List[Dict[str, Any]] = []
        self.bookmarks: Set[str] = set()
        self.document = None
        self.raw_xml = None
        self.issues: List[StructureIssue] = []

    def analyze_docx(self, filepath: str) -> Dict[str, Any]:
        """
        Analyze a DOCX file structure.

        Args:
            filepath: Path to the DOCX file

        Returns:
            Dictionary with structure analysis results
        """
        if not os.path.exists(filepath):
            return {'error': f'File not found: {filepath}'}

        if not filepath.lower().endswith('.docx'):
            return {'error': 'File must be a .docx file'}

        # Reset state
        self._reset()

        # Extract using python-docx if available
        if DOCX_AVAILABLE:
            self._analyze_with_python_docx(filepath)
        else:
            self._analyze_raw_xml(filepath)

        # Validate structure
        self._validate_structure()

        return self._build_results()

    def _reset(self):
        """Reset analyzer state."""
        self.headings = []
        self.cross_references = []
        self.figures = []
        self.tables = []
        self.toc_entries = []
        self.bookmarks = set()
        self.document = None
        self.raw_xml = None
        self.issues = []

    def _analyze_with_python_docx(self, filepath: str):
        """Analyze document using python-docx library."""
        self.document = Document(filepath)

        # Extract headings
        self._extract_headings_docx()

        # Extract cross-references from text
        self._extract_cross_references_docx()

        # Extract figures and tables
        self._extract_figures_tables_docx()

        # Extract bookmarks
        self._extract_bookmarks_docx()

        # Try to extract TOC if present
        self._extract_toc_docx()

    def _extract_headings_docx(self):
        """Extract headings using python-docx."""
        for i, para in enumerate(self.document.paragraphs):
            style_name = para.style.name if para.style else ''

            # Check if it's a heading style
            level = self._get_heading_level(style_name)

            if level > 0:
                # Extract any numbering
                number = self._extract_number(para.text)
                text = self._clean_heading_text(para.text)

                if text.strip():
                    heading = Heading(
                        text=text,
                        level=level,
                        number=number,
                        paragraph_index=i,
                        style_name=style_name
                    )
                    self.headings.append(heading)

    def _get_heading_level(self, style_name: str) -> int:
        """Get heading level from style name."""
        if not style_name:
            return 0

        # Direct match
        if style_name in self.HEADING_STYLES:
            # Extract number from "Heading X"
            match = re.search(r'(\d+)', style_name)
            if match:
                return int(match.group(1))
            elif style_name == 'Title':
                return 0  # Title is special
            elif style_name == 'Subtitle':
                return 1

        # Pattern match
        for pattern in self.HEADING_PATTERNS:
            match = re.match(pattern, style_name)
            if match:
                return int(match.group(1))

        return 0

    def _extract_number(self, text: str) -> Optional[str]:
        """Extract section number from heading text."""
        for pattern_name, pattern in self.NUMBERING_PATTERNS.items():
            match = re.match(pattern, text.strip())
            if match:
                return match.group(0).strip()
        return None

    def _clean_heading_text(self, text: str) -> str:
        """Remove numbering from heading text."""
        cleaned = text.strip()
        for pattern in self.NUMBERING_PATTERNS.values():
            cleaned = re.sub('^' + pattern, '', cleaned).strip()
        return cleaned

    def _extract_cross_references_docx(self):
        """Extract cross-references from document text."""
        for i, para in enumerate(self.document.paragraphs):
            text = para.text

            for ref_type, patterns in self.XREF_PATTERNS.items():
                for pattern in patterns:
                    for match in re.finditer(pattern, text, re.IGNORECASE):
                        ref = CrossReference(
                            ref_type=ref_type,
                            ref_text=match.group(0),
                            target=match.group(1) if match.groups() else None,
                            paragraph_index=i,
                            context=text[:100]
                        )
                        self.cross_references.append(ref)

    def _extract_figures_tables_docx(self):
        """Extract figure and table captions."""
        for i, para in enumerate(self.document.paragraphs):
            text = para.text.strip()

            # Check for figure caption
            fig_match = re.match(self.CAPTION_PATTERNS['figure'], text, re.IGNORECASE)
            if fig_match:
                self.figures.append({
                    'number': fig_match.group(1),
                    'caption': fig_match.group(2).strip(),
                    'paragraph_index': i,
                    'full_text': text
                })

            # Check for table caption
            tbl_match = re.match(self.CAPTION_PATTERNS['table'], text, re.IGNORECASE)
            if tbl_match:
                self.tables.append({
                    'number': tbl_match.group(1),
                    'caption': tbl_match.group(2).strip(),
                    'paragraph_index': i,
                    'full_text': text
                })

    def _extract_bookmarks_docx(self):
        """Extract bookmarks from document XML."""
        try:
            # Access the document's XML to find bookmarks
            for para in self.document.paragraphs:
                xml_str = para._element.xml
                # Look for bookmark start elements
                for match in re.finditer(r'w:bookmarkStart[^>]+w:name="([^"]+)"', xml_str):
                    self.bookmarks.add(match.group(1))
        except Exception:
            pass  # Bookmarks extraction is optional

    def _extract_toc_docx(self):
        """Extract Table of Contents entries."""
        in_toc = False
        toc_level_pattern = r'^(?:TOC\s*)?(\d+)'

        for para in self.document.paragraphs:
            style_name = para.style.name if para.style else ''

            # Check if we're entering/leaving TOC
            if 'TOC' in style_name.upper() or 'CONTENTS' in para.text.upper():
                in_toc = True

            if in_toc and style_name.startswith('TOC'):
                # Extract TOC level
                level_match = re.search(toc_level_pattern, style_name)
                level = int(level_match.group(1)) if level_match else 1

                # Parse TOC entry (text and page number)
                text = para.text.strip()
                # TOC entries often end with page number after dots
                page_match = re.search(r'\.{2,}\s*(\d+)\s*$', text)

                entry = {
                    'text': re.sub(r'\.{2,}\s*\d+\s*$', '', text).strip(),
                    'level': level,
                    'page': int(page_match.group(1)) if page_match else None
                }
                self.toc_entries.append(entry)

    def _analyze_raw_xml(self, filepath: str):
        """Fallback: Analyze document by parsing raw XML."""
        try:
            with ZipFile(filepath, 'r') as zf:
                # Read document.xml
                with zf.open('word/document.xml') as doc_xml:
                    self.raw_xml = doc_xml.read()

                self._parse_raw_xml()
        except Exception as e:
            self.issues.append(StructureIssue(
                issue_type=StructureIssueType.MISSING_HEADING,
                message=f'Error reading document: {str(e)}',
                severity='error'
            ))

    def _parse_raw_xml(self):
        """Parse the raw XML to extract structure."""
        if not self.raw_xml:
            return

        # Define namespace
        ns = {
            'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
        }

        root = ET.fromstring(self.raw_xml)

        # Find all paragraphs
        para_index = 0
        for para in root.findall('.//w:p', ns):
            # Get paragraph style
            style_elem = para.find('.//w:pStyle', ns)
            style_name = style_elem.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val') if style_elem is not None else ''

            # Get text content
            text_parts = []
            for t in para.findall('.//w:t', ns):
                if t.text:
                    text_parts.append(t.text)
            text = ''.join(text_parts)

            # Check if heading
            level = self._get_heading_level(style_name)
            if level > 0 and text.strip():
                heading = Heading(
                    text=self._clean_heading_text(text),
                    level=level,
                    number=self._extract_number(text),
                    paragraph_index=para_index,
                    style_name=style_name
                )
                self.headings.append(heading)

            para_index += 1

    def _validate_structure(self):
        """Validate document structure and find issues."""
        self._check_heading_hierarchy()
        self._check_numbering_consistency()
        self._check_duplicate_headings()
        self._check_cross_references()
        self._check_toc_matches()
        self._check_figure_table_references()

    def _check_heading_hierarchy(self):
        """Check for skipped heading levels."""
        prev_level = 0

        for heading in self.headings:
            if heading.level > prev_level + 1:
                # Skipped a level
                self.issues.append(StructureIssue(
                    issue_type=StructureIssueType.SKIPPED_LEVEL,
                    message=f'Heading level skipped from {prev_level} to {heading.level}',
                    location=f'"{heading.text}"',
                    suggestion=f'Add a level {prev_level + 1} heading before this, or change this to level {prev_level + 1}'
                ))

            prev_level = heading.level

    def _check_numbering_consistency(self):
        """Check for inconsistent numbering patterns."""
        numbered_headings = [h for h in self.headings if h.number]

        if not numbered_headings:
            return

        # Group by level
        by_level = defaultdict(list)
        for h in numbered_headings:
            by_level[h.level].append(h)

        # Check each level for consistency
        for level, headings in by_level.items():
            patterns_used = set()
            for h in headings:
                for pattern_name, pattern in self.NUMBERING_PATTERNS.items():
                    if re.match(pattern, h.number):
                        patterns_used.add(pattern_name)
                        break

            if len(patterns_used) > 1:
                self.issues.append(StructureIssue(
                    issue_type=StructureIssueType.INCONSISTENT_NUMBERING,
                    message=f'Inconsistent numbering at level {level}',
                    location=f'Patterns found: {", ".join(patterns_used)}',
                    severity='warning',
                    suggestion='Use consistent numbering throughout each heading level'
                ))

    def _check_duplicate_headings(self):
        """Check for duplicate heading text."""
        seen = {}
        for heading in self.headings:
            key = (heading.level, heading.text.lower().strip())
            if key in seen:
                self.issues.append(StructureIssue(
                    issue_type=StructureIssueType.DUPLICATE_HEADING,
                    message=f'Duplicate heading at level {heading.level}',
                    location=f'"{heading.text}" appears multiple times',
                    severity='info',
                    suggestion='Consider making heading text unique for clarity'
                ))
            else:
                seen[key] = heading

    def _check_cross_references(self):
        """Validate cross-references have targets."""
        # Build set of valid targets
        valid_figures = {f['number'] for f in self.figures}
        valid_tables = {t['number'] for t in self.tables}
        valid_sections = set()

        for h in self.headings:
            if h.number:
                valid_sections.add(h.number.rstrip('.'))

        # Check each reference
        for ref in self.cross_references:
            is_valid = True

            if ref.ref_type == 'figure' and ref.target:
                is_valid = ref.target in valid_figures
            elif ref.ref_type == 'table' and ref.target:
                is_valid = ref.target in valid_tables
            elif ref.ref_type == 'section' and ref.target:
                # Normalize and check
                normalized = ref.target.rstrip('.')
                is_valid = normalized in valid_sections or any(
                    normalized in s for s in valid_sections
                )

            ref.is_valid = is_valid

            if not is_valid:
                self.issues.append(StructureIssue(
                    issue_type=StructureIssueType.BROKEN_REFERENCE,
                    message=f'Potentially broken {ref.ref_type} reference',
                    location=f'"{ref.ref_text}" - target "{ref.target}" not found',
                    severity='warning',
                    suggestion=f'Verify that {ref.ref_type} {ref.target} exists in the document'
                ))

    def _check_toc_matches(self):
        """Check that TOC entries match actual headings."""
        if not self.toc_entries:
            return

        # Build heading text set
        heading_texts = {h.text.lower().strip() for h in self.headings}

        for entry in self.toc_entries:
            entry_text = entry['text'].lower().strip()
            # Remove common prefixes like numbers
            cleaned = self._clean_heading_text(entry_text)

            if cleaned and cleaned.lower() not in heading_texts:
                # Check for partial match
                matches = [h for h in heading_texts if cleaned.lower() in h or h in cleaned.lower()]
                if not matches:
                    self.issues.append(StructureIssue(
                        issue_type=StructureIssueType.TOC_MISMATCH,
                        message='TOC entry may not match document heading',
                        location=f'TOC: "{entry["text"]}"',
                        severity='warning',
                        suggestion='Update TOC to reflect current document headings'
                    ))

    def _check_figure_table_references(self):
        """Check that figures and tables are referenced in text."""
        # Get all referenced figures/tables
        referenced_figures = {ref.target for ref in self.cross_references if ref.ref_type == 'figure' and ref.target}
        referenced_tables = {ref.target for ref in self.cross_references if ref.ref_type == 'table' and ref.target}

        # Check figures
        for fig in self.figures:
            if fig['number'] not in referenced_figures:
                self.issues.append(StructureIssue(
                    issue_type=StructureIssueType.MISSING_FIGURE_CAPTION,
                    message='Figure not referenced in text',
                    location=f'Figure {fig["number"]}: {fig["caption"][:50]}...',
                    severity='info',
                    suggestion='Add a reference to this figure in the document text'
                ))

        # Check tables
        for tbl in self.tables:
            if tbl['number'] not in referenced_tables:
                self.issues.append(StructureIssue(
                    issue_type=StructureIssueType.MISSING_TABLE_CAPTION,
                    message='Table not referenced in text',
                    location=f'Table {tbl["number"]}: {tbl["caption"][:50]}...',
                    severity='info',
                    suggestion='Add a reference to this table in the document text'
                ))

    def _build_results(self) -> Dict[str, Any]:
        """Build the final results dictionary."""
        return {
            'version': self.VERSION,
            'headings': [h.to_dict() for h in self.headings],
            'heading_count': len(self.headings),
            'heading_tree': self._build_heading_tree(),
            'cross_references': [r.to_dict() for r in self.cross_references],
            'cross_reference_count': len(self.cross_references),
            'broken_references': len([r for r in self.cross_references if not r.is_valid]),
            'figures': self.figures,
            'figure_count': len(self.figures),
            'tables': self.tables,
            'table_count': len(self.tables),
            'toc_entries': self.toc_entries,
            'toc_count': len(self.toc_entries),
            'bookmarks': list(self.bookmarks),
            'bookmark_count': len(self.bookmarks),
            'issues': [i.to_dict() for i in self.issues],
            'issue_count': len(self.issues),
            'statistics': self._calculate_statistics(),
            'outline': self._generate_outline()
        }

    def _build_heading_tree(self) -> List[Dict[str, Any]]:
        """Build hierarchical heading tree."""
        if not self.headings:
            return []

        root_nodes = []
        stack = []

        for heading in self.headings:
            node = Heading(
                text=heading.text,
                level=heading.level,
                number=heading.number,
                page=heading.page,
                paragraph_index=heading.paragraph_index,
                style_name=heading.style_name
            )

            # Find parent
            while stack and stack[-1].level >= heading.level:
                stack.pop()

            if stack:
                stack[-1].children.append(node)
            else:
                root_nodes.append(node)

            stack.append(node)

        return [n.to_dict() for n in root_nodes]

    def _calculate_statistics(self) -> Dict[str, Any]:
        """Calculate structure statistics."""
        level_counts = defaultdict(int)
        for h in self.headings:
            level_counts[h.level] += 1

        max_depth = max(level_counts.keys()) if level_counts else 0

        return {
            'max_heading_depth': max_depth,
            'headings_by_level': dict(level_counts),
            'avg_section_length': self._calculate_avg_section_length(),
            'has_toc': len(self.toc_entries) > 0,
            'reference_density': len(self.cross_references) / len(self.headings) if self.headings else 0,
            'issues_by_type': self._count_issues_by_type()
        }

    def _calculate_avg_section_length(self) -> float:
        """Calculate average number of paragraphs between headings."""
        if len(self.headings) < 2:
            return 0

        gaps = []
        for i in range(1, len(self.headings)):
            gap = self.headings[i].paragraph_index - self.headings[i-1].paragraph_index
            gaps.append(gap)

        return round(sum(gaps) / len(gaps), 1) if gaps else 0

    def _count_issues_by_type(self) -> Dict[str, int]:
        """Count issues by type."""
        counts = defaultdict(int)
        for issue in self.issues:
            counts[issue.issue_type.value] += 1
        return dict(counts)

    def _generate_outline(self) -> str:
        """Generate a text outline of the document."""
        lines = []

        for heading in self.headings:
            indent = "  " * (heading.level - 1)
            number_str = f"{heading.number} " if heading.number else ""
            lines.append(f"{indent}{number_str}{heading.text}")

        return "\n".join(lines)

    # ==========================================================================
    # PUBLIC UTILITY METHODS
    # ==========================================================================

    def get_heading_tree(self) -> List[Dict[str, Any]]:
        """Get the heading tree structure."""
        return self._build_heading_tree()

    def get_outline(self) -> str:
        """Get text outline of document structure."""
        return self._generate_outline()

    def get_issues(self, severity: Optional[str] = None) -> List[Dict[str, Any]]:
        """
        Get structure issues, optionally filtered by severity.

        Args:
            severity: Filter by 'error', 'warning', or 'info'

        Returns:
            List of issue dictionaries
        """
        issues = self.issues
        if severity:
            issues = [i for i in issues if i.severity == severity]
        return [i.to_dict() for i in issues]

    def get_broken_references(self) -> List[Dict[str, Any]]:
        """Get list of broken cross-references."""
        return [r.to_dict() for r in self.cross_references if not r.is_valid]

    def validate_structure(self) -> bool:
        """
        Check if document structure is valid.

        Returns:
            True if no errors found
        """
        errors = [i for i in self.issues if i.severity == 'error']
        return len(errors) == 0


# =============================================================================
# CONVENIENCE FUNCTIONS
# =============================================================================

def analyze_document_structure(filepath: str) -> Dict[str, Any]:
    """
    Quick function to analyze document structure.

    Args:
        filepath: Path to DOCX file

    Returns:
        Structure analysis results
    """
    analyzer = StructureAnalyzer()
    return analyzer.analyze_docx(filepath)


def get_document_outline(filepath: str) -> str:
    """
    Get a text outline of the document.

    Args:
        filepath: Path to DOCX file

    Returns:
        Text outline string
    """
    analyzer = StructureAnalyzer()
    analyzer.analyze_docx(filepath)
    return analyzer.get_outline()


def validate_document_structure(filepath: str) -> Tuple[bool, List[Dict[str, Any]]]:
    """
    Validate document structure.

    Args:
        filepath: Path to DOCX file

    Returns:
        Tuple of (is_valid, list_of_issues)
    """
    analyzer = StructureAnalyzer()
    analyzer.analyze_docx(filepath)
    return analyzer.validate_structure(), analyzer.get_issues()


# =============================================================================
# CLI INTERFACE
# =============================================================================

if __name__ == '__main__':
    import sys

    if len(sys.argv) > 1:
        filepath = sys.argv[1]
        print(f"Analyzing: {filepath}\n")

        results = analyze_document_structure(filepath)

        if 'error' in results:
            print(f"Error: {results['error']}")
            sys.exit(1)

        print("=== Document Outline ===\n")
        print(results['outline'])

        print(f"\n=== Statistics ===")
        print(f"Headings: {results['heading_count']}")
        print(f"Max depth: {results['statistics']['max_heading_depth']}")
        print(f"Figures: {results['figure_count']}")
        print(f"Tables: {results['table_count']}")
        print(f"Cross-references: {results['cross_reference_count']}")
        print(f"Broken references: {results['broken_references']}")

        if results['issues']:
            print(f"\n=== Issues ({results['issue_count']}) ===")
            for issue in results['issues'][:10]:
                print(f"[{issue['severity'].upper()}] {issue['message']}")
                if issue['location']:
                    print(f"  Location: {issue['location']}")
                if issue['suggestion']:
                    print(f"  Suggestion: {issue['suggestion']}")
    else:
        print("Usage: python structure_analyzer.py <document.docx>")
        print("\nProvides document structure analysis including:")
        print("  - Heading hierarchy")
        print("  - Cross-reference validation")
        print("  - Figure/table tracking")
        print("  - Structure issue detection")
