"""
Statement Forge Export Functions
================================
Export statements to various formats:
- TIBCO Nimbus CSV (CRITICAL: format must be preserved exactly)
- Excel (.xlsx)
- JSON
- Word (.docx)

v2.9.3 F10: Additional export formats
v3.0.30: Flat mode imports (no package directory)
"""

import csv
import json
from datetime import datetime
from pathlib import Path
from typing import List, Dict, Any, Optional
from io import StringIO, BytesIO

# v3.0.49: Support both package and flat import layouts
try:
    from statement_forge.models import Statement
except ImportError:
    from statement_forge__models import Statement


def export_to_nimbus_csv(statements: List[Statement], 
                          filepath: str = None,
                          return_content: bool = False) -> Dict[str, Any]:
    """
    Export statements to TIBCO Nimbus CSV format.
    
    CRITICAL: This format is used in production systems. DO NOT MODIFY.
    
    Format:
    - 12 columns: Level 1-6 with descriptions
    - UTF-8 with BOM (utf-8-sig encoding)
    - Role on own line with colon, blank line after
    """
    try:
        headers = [
            'Level 1', 'Level 1 Description',
            'Level 2', 'Level 2 Description',
            'Level 3', 'Level 3 Description',
            'Level 4', 'Level 4 Description',
            'Level 5', 'Level 5 Description',
            'Level 6', 'Level 6 Description'
        ]
        
        rows = []

        for stmt in statements:
            if not stmt.title and not stmt.description:
                continue

            row = [''] * 12
            level = max(1, min(6, stmt.level))
            level_idx = (level - 1) * 2

            # Build title column - handle cases where title equals number
            number = stmt.number.strip() if stmt.number else ""
            title = stmt.title.strip() if stmt.title else ""

            if not title or title == number:
                row[level_idx] = number if number else title
            elif number:
                row[level_idx] = f"{number}  {title}"
            else:
                row[level_idx] = title

            # Description column - role on own line, blank line, then statement
            description_parts = []
            if stmt.role:
                description_parts.append(stmt.role + ":")
                description_parts.append("")  # Blank line after role
            if stmt.description:
                description_parts.append(stmt.description)
            if stmt.notes:
                description_parts.append(' '.join(stmt.notes))
            row[level_idx + 1] = '\n'.join(description_parts)

            rows.append(row)
        
        if return_content:
            output = StringIO()
            writer = csv.writer(output, quoting=csv.QUOTE_MINIMAL)
            writer.writerow(headers)
            writer.writerows(rows)
            return {'success': True, 'content': output.getvalue(), 'count': len(statements), 'format': 'nimbus_csv'}
        else:
            with open(filepath, 'w', newline='', encoding='utf-8-sig') as f:
                writer = csv.writer(f, quoting=csv.QUOTE_MINIMAL)
                writer.writerow(headers)
                writer.writerows(rows)
            return {'success': True, 'filepath': filepath, 'count': len(statements), 'format': 'nimbus_csv'}
            
    except Exception as e:
        return {'success': False, 'error': str(e), 'format': 'nimbus_csv'}


def export_to_excel(statements: List[Statement],
                     filepath: str = None,
                     source_document: str = "",
                     return_bytes: bool = False) -> Dict[str, Any]:
    """Export statements to Excel with formatting and summary sheet."""
    try:
        import openpyxl
        from openpyxl.styles import Font, PatternFill, Alignment
    except ImportError:
        return {'success': False, 'error': 'openpyxl required for Excel export', 'format': 'xlsx'}
    
    try:
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "Statements"
        
        headers = ['Level 1', 'Level 1 Description', 'Level 2', 'Level 2 Description',
                   'Level 3', 'Level 3 Description', 'Level 4', 'Level 4 Description',
                   'Level 5', 'Level 5 Description', 'Level 6', 'Level 6 Description']
        
        header_font = Font(bold=True, color='FFFFFF')
        header_fill = PatternFill(start_color='1a1a2e', end_color='1a1a2e', fill_type='solid')
        
        for col, header in enumerate(headers, 1):
            cell = ws.cell(row=1, column=col, value=header)
            cell.font = header_font
            cell.fill = header_fill
        
        row_num = 2

        for stmt in statements:
            if not stmt.title and not stmt.description:
                continue

            level = max(1, min(6, stmt.level))
            level_idx = (level - 1) * 2 + 1

            # Build title column - handle cases where title equals number
            number = stmt.number.strip() if stmt.number else ""
            title = stmt.title.strip() if stmt.title else ""

            if not title or title == number:
                title_value = number if number else title
            elif number:
                title_value = f"{number}  {title}"
            else:
                title_value = title

            ws.cell(row=row_num, column=level_idx, value=title_value)

            # Description column - role on own line, blank line, then statement
            description_parts = []
            if stmt.role:
                description_parts.append(stmt.role + ":")
                description_parts.append("")  # Blank line after role
            if stmt.description:
                description_parts.append(stmt.description)
            if stmt.notes:
                description_parts.append(' '.join(stmt.notes))
            ws.cell(row=row_num, column=level_idx + 1, value='\n'.join(description_parts))
            row_num += 1
        
        # Summary sheet
        ws_summary = wb.create_sheet(title="Summary")
        directive_counts = {'shall': 0, 'must': 0, 'will': 0, 'should': 0, 'may': 0}
        for stmt in statements:
            if stmt.directive and stmt.directive.lower() in directive_counts:
                directive_counts[stmt.directive.lower()] += 1
        
        summary_data = [
            ['Source Document', source_document or 'Unknown'],
            ['Export Date', datetime.now().strftime('%Y-%m-%d %H:%M:%S')],
            ['Total Statements', len(statements)],
            ['Shall', directive_counts['shall']],
            ['Must', directive_counts['must']],
            ['Will', directive_counts['will']],
        ]
        for r, data in enumerate(summary_data, 1):
            for c, val in enumerate(data, 1):
                ws_summary.cell(row=r, column=c, value=val)
        
        if return_bytes:
            output = BytesIO()
            wb.save(output)
            output.seek(0)
            return {'success': True, 'content': output.getvalue(), 'count': len(statements), 'format': 'xlsx'}
        else:
            wb.save(filepath)
            return {'success': True, 'filepath': filepath, 'count': len(statements), 'format': 'xlsx'}
            
    except Exception as e:
        return {'success': False, 'error': str(e), 'format': 'xlsx'}


def export_to_json(statements: List[Statement],
                    filepath: str = None,
                    source_document: str = "",
                    return_content: bool = False) -> Dict[str, Any]:
    """Export statements to JSON with full metadata."""
    try:
        directive_counts = {'shall': 0, 'must': 0, 'will': 0, 'should': 0, 'may': 0}
        section_count = sum(1 for s in statements if s.is_header)
        for stmt in statements:
            if stmt.directive and stmt.directive.lower() in directive_counts:
                directive_counts[stmt.directive.lower()] += 1
        
        export_data = {
            'metadata': {
                'source_document': source_document or 'Unknown',
                'extracted_at': datetime.now().isoformat(),
                'statement_count': len(statements),
                'section_count': section_count,
                'directive_counts': directive_counts
            },
            'statements': [stmt.to_dict() for stmt in statements]
        }
        
        if return_content:
            return {'success': True, 'content': json.dumps(export_data, indent=2), 'count': len(statements), 'format': 'json'}
        else:
            with open(filepath, 'w', encoding='utf-8') as f:
                json.dump(export_data, f, indent=2)
            return {'success': True, 'filepath': filepath, 'count': len(statements), 'format': 'json'}
            
    except Exception as e:
        return {'success': False, 'error': str(e), 'format': 'json'}


def export_to_word(statements: List[Statement],
                    filepath: str = None,
                    source_document: str = "",
                    return_bytes: bool = False) -> Dict[str, Any]:
    """Export statements to Word document with hierarchical structure."""
    try:
        from docx import Document
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        return {'success': False, 'error': 'python-docx required for Word export', 'format': 'docx'}
    
    try:
        doc = Document()
        
        title = doc.add_heading('Statement Extraction Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph(f"Source: {source_document or 'Unknown'}")
        doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        doc.add_paragraph(f"Total Statements: {len(statements)}")
        doc.add_page_break()
        
        doc.add_heading('Extracted Statements', level=1)
        
        current_role = ""
        for stmt in statements:
            if not stmt.title and not stmt.description:
                continue
            
            if stmt.role and stmt.role != current_role:
                doc.add_heading(stmt.role, level=2)
                current_role = stmt.role
            
            if stmt.is_header:
                level = min(stmt.level + 1, 9)
                heading_text = f"{stmt.number} {stmt.title}".strip() if stmt.number else stmt.title
                doc.add_heading(heading_text, level=level)
            else:
                p = doc.add_paragraph()
                if stmt.number:
                    run = p.add_run(f"{stmt.number} ")
                    run.bold = True
                if stmt.title:
                    run = p.add_run(f"{stmt.title}: ")
                    run.italic = True
                if stmt.description:
                    p.add_run(stmt.description)
                if stmt.notes:
                    run = p.add_run(f" [NOTE: {'; '.join(stmt.notes)}]")
                    run.italic = True
                if stmt.directive:
                    run = p.add_run(f" [{stmt.directive.upper()}]")
                    run.bold = True
        
        if return_bytes:
            output = BytesIO()
            doc.save(output)
            output.seek(0)
            return {'success': True, 'content': output.getvalue(), 'count': len(statements), 'format': 'docx'}
        else:
            doc.save(filepath)
            return {'success': True, 'filepath': filepath, 'count': len(statements), 'format': 'docx'}
            
    except Exception as e:
        return {'success': False, 'error': str(e), 'format': 'docx'}


def validate_csv_format(filepath: str) -> Dict[str, Any]:
    """
    Validate exported CSV matches TIBCO Nimbus format.

    Returns:
        dict: {'valid': bool, 'message': str}
    """
    try:
        with open(filepath, 'r', encoding='utf-8-sig') as f:
            reader = csv.reader(f)
            header = next(reader)

            expected = []
            for i in range(1, 7):
                expected.extend([f'Level {i}', f'Level {i} Description'])

            if header != expected:
                return {'valid': False, 'message': f"Invalid header format. Expected: {expected[:4]}..."}

            row_count = sum(1 for _ in reader)
            if row_count == 0:
                return {'valid': False, 'message': "CSV has no data rows."}

            return {'valid': True, 'message': f"Valid TIBCO Nimbus format with {row_count} statements."}

    except Exception as e:
        return {'valid': False, 'message': f"Validation error: {e}"}


def get_export_stats(statements: List[Statement], text: str = None) -> Dict[str, Any]:
    """
    Get statistics about statements for export summary.

    v3.5.0: Added document_type classification and guidance_vs_requirement analysis.
    """
    directive_counts = {'shall': 0, 'must': 0, 'will': 0, 'should': 0, 'may': 0}
    section_count = sum(1 for s in statements if s.is_header)
    roles_seen = set()

    for stmt in statements:
        if stmt.directive and stmt.directive.lower() in directive_counts:
            directive_counts[stmt.directive.lower()] += 1
        if stmt.role:
            roles_seen.add(stmt.role)

    # v3.5.0: Classify document type based on directive distribution
    shall_must = directive_counts.get('shall', 0) + directive_counts.get('must', 0)
    should_may = directive_counts.get('should', 0) + directive_counts.get('may', 0)
    will_count = directive_counts.get('will', 0)
    total_directives = shall_must + should_may + will_count

    if total_directives == 0:
        doc_classification = 'informational'
        doc_classification_detail = 'No directive statements found — informational/descriptive document'
    elif shall_must > should_may and shall_must > will_count:
        doc_classification = 'requirements'
        doc_classification_detail = (
            f'Requirements document — {shall_must} mandatory (shall/must) vs '
            f'{should_may} guidance (should/may) statements'
        )
    elif should_may > shall_must:
        doc_classification = 'guidance'
        doc_classification_detail = (
            f'Guidance/handbook — {should_may} guidance (should/may) vs '
            f'{shall_must} mandatory (shall/must) statements. '
            f'"Should" statements are recommendations, not requirements.'
        )
    else:
        doc_classification = 'descriptive'
        doc_classification_detail = (
            f'Descriptive/process document — {will_count} descriptive (will) vs '
            f'{shall_must} mandatory and {should_may} guidance statements'
        )

    return {
        'total_statements': len(statements),
        'section_count': section_count,
        'role_count': len(roles_seen),
        'roles': list(roles_seen),
        'directive_counts': directive_counts,
        'document_classification': doc_classification,
        'document_classification_detail': doc_classification_detail,
    }
